"""
Phase 6C — safe, local Evidence Text Extraction (read-only preview).

Answers a single technical question: *what readable text can we safely extract
from this evidence file?* It is preparation for later AI Evidence Analyzer /
Rule Engine phases — it does NOT analyze evidence sufficiency, does NOT decide
compliance, and does NOT run OCR or AI.

Guarantees:
* No AI, no network, no external services.
* No OCR (scanned images / image-only PDFs return `no_text_extracted`).
* Read-only — never writes, never executes file content / macros / formulas.
* Hard size limit, character cap, whitespace normalization, and defensive
  error handling (parser errors become a status + warning, never a traceback or
  a leaked file path).

Uses only already-installed, safe parsers: pdfplumber (PDF text layer),
python-docx (DOCX), openpyxl (XLSX), and plain UTF-8 reads (TXT/CSV/MD).
"""
import os
import re
from dataclasses import dataclass, field

# Limits.
MAX_EXTRACTION_BYTES = 25 * 1024 * 1024   # 25 MB cap for extraction work
MAX_TEXT_CHARS = 50_000                    # cap stored/previewed text
MAX_PDF_PAGES = 300
MAX_XLSX_CELLS = 20_000

EXTRACTABLE_EXTS = {'pdf', 'docx', 'xlsx', 'txt', 'csv', 'md'}
IMAGE_EXTS = {'png', 'jpg', 'jpeg', 'tiff', 'tif', 'bmp', 'gif', 'webp'}

# Statuses.
EXTRACTED = 'extracted'
NO_TEXT = 'no_text_extracted'
UNSUPPORTED = 'unsupported_type'
TOO_LARGE = 'too_large'
FAILED = 'failed'

STATUS_AR = {
    EXTRACTED: 'تم الاستخراج',
    NO_TEXT: 'تعذّر استخراج نص كافٍ',
    UNSUPPORTED: 'نوع ملف غير مدعوم',
    TOO_LARGE: 'الملف كبير جدًا',
    FAILED: 'فشل الاستخراج',
}


@dataclass(frozen=True)
class ExtractionResult:
    status: str
    extracted_text: str = ''
    char_count: int = 0
    page_count: object = None
    extraction_method: str = ''
    warnings: list = field(default_factory=list)
    error_message: object = None
    truncated: bool = False

    @property
    def status_ar(self):
        return STATUS_AR.get(self.status, self.status)

    @property
    def has_text(self):
        return self.status == EXTRACTED and self.char_count > 0


def _safe_ext(filename):
    """Extension from the declared filename only (never from a server path)."""
    name = (filename or '').strip().lower()
    if '.' not in name:
        return ''
    return name.rsplit('.', 1)[1]


def _normalize(text):
    """Collapse runs of whitespace, trim, and cap length. Returns (text, truncated)."""
    if not text:
        return '', False
    text = text.replace('\x00', ' ')
    text = re.sub(r'[ \t ]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = '\n'.join(line.strip() for line in text.split('\n')).strip()
    if len(text) > MAX_TEXT_CHARS:
        return text[:MAX_TEXT_CHARS], True
    return text, False


def _read_text_file(path):
    with open(path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read(MAX_TEXT_CHARS + 1)


def _read_pdf(path):
    import pdfplumber  # lazy; text layer only, no OCR
    parts, pages = [], 0
    warnings = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            if i >= MAX_PDF_PAGES:
                warnings.append(f'تم قصر الاستخراج على أول {MAX_PDF_PAGES} صفحة.')
                break
            pages += 1
            parts.append(page.extract_text() or '')
    return '\n'.join(parts), pages, warnings


def _read_docx(path):
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text]
            if cells:
                parts.append(' | '.join(cells))
    return '\n'.join(parts)


def _read_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    parts, cells = [], 0
    warnings = []
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                values = [str(v) for v in row if v not in (None, '')]
                if values:
                    parts.append(' | '.join(values))
                cells += len(row)
                if cells >= MAX_XLSX_CELLS:
                    warnings.append('تم قصر الاستخراج على جزء من الخلايا.')
                    raise StopIteration
    except StopIteration:
        pass
    finally:
        wb.close()
    return '\n'.join(parts), warnings


def extract_text_from_file(file_path, filename, content_type=None) -> ExtractionResult:
    """Safely extract readable text from a single file. Pure, deterministic, read-only."""
    ext = _safe_ext(filename)

    if not file_path or not os.path.exists(file_path):
        return ExtractionResult(status=FAILED, extraction_method='none',
                                error_message='الملف غير متوفر.',
                                warnings=['تعذّر العثور على الملف.'])
    try:
        size = os.path.getsize(file_path)
    except OSError:
        return ExtractionResult(status=FAILED, extraction_method='none',
                                error_message='تعذّر قراءة الملف.')
    if size > MAX_EXTRACTION_BYTES:
        return ExtractionResult(status=TOO_LARGE, extraction_method='none',
                                warnings=['حجم الملف يتجاوز حد الاستخراج.'])

    if ext in IMAGE_EXTS:
        return ExtractionResult(status=NO_TEXT, extraction_method='none',
                                warnings=['ملف صورة — استخراج النص بالـ OCR مُخطّط لمرحلة لاحقة.'])
    if ext not in EXTRACTABLE_EXTS:
        return ExtractionResult(status=UNSUPPORTED, extraction_method='none',
                                warnings=['نوع الملف غير مدعوم للاستخراج النصّي.'])

    method = ext
    page_count = None
    warnings = []
    try:
        if ext in ('txt', 'csv', 'md'):
            raw = _read_text_file(file_path)
            method = 'plain_text'
        elif ext == 'pdf':
            raw, page_count, warnings = _read_pdf(file_path)
            method = 'pdf_text_layer'
        elif ext == 'docx':
            raw = _read_docx(file_path)
            method = 'docx'
        elif ext == 'xlsx':
            raw, warnings = _read_xlsx(file_path)
            method = 'xlsx'
        else:  # pragma: no cover - guarded above
            return ExtractionResult(status=UNSUPPORTED, extraction_method='none')
    except ImportError:
        return ExtractionResult(status=UNSUPPORTED, extraction_method=method,
                                warnings=['مكتبة الاستخراج غير متوفّرة لهذا النوع.'],
                                error_message='مكتبة الاستخراج غير متوفّرة.')
    except Exception:
        # Never leak a traceback or path to the user.
        return ExtractionResult(status=FAILED, extraction_method=method,
                                warnings=['تعذّر تحليل محتوى الملف.'],
                                error_message='حدث خطأ أثناء قراءة محتوى الملف.')

    text, truncated = _normalize(raw)
    if truncated:
        warnings = list(warnings) + ['تم اقتصار النص المستخرج على الحد الأقصى للأحرف.']
    if not text:
        if ext == 'pdf':
            warnings = list(warnings) + ['قد يكون المستند ممسوحًا ضوئيًا (صورة) — OCR مُخطّط لمرحلة لاحقة.']
        return ExtractionResult(status=NO_TEXT, extraction_method=method,
                                page_count=page_count, warnings=warnings)

    return ExtractionResult(status=EXTRACTED, extracted_text=text, char_count=len(text),
                            page_count=page_count, extraction_method=method,
                            warnings=warnings, truncated=truncated)


def extract_text_from_evidence(evidence) -> ExtractionResult:
    """Extract text from an EvidenceSubmission (or legacy Evidence) instance (read-only)."""
    file_field = getattr(evidence, 'uploaded_file', None) or getattr(evidence, 'file', None)
    filename = getattr(evidence, 'original_filename', '') or getattr(file_field, 'name', '')
    if file_field is None:
        return ExtractionResult(status=FAILED, extraction_method='none',
                                error_message='لا يوجد ملف مرتبط بهذا الدليل.')
    try:
        path = file_field.path
    except Exception:
        return ExtractionResult(status=FAILED, extraction_method='none',
                                error_message='تعذّر الوصول إلى الملف.')
    return extract_text_from_file(path, filename)
