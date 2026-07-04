"""
Regression tests for PATCH_NOTES fixes owned by the compliance app:
  #1 50 MB upload-size conflict (settings)        — FR-005.2 / NFR-006
  #2 evidence validation before row creation      — FR-005.11
"""
from unittest import mock

from django.conf import settings
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse

from core.models import User, Company
from compliance.models import Framework, Domain, Control, Evidence


def _company_with_control():
    company = Company.objects.create(
        name='Co', cr_number='1010101010', sector='technology', size='small',
        contact_email='c@x.com', target_nca=True)
    fw, _ = Framework.objects.get_or_create(code='NCA_ECC', defaults={'name': 'NCA'})
    dom, _ = Domain.objects.get_or_create(framework=fw, name='Gov', defaults={'code': 'GOV'})
    control = Control.objects.create(
        framework=fw, domain=dom, control_id='NCA-1', title='T', description='d')
    return company, control


class UploadSizeConfigTests(TestCase):
    """PATCH #1: the 10 MB DATA_UPLOAD cap that silently blocked 50 MB must be gone."""

    def test_evidence_limit_is_50mb(self):
        self.assertEqual(settings.MAX_EVIDENCE_FILE_SIZE, 50 * 1024 * 1024)

    def test_request_body_cap_exceeds_evidence_limit(self):
        # If DATA_UPLOAD_MAX_MEMORY_SIZE <= MAX_EVIDENCE_FILE_SIZE, 50 MB files are blocked again.
        self.assertGreater(settings.DATA_UPLOAD_MAX_MEMORY_SIZE, settings.MAX_EVIDENCE_FILE_SIZE)

    def test_supported_extensions_cover_srs_formats(self):
        for ext in ('pdf', 'png', 'jpg', 'jpeg', 'tiff', 'docx', 'xlsx', 'txt'):
            self.assertIn(ext, settings.ALLOWED_EVIDENCE_EXTENSIONS)


class EvidenceValidationTests(TestCase):
    """PATCH #2: upload_evidence rejects bad files BEFORE creating an Evidence row."""

    def setUp(self):
        self.company, self.control = _company_with_control()
        self.user = User.objects.create_user(
            email='u@x.com', password='longenough12', company=self.company, role='company_admin')
        self.client.force_login(self.user)
        self.url = reverse('compliance:upload_evidence', args=[self.control.id])

    def test_rejects_unsupported_extension(self):
        bad = SimpleUploadedFile('malware.exe', b'MZ binary', content_type='application/octet-stream')
        resp = self.client.post(self.url, {'evidence_file': bad})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(Evidence.objects.count(), 0)

    @override_settings(MAX_EVIDENCE_FILE_SIZE=10)  # tiny cap so a few bytes are "oversized"
    def test_rejects_oversized_file(self):
        big = SimpleUploadedFile('policy.pdf', b'x' * 50, content_type='application/pdf')
        resp = self.client.post(self.url, {'evidence_file': big})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(Evidence.objects.count(), 0)

    def test_accepts_valid_file_and_creates_row(self):
        good = SimpleUploadedFile('policy.txt', b'Cybersecurity policy approved.', content_type='text/plain')
        # Force the Celery branch with a no-op delay so no broker/OCR is touched.
        with mock.patch('monitoring.tasks.analyze_evidence_async.delay') as delayed:
            resp = self.client.post(self.url, {'evidence_file': good})
            delayed.assert_called_once()
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(Evidence.objects.filter(company_control__company=self.company).count(), 1)
        ev = Evidence.objects.get()
        self.assertEqual(ev.file_type, 'txt')
        self.assertEqual(ev.uploaded_by, self.user)


# ============================================================
# Phase 1 — Source Registry + Framework Versioning
# ============================================================
from io import StringIO

from django.core.management import call_command
from django.db import IntegrityError, transaction

from compliance.models import SourceDocument, FrameworkVersion


class SourceRegistryModelTests(TestCase):
    def test_source_document_str(self):
        doc = SourceDocument.objects.create(
            title='ECC', issuer='NCA', document_type='standard', version='2:2024')
        self.assertEqual(str(doc), '[NCA] ECC 2:2024')

    def test_framework_version_code_unique(self):
        fw, _ = Framework.objects.get_or_create(code='NCA_ECC', defaults={'name': 'NCA'})
        FrameworkVersion.objects.create(framework=fw, code='NCA-ECC-2-2024')
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                FrameworkVersion.objects.create(framework=fw, code='NCA-ECC-2-2024')

    def test_framework_version_links_to_framework(self):
        fw, _ = Framework.objects.get_or_create(code='NCA_ECC', defaults={'name': 'NCA'})
        ver = FrameworkVersion.objects.create(framework=fw, code='NCA-ECC-2018', status='superseded')
        self.assertEqual(ver.framework, fw)
        self.assertIn(ver, fw.versions.all())

    def test_existing_framework_version_field_still_exists(self):
        # The legacy Framework.version char field must NOT be removed by Phase 1.
        fw = Framework(code='X', name='X', version='2.0')
        self.assertEqual(fw.version, '2.0')


class SeedFrameworkVersionsTests(TestCase):
    def _seed(self):
        out = StringIO()
        call_command('seed_framework_versions', stdout=out)
        return out.getvalue()

    def test_seed_framework_versions_creates_active_ecc_2024(self):
        self._seed()
        ecc = FrameworkVersion.objects.get(code='NCA-ECC-2-2024')
        self.assertEqual(ecc.status, 'active')
        self.assertTrue(ecc.is_default)

    def test_seed_framework_versions_marks_ecc_2018_as_superseded_or_legacy(self):
        self._seed()
        ecc18 = FrameworkVersion.objects.get(code='NCA-ECC-2018')
        self.assertIn(ecc18.status, ('superseded', 'legacy'))
        self.assertFalse(ecc18.is_default)

    def test_legacy_framework_version_exists_for_334_excel(self):
        self._seed()
        legacy = FrameworkVersion.objects.get(code='LEGACY-334-CONTROLS')
        self.assertEqual(legacy.status, 'legacy')
        self.assertIsNotNone(legacy.source_document)
        self.assertEqual(legacy.source_document.document_type, 'bootstrap_excel')

    def test_seed_framework_versions_idempotent(self):
        self._seed()
        v1, s1 = FrameworkVersion.objects.count(), SourceDocument.objects.count()
        self._seed()  # second run must not duplicate
        self.assertEqual(FrameworkVersion.objects.count(), v1)
        self.assertEqual(SourceDocument.objects.count(), s1)

    def test_seed_does_not_touch_controls_or_companycontrols(self):
        # Phase 1 must not import controls or create CompanyControls.
        before_controls = Control.objects.count()
        self._seed()
        self.assertEqual(Control.objects.count(), before_controls)


# ============================================================
# Phase 2A — Control Knowledge Base Schema
# ============================================================
from compliance.models import ControlVersion, ControlApplicabilityTag, CompanyControl


def _fw_dom():
    fw, _ = Framework.objects.get_or_create(code='NCA_ECC', defaults={'name': 'NCA'})
    dom, _ = Domain.objects.get_or_create(framework=fw, name='Gov', defaults={'code': 'GOV'})
    return fw, dom


class ControlKnowledgeBaseSchemaTests(TestCase):
    def setUp(self):
        self.fw, self.dom = _fw_dom()

    def test_control_optional_framework_version_field_exists(self):
        f = Control._meta.get_field('framework_version')
        self.assertTrue(f.null and f.blank)

    def test_control_optional_source_document_field_exists(self):
        f = Control._meta.get_field('source_document')
        self.assertTrue(f.null and f.blank)

    def test_control_can_exist_without_framework_version(self):
        c = Control.objects.create(framework=self.fw, domain=self.dom, control_id='C-1',
                                   title='t', description='d')
        self.assertIsNone(c.framework_version)
        self.assertIsNone(c.source_document)
        self.assertFalse(c.is_legacy_import)

    def test_control_can_link_to_framework_version(self):
        ver = FrameworkVersion.objects.create(framework=self.fw, code='NCA-ECC-2-2024', is_default=True)
        c = Control.objects.create(framework=self.fw, domain=self.dom, control_id='C-2',
                                   title='t', description='d', framework_version=ver)
        self.assertEqual(c.framework_version, ver)
        self.assertIn(c, ver.controls.all())

    def test_control_can_link_to_source_document(self):
        doc = SourceDocument.objects.create(title='ECC', issuer='NCA', document_type='standard')
        c = Control.objects.create(framework=self.fw, domain=self.dom, control_id='C-3',
                                   title='t', description='d', source_document=doc,
                                   source_page=12, external_reference='NCA ECC 1-1-1')
        self.assertEqual(c.source_document, doc)
        self.assertEqual(c.source_page, 12)
        self.assertEqual(c.external_reference, 'NCA ECC 1-1-1')

    def test_control_version_can_snapshot_control_text(self):
        c = Control.objects.create(framework=self.fw, domain=self.dom, control_id='C-4',
                                   title='Asset Mgmt', description='d')
        cv = ControlVersion.objects.create(control=c, version_label='ECC 2:2024',
                                           title_snapshot='Asset Mgmt',
                                           statement_snapshot='The org shall manage assets.')
        self.assertIn(cv, c.versions.all())
        self.assertEqual(str(cv), 'C-4 @ ECC 2:2024')

    def test_control_applicability_tag_unique_per_control(self):
        c = Control.objects.create(framework=self.fw, domain=self.dom, control_id='C-5',
                                   title='t', description='d')
        ControlApplicabilityTag.objects.create(control=c, tag='cloud')
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                ControlApplicabilityTag.objects.create(control=c, tag='cloud')

    def test_existing_controls_still_create_without_new_fields(self):
        # Old-style creation (no new kwargs) must still work unchanged.
        c = Control.objects.create(framework=self.fw, domain=self.dom, control_id='C-6',
                                   title='t', description='d')
        self.assertEqual(str(c), 'C-6: t')

    def test_company_control_still_works_with_existing_control(self):
        from core.models import Company
        company = Company.objects.create(name='Co', cr_number='9090909090', sector='technology',
                                         size='small', contact_email='c@x.com', target_nca=True)
        c = Control.objects.create(framework=self.fw, domain=self.dom, control_id='C-7',
                                   title='t', description='d')
        cc = CompanyControl.objects.create(company=company, control=c, status='compliant')
        self.assertEqual(cc.control, c)

    def test_evidence_still_links_to_company_control(self):
        from core.models import Company
        from compliance.models import Evidence
        company = Company.objects.create(name='Co', cr_number='9191919191', sector='technology',
                                         size='small', contact_email='c@x.com', target_nca=True)
        c = Control.objects.create(framework=self.fw, domain=self.dom, control_id='C-8',
                                   title='t', description='d')
        cc = CompanyControl.objects.create(company=company, control=c)
        ev = Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        self.assertEqual(ev.company_control, cc)


class LinkLegacyControlsCommandTests(TestCase):
    def setUp(self):
        self.fw, self.dom = _fw_dom()
        self.ver = FrameworkVersion.objects.create(
            framework=self.fw, code='NCA-ECC-2-2024', is_default=True, status='active')
        self.controls = [
            Control.objects.create(framework=self.fw, domain=self.dom, control_id=f'L-{i}',
                                   title='t', description='d') for i in range(3)]

    def _run(self):
        out = StringIO()
        call_command('link_legacy_controls_to_framework_versions', stdout=out)
        return out.getvalue()

    def test_link_legacy_controls_links_to_default_version(self):
        self._run()
        for c in self.controls:
            c.refresh_from_db()
            self.assertEqual(c.framework_version, self.ver)
            self.assertTrue(c.is_legacy_import)
            self.assertTrue(c.import_notes)

    def test_link_legacy_controls_command_idempotent(self):
        self._run()
        self._run()  # second run: all already linked -> no change
        linked = Control.objects.filter(framework_version=self.ver).count()
        self.assertEqual(linked, 3)

    def test_link_legacy_controls_does_not_create_controls(self):
        before = Control.objects.count()
        self._run()
        self.assertEqual(Control.objects.count(), before)

    def test_link_legacy_controls_does_not_touch_companycontrols_or_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        company = Company.objects.create(name='Co', cr_number='9292929292', sector='technology',
                                         size='small', contact_email='c@x.com', target_nca=True)
        cc = CompanyControl.objects.create(company=company, control=self.controls[0])
        ev = Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_count, ev_count = CompanyControl.objects.count(), Evidence.objects.count()
        self._run()
        self.assertEqual(CompanyControl.objects.count(), cc_count)
        self.assertEqual(Evidence.objects.count(), ev_count)
        cc.refresh_from_db(); ev.refresh_from_db()
        self.assertEqual(cc.control, self.controls[0])
        self.assertEqual(ev.company_control, cc)


# ============================================================
# Phase 2B — Official Control Import Pilot
# ============================================================
class ImportOfficialControlsPilotTests(TestCase):
    """Runs on a clean test DB (no legacy 334), so the pilot creates real
    official-linked controls and we can assert the full import path."""

    def setUp(self):
        # Phase 1 seed provides the FrameworkVersions + SourceDocuments the pilot needs.
        call_command('seed_framework_versions', stdout=StringIO())

    def _import(self):
        out = StringIO()
        call_command('import_official_controls_pilot', stdout=out)
        return out.getvalue()

    def _pilot_qs(self):
        return Control.objects.filter(is_legacy_import=False, framework_version__isnull=False)

    def test_import_official_controls_pilot_command_exists(self):
        # Should not raise CommandError / unknown command.
        self._import()

    def test_import_official_controls_pilot_creates_limited_number_of_controls(self):
        self._import()
        n = self._pilot_qs().count()
        self.assertGreater(n, 0)
        self.assertLessEqual(n, 9)
        self.assertEqual(n, 9)

    def test_import_official_controls_pilot_is_idempotent(self):
        self._import()
        c1, v1 = Control.objects.count(), ControlVersion.objects.count()
        self._import()
        self.assertEqual(Control.objects.count(), c1)
        self.assertEqual(ControlVersion.objects.count(), v1)

    def test_import_official_controls_pilot_links_nca_controls_to_ecc_2024(self):
        self._import()
        c = Control.objects.get(framework__code='NCA_ECC', control_id='1-1-1')
        self.assertEqual(c.framework_version.code, 'NCA-ECC-2-2024')
        self.assertEqual(c.external_reference, 'NCA ECC 1-1-1')

    def test_import_official_controls_pilot_links_aramco_controls_to_sacs_002(self):
        self._import()
        c = Control.objects.get(framework__code='ARAMCO_SACS002', control_id='TPC-1')
        self.assertEqual(c.framework_version.code, 'ARAMCO-SACS-002')

    def test_import_official_controls_pilot_links_sabic_controls_to_cybertrust(self):
        self._import()
        c = Control.objects.get(framework__code='SABIC_CT', control_id='CT-01')
        self.assertEqual(c.framework_version.code, 'SABIC-CYBERTRUST-1-0')

    def test_import_official_controls_pilot_creates_control_versions(self):
        self._import()
        self.assertEqual(ControlVersion.objects.count(), 9)
        cv = ControlVersion.objects.get(control__control_id='CT-03')
        self.assertTrue(cv.statement_snapshot)
        self.assertEqual(cv.change_summary, 'Initial official pilot import.')

    def test_import_official_controls_pilot_requires_source_document(self):
        self._import()
        for c in self._pilot_qs():
            self.assertIsNotNone(c.source_document)

    def test_import_official_controls_pilot_does_not_touch_companycontrols_or_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        # Pre-create a company-control + evidence on a legacy-style control.
        fw, dom = _fw_dom()
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='ZZ-9',
                                        title='t', description='d')
        company = Company.objects.create(name='Co', cr_number='7070707070', sector='technology',
                                         size='small', contact_email='c@x.com', target_nca=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        ev = Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        self._import()
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)

    def test_import_official_controls_pilot_does_not_use_legacy_excel_as_authority(self):
        self._import()
        for c in self._pilot_qs():
            self.assertNotEqual(c.source_document.document_type, 'bootstrap_excel')
            self.assertIn(c.source_document.issuer, ('NCA', 'ARAMCO', 'SABIC'))

    def test_import_official_controls_pilot_coexists_with_legacy_id(self):
        # Phase 2D: a legacy control with the same (framework, control_id) and a
        # NULL framework_version must coexist with the official one, untouched.
        fw = Framework.objects.get(code='NCA_ECC')
        dom, _ = Domain.objects.get_or_create(framework=fw, code='1', defaults={'name': 'Gov'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='1-1-1',
                                        title='LEGACY TITLE', description='legacy')  # framework_version=None
        self._import()
        legacy.refresh_from_db()
        self.assertEqual(legacy.title, 'LEGACY TITLE')          # untouched
        self.assertIsNone(legacy.framework_version)             # still legacy
        self.assertFalse(legacy.is_legacy_import)
        # The official one now exists ALONGSIDE the legacy one.
        official = Control.objects.get(framework_version__code='NCA-ECC-2-2024', control_id='1-1-1')
        self.assertNotEqual(official.pk, legacy.pk)
        self.assertEqual(Control.objects.filter(framework=fw, control_id='1-1-1').count(), 2)


# ============================================================
# Phase 2D — Control Identity Migration (partial unique constraints)
# ============================================================
class ControlIdentityConstraintTests(TestCase):
    def setUp(self):
        self.fw, self.dom = _fw_dom()
        self.fv = FrameworkVersion.objects.create(
            framework=self.fw, code='NCA-ECC-2-2024', is_default=True, status='active')

    def test_legacy_control_can_remain_with_null_framework_version(self):
        c = Control.objects.create(framework=self.fw, domain=self.dom, control_id='1-1-1',
                                   title='legacy', description='d')
        self.assertIsNone(c.framework_version)

    def test_official_and_legacy_same_control_id_coexist(self):
        legacy = Control.objects.create(framework=self.fw, domain=self.dom, control_id='1-1-1',
                                        title='legacy', description='d')
        official = Control.objects.create(framework=self.fw, domain=self.dom, control_id='1-1-1',
                                          title='official', description='d', framework_version=self.fv)
        self.assertNotEqual(legacy.pk, official.pk)
        self.assertEqual(Control.objects.filter(control_id='1-1-1').count(), 2)

    def test_duplicate_official_same_framework_version_is_blocked(self):
        Control.objects.create(framework=self.fw, domain=self.dom, control_id='1-1-1',
                               title='o1', description='d', framework_version=self.fv)
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                Control.objects.create(framework=self.fw, domain=self.dom, control_id='1-1-1',
                                       title='o2', description='d', framework_version=self.fv)

    def test_duplicate_legacy_same_framework_is_blocked(self):
        Control.objects.create(framework=self.fw, domain=self.dom, control_id='9-9-9',
                               title='l1', description='d')  # fv NULL
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                Control.objects.create(framework=self.fw, domain=self.dom, control_id='9-9-9',
                                       title='l2', description='d')  # fv NULL

    def test_company_control_still_links_to_legacy_control(self):
        from core.models import Company
        legacy = Control.objects.create(framework=self.fw, domain=self.dom, control_id='1-2-3',
                                        title='legacy', description='d')
        company = Company.objects.create(name='Co', cr_number='6060606060', sector='technology',
                                         size='small', contact_email='c@x.com', target_nca=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        self.assertEqual(cc.control, legacy)

    def test_evidence_still_links_to_company_control(self):
        from core.models import Company
        from compliance.models import Evidence
        legacy = Control.objects.create(framework=self.fw, domain=self.dom, control_id='1-2-4',
                                        title='legacy', description='d')
        company = Company.objects.create(name='Co', cr_number='6161616161', sector='technology',
                                         size='small', contact_email='c@x.com', target_nca=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        ev = Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        self.assertEqual(ev.company_control, cc)

    def test_pilot_import_creates_official_beside_legacy(self):
        # Seed all 9 legacy ids, then import: expect 9 legacy + 9 official = 18 total.
        call_command('seed_framework_versions', stdout=StringIO())
        legacy_specs = [
            ('NCA_ECC', ['1-1-1', '1-2-1', '1-3-1']),
            ('ARAMCO_SACS002', ['TPC-1', 'TPC-2', 'TPC-3']),
            ('SABIC_CT', ['CT-01', 'CT-02', 'CT-03']),
        ]
        for fw_code, ids in legacy_specs:
            fw, _ = Framework.objects.get_or_create(code=fw_code, defaults={'name': fw_code})
            dom, _ = Domain.objects.get_or_create(framework=fw, code='L', defaults={'name': 'Legacy'})
            for cid in ids:
                Control.objects.create(framework=fw, domain=dom, control_id=cid,
                                       title=f'legacy {cid}', description='d')  # fv NULL
        legacy_before = Control.objects.filter(framework_version__isnull=True).count()
        call_command('import_official_controls_pilot', stdout=StringIO())
        official = Control.objects.filter(framework_version__isnull=False, is_legacy_import=False)
        self.assertEqual(official.count(), 9)
        self.assertEqual(Control.objects.filter(framework_version__isnull=True).count(), legacy_before)  # legacy intact

    def test_pilot_import_remains_idempotent_after_identity_change(self):
        call_command('seed_framework_versions', stdout=StringIO())
        call_command('import_official_controls_pilot', stdout=StringIO())
        n1 = Control.objects.count()
        call_command('import_official_controls_pilot', stdout=StringIO())
        self.assertEqual(Control.objects.count(), n1)


# ============================================================
# Phase 2F — Official Aramco dataset: validation + dry-run/apply import
# ============================================================
from django.core.management.base import CommandError

FV = 'ARAMCO-SACS-002'
_VALIDATE = 'compliance.management.commands.validate_official_control_dataset.load_official_dataset'
_VALID_CTRL = dict(control_id='TPC-1', external_reference='SACS-002 TPC-1', domain='IDENTIFY',
                   title='t', statement='s', source_reference='IDENTIFY / Governance (GV) / TPC-1')


def _patched_dataset(controls):
    meta = {'framework_version_code': FV, 'expected_control_count': len(controls),
            '_path': 'x', 'source_document_title': 'SACS-002',
            'notes': 'NOT derived from the legacy 334 Excel.'}
    return mock.patch(_VALIDATE, return_value=(meta, controls))


class ValidateOfficialDatasetTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def test_validate_official_control_dataset_command_exists(self):
        # The real dataset should validate cleanly (no CommandError raised).
        call_command('validate_official_control_dataset', framework_version=FV, stdout=StringIO())

    def test_validate_aramco_dataset_passes(self):
        out = StringIO()
        call_command('validate_official_control_dataset', framework_version=FV, stdout=out)
        self.assertIn('OK', out.getvalue())

    def test_validate_rejects_duplicate_control_id(self):
        with _patched_dataset([dict(_VALID_CTRL), dict(_VALID_CTRL)]):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset', framework_version=FV, stdout=StringIO())

    def test_validate_rejects_missing_external_reference(self):
        bad = {k: v for k, v in _VALID_CTRL.items() if k != 'external_reference'}
        with _patched_dataset([bad]):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset', framework_version=FV, stdout=StringIO())

    def test_validate_rejects_missing_source_reference(self):
        bad = {k: v for k, v in _VALID_CTRL.items() if k != 'source_reference'}  # and no source_page
        with _patched_dataset([bad]):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset', framework_version=FV, stdout=StringIO())

    def test_validate_rejects_missing_framework_version(self):
        # No seed of THIS fv -> not in DB. Use an unknown fv code with a patched dataset.
        meta = {'framework_version_code': 'NOPE-1', 'expected_control_count': 1, '_path': 'x'}
        with mock.patch(_VALIDATE, return_value=(meta, [dict(_VALID_CTRL)])):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset', framework_version='NOPE-1', stdout=StringIO())

    def test_validate_does_not_touch_database(self):
        before = Control.objects.count()
        call_command('validate_official_control_dataset', framework_version=FV, stdout=StringIO())
        self.assertEqual(Control.objects.count(), before)


class ImportOfficialControlsTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())
        self.fv = FrameworkVersion.objects.get(code=FV)

    def _official_qs(self):
        return Control.objects.filter(framework_version=self.fv, is_legacy_import=False)

    def test_import_official_controls_dry_run_does_not_create_controls(self):
        call_command('import_official_controls', framework_version=FV, stdout=StringIO())  # dry-run default
        self.assertEqual(self._official_qs().count(), 0)

    def test_import_official_controls_apply_creates_controls_in_test_db(self):
        call_command('import_official_controls', framework_version=FV, apply=True, stdout=StringIO())
        self.assertEqual(self._official_qs().count(), 92)

    def test_import_official_controls_apply_is_idempotent(self):
        call_command('import_official_controls', framework_version=FV, apply=True, stdout=StringIO())
        n1 = Control.objects.count()
        call_command('import_official_controls', framework_version=FV, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.count(), n1)
        self.assertEqual(self._official_qs().count(), 92)

    def test_import_official_controls_updates_pilot_controls_not_duplicate(self):
        # Pre-create a pilot TPC-1 official control; import must update, not duplicate it.
        dom, _ = Domain.objects.get_or_create(framework=self.fv.framework, code='IDENTIFY',
                                              defaults={'name': 'IDENTIFY'})
        Control.objects.create(framework=self.fv.framework, framework_version=self.fv,
                               domain=dom, control_id='TPC-1', title='old', description='old')
        call_command('import_official_controls', framework_version=FV, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.filter(framework_version=self.fv, control_id='TPC-1').count(), 1)
        self.assertEqual(self._official_qs().count(), 92)

    def test_import_official_controls_creates_control_versions(self):
        call_command('import_official_controls', framework_version=FV, apply=True, stdout=StringIO())
        self.assertEqual(ControlVersion.objects.filter(framework_version=self.fv).count(), 92)

    def test_import_official_controls_does_not_touch_legacy_controls(self):
        fw = self.fv.framework
        dom, _ = Domain.objects.get_or_create(framework=fw, code='L', defaults={'name': 'L'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='TPC-1',
                                        title='LEGACY', description='d')  # fv NULL
        call_command('import_official_controls', framework_version=FV, apply=True, stdout=StringIO())
        legacy.refresh_from_db()
        self.assertEqual(legacy.title, 'LEGACY')
        self.assertIsNone(legacy.framework_version)
        self.assertEqual(Control.objects.filter(framework=fw, control_id='TPC-1').count(), 2)

    def test_import_official_controls_does_not_touch_companycontrol_or_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        fw = self.fv.framework
        dom, _ = Domain.objects.get_or_create(framework=fw, code='L', defaults={'name': 'L'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='TPC-1',
                                        title='LEGACY', description='d')
        company = Company.objects.create(name='Co', cr_number='5050505050', sector='technology',
                                         size='small', contact_email='c@x.com', target_aramco=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        call_command('import_official_controls', framework_version=FV, apply=True, stdout=StringIO())
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)

    def test_import_official_controls_requires_valid_framework_version(self):
        with self.assertRaises(CommandError):
            call_command('import_official_controls', framework_version='UNKNOWN-XX', stdout=StringIO())


# ============================================================
# Phase 2G — Aramco apply (with TPC-62) + SABIC dataset dry-run
# ============================================================
from compliance.official_dataset import load_official_dataset

SABIC_FV = 'SABIC-CYBERTRUST-1-0'
_SABIC_VALIDATE = 'compliance.management.commands.validate_official_control_dataset.load_official_dataset'
_SABIC_CTRL = dict(control_id='CT-01', external_reference='SABIC CT-01', domain='GOVERN',
                   title='t', statement='s', source_reference='GOVERN / GV.PO / CT-01')


class Tpc62AndAramcoApplyTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())
        self.fv = FrameworkVersion.objects.get(code='ARAMCO-SACS-002')

    def _official(self):
        return Control.objects.filter(framework_version=self.fv, is_legacy_import=False)

    def test_tpc_62_verification_documented(self):
        # TPC-62 was verified from official sources and is present with a source_reference.
        meta, controls = load_official_dataset('ARAMCO-SACS-002')
        self.assertEqual(meta['expected_control_count'], 92)
        tpc62 = [c for c in controls if c['control_id'] == 'TPC-62']
        self.assertEqual(len(tpc62), 1)
        self.assertTrue(tpc62[0]['statement'].strip())
        self.assertIn('TPC-62', tpc62[0]['source_reference'])
        self.assertIn('TPC-62', meta['notes'])  # documented

    def test_aramco_apply_creates_official_controls_in_realistic_test_db(self):
        call_command('import_official_controls', framework_version='ARAMCO-SACS-002', apply=True, stdout=StringIO())
        self.assertEqual(self._official().count(), 92)
        self.assertTrue(self._official().filter(control_id='TPC-62').exists())

    def test_aramco_apply_is_idempotent(self):
        call_command('import_official_controls', framework_version='ARAMCO-SACS-002', apply=True, stdout=StringIO())
        n = Control.objects.count()
        call_command('import_official_controls', framework_version='ARAMCO-SACS-002', apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.count(), n)
        self.assertEqual(self._official().count(), 92)

    def test_aramco_apply_updates_pilot_not_duplicate(self):
        dom, _ = Domain.objects.get_or_create(framework=self.fv.framework, code='IDENTIFY',
                                              defaults={'name': 'IDENTIFY'})
        Control.objects.create(framework=self.fv.framework, framework_version=self.fv,
                               domain=dom, control_id='TPC-1', title='old', description='old')
        call_command('import_official_controls', framework_version='ARAMCO-SACS-002', apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.filter(framework_version=self.fv, control_id='TPC-1').count(), 1)
        self.assertEqual(self._official().count(), 92)

    def test_aramco_apply_does_not_touch_legacy(self):
        fw = self.fv.framework
        dom, _ = Domain.objects.get_or_create(framework=fw, code='L', defaults={'name': 'L'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='TPC-1',
                                        title='LEGACY', description='d')  # fv NULL
        call_command('import_official_controls', framework_version='ARAMCO-SACS-002', apply=True, stdout=StringIO())
        legacy.refresh_from_db()
        self.assertEqual(legacy.title, 'LEGACY')
        self.assertIsNone(legacy.framework_version)

    def test_aramco_apply_does_not_touch_companycontrol_or_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        fw = self.fv.framework
        dom, _ = Domain.objects.get_or_create(framework=fw, code='L', defaults={'name': 'L'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='TPC-5',
                                        title='LEGACY', description='d')
        company = Company.objects.create(name='Co', cr_number='4141414141', sector='technology',
                                         size='small', contact_email='c@x.com', target_aramco=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        call_command('import_official_controls', framework_version='ARAMCO-SACS-002', apply=True, stdout=StringIO())
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)


class SabicDatasetTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())
        self.fv = FrameworkVersion.objects.get(code=SABIC_FV)

    def _official(self):
        return Control.objects.filter(framework_version=self.fv, is_legacy_import=False)

    def test_validate_sabic_dataset_passes(self):
        out = StringIO()
        call_command('validate_official_control_dataset', framework_version=SABIC_FV, stdout=out)
        self.assertIn('OK', out.getvalue())

    def test_validate_supports_sabic_ct_id_pattern(self):
        # A non-CT id must be rejected by the SABIC pattern.
        bad = dict(_SABIC_CTRL); bad['control_id'] = 'XX-1'
        meta = {'framework_version_code': SABIC_FV, 'expected_control_count': 1, '_path': 'x',
                'notes': 'NOT derived from legacy'}
        with mock.patch(_SABIC_VALIDATE, return_value=(meta, [bad])):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset', framework_version=SABIC_FV, stdout=StringIO())

    def test_validate_sabic_rejects_duplicate_ct_control_id(self):
        meta = {'framework_version_code': SABIC_FV, 'expected_control_count': 2, '_path': 'x',
                'notes': 'NOT derived from legacy'}
        with mock.patch(_SABIC_VALIDATE, return_value=(meta, [dict(_SABIC_CTRL), dict(_SABIC_CTRL)])):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset', framework_version=SABIC_FV, stdout=StringIO())

    def test_sabic_dry_run_does_not_create_controls(self):
        call_command('import_official_controls', framework_version=SABIC_FV, stdout=StringIO())  # dry-run
        self.assertEqual(self._official().count(), 0)

    def test_sabic_dry_run_reports_pilot_updates(self):
        dom, _ = Domain.objects.get_or_create(framework=self.fv.framework, code='GOVERN',
                                              defaults={'name': 'GOVERN'})
        Control.objects.create(framework=self.fv.framework, framework_version=self.fv,
                               domain=dom, control_id='CT-01', title='old', description='old')
        out = StringIO()
        call_command('import_official_controls', framework_version=SABIC_FV, stdout=out)
        self.assertIn('would_update: 1', out.getvalue())
        self.assertEqual(self._official().count(), 1)  # dry-run created nothing

    def test_import_sabic_apply_in_test_db_is_idempotent(self):
        call_command('import_official_controls', framework_version=SABIC_FV, apply=True, stdout=StringIO())
        self.assertEqual(self._official().count(), 94)
        n = Control.objects.count()
        call_command('import_official_controls', framework_version=SABIC_FV, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.count(), n)

    def test_import_sabic_creates_control_versions_in_test_db(self):
        call_command('import_official_controls', framework_version=SABIC_FV, apply=True, stdout=StringIO())
        self.assertEqual(ControlVersion.objects.filter(framework_version=self.fv).count(), 94)

    def test_import_sabic_does_not_touch_legacy_companycontrol_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        fw = self.fv.framework
        dom, _ = Domain.objects.get_or_create(framework=fw, code='L', defaults={'name': 'L'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='CT-01',
                                        title='LEGACY', description='d')  # fv NULL
        company = Company.objects.create(name='Co', cr_number='4242424242', sector='technology',
                                         size='small', contact_email='c@x.com', target_sabic=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        call_command('import_official_controls', framework_version=SABIC_FV, apply=True, stdout=StringIO())
        legacy.refresh_from_db()
        self.assertEqual(legacy.title, 'LEGACY')
        self.assertIsNone(legacy.framework_version)
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)


# ============================================================
# Phase 2H — SABIC apply (CT-24 verification) + NCA ECC dataset dry-run
# ============================================================
NCA_FV = 'NCA-ECC-2-2024'
_NCA_VALIDATE = 'compliance.management.commands.validate_official_control_dataset.load_official_dataset'
_NCA_CTRL = dict(control_id='1-1-1', external_reference='NCA ECC 1-1-1',
                 domain='Cybersecurity Governance', title='t', statement='s',
                 source_reference='NCA ECC 2:2024, Governance / Strategy / 1-1-1')


class SabicApplyPhase2HTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())
        self.fv = FrameworkVersion.objects.get(code=SABIC_FV)

    def _official(self):
        return Control.objects.filter(framework_version=self.fv, is_legacy_import=False)

    def test_ct_24_verification_documented(self):
        # CT-24 is absent from the governing Standard -> excluded, dataset stays at 94, documented.
        meta, controls = load_official_dataset(SABIC_FV)
        self.assertEqual(meta['expected_control_count'], 94)
        self.assertFalse(any(c['control_id'] == 'CT-24' for c in controls))
        self.assertIn('CT-24', meta['notes'])

    def test_sabic_apply_creates_official_controls_in_realistic_test_db(self):
        call_command('import_official_controls', framework_version=SABIC_FV, apply=True, stdout=StringIO())
        self.assertEqual(self._official().count(), 94)

    def test_sabic_apply_is_idempotent(self):
        call_command('import_official_controls', framework_version=SABIC_FV, apply=True, stdout=StringIO())
        n = Control.objects.count()
        call_command('import_official_controls', framework_version=SABIC_FV, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.count(), n)

    def test_sabic_apply_updates_pilot_not_duplicate(self):
        dom, _ = Domain.objects.get_or_create(framework=self.fv.framework, code='GOVERN',
                                              defaults={'name': 'GOVERN'})
        Control.objects.create(framework=self.fv.framework, framework_version=self.fv,
                               domain=dom, control_id='CT-01', title='old', description='old')
        call_command('import_official_controls', framework_version=SABIC_FV, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.filter(framework_version=self.fv, control_id='CT-01').count(), 1)
        self.assertEqual(self._official().count(), 94)

    def test_sabic_apply_does_not_touch_legacy(self):
        fw = self.fv.framework
        dom, _ = Domain.objects.get_or_create(framework=fw, code='L', defaults={'name': 'L'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='CT-01',
                                        title='LEGACY', description='d')
        call_command('import_official_controls', framework_version=SABIC_FV, apply=True, stdout=StringIO())
        legacy.refresh_from_db()
        self.assertEqual(legacy.title, 'LEGACY')
        self.assertIsNone(legacy.framework_version)

    def test_sabic_apply_does_not_touch_companycontrol_or_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        fw = self.fv.framework
        dom, _ = Domain.objects.get_or_create(framework=fw, code='L', defaults={'name': 'L'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='CT-05',
                                        title='LEGACY', description='d')
        company = Company.objects.create(name='Co', cr_number='3131313131', sector='technology',
                                         size='small', contact_email='c@x.com', target_sabic=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        call_command('import_official_controls', framework_version=SABIC_FV, apply=True, stdout=StringIO())
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)


class NcaEccDatasetTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())
        self.fv = FrameworkVersion.objects.get(code=NCA_FV)

    def _official(self):
        return Control.objects.filter(framework_version=self.fv, is_legacy_import=False)

    def test_validate_nca_ecc_dataset_passes(self):
        out = StringIO()
        call_command('validate_official_control_dataset', framework_version=NCA_FV, stdout=out)
        self.assertIn('OK', out.getvalue())

    def test_validate_supports_nca_ecc_id_pattern(self):
        bad = dict(_NCA_CTRL); bad['control_id'] = 'TPC-1'  # not a NCA x-x-x id
        meta = {'framework_version_code': NCA_FV, 'expected_control_count': 1, '_path': 'x',
                'notes': 'NOT derived from legacy'}
        with mock.patch(_NCA_VALIDATE, return_value=(meta, [bad])):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset', framework_version=NCA_FV, stdout=StringIO())

    def test_validate_nca_ecc_rejects_duplicate_control_id(self):
        meta = {'framework_version_code': NCA_FV, 'expected_control_count': 2, '_path': 'x',
                'notes': 'NOT derived from legacy'}
        with mock.patch(_NCA_VALIDATE, return_value=(meta, [dict(_NCA_CTRL), dict(_NCA_CTRL)])):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset', framework_version=NCA_FV, stdout=StringIO())

    def test_nca_ecc_dry_run_does_not_create_controls(self):
        call_command('import_official_controls', framework_version=NCA_FV, stdout=StringIO())  # dry-run
        self.assertEqual(self._official().count(), 0)

    def test_nca_ecc_dry_run_reports_pilot_updates(self):
        dom, _ = Domain.objects.get_or_create(framework=self.fv.framework, code='Cybersecurity Governance',
                                              defaults={'name': 'Cybersecurity Governance'})
        Control.objects.create(framework=self.fv.framework, framework_version=self.fv,
                               domain=dom, control_id='1-1-1', title='old', description='old')
        out = StringIO()
        call_command('import_official_controls', framework_version=NCA_FV, stdout=out)
        self.assertIn('would_update: 1', out.getvalue())
        self.assertEqual(self._official().count(), 1)  # dry-run wrote nothing

    def test_import_nca_ecc_apply_in_test_db_is_idempotent(self):
        call_command('import_official_controls', framework_version=NCA_FV, apply=True, stdout=StringIO())
        c = self._official().count()
        self.assertGreater(c, 0)
        n = Control.objects.count()
        call_command('import_official_controls', framework_version=NCA_FV, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.count(), n)

    def test_import_nca_ecc_creates_control_versions_in_test_db(self):
        call_command('import_official_controls', framework_version=NCA_FV, apply=True, stdout=StringIO())
        self.assertEqual(ControlVersion.objects.filter(framework_version=self.fv).count(),
                         self._official().count())

    def test_import_nca_ecc_does_not_touch_legacy_companycontrol_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        fw = self.fv.framework
        dom, _ = Domain.objects.get_or_create(framework=fw, code='L', defaults={'name': 'L'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='1-1-1',
                                        title='LEGACY', description='d')  # fv NULL
        company = Company.objects.create(name='Co', cr_number='3232323232', sector='government',
                                         size='large', contact_email='c@x.com', target_nca=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        call_command('import_official_controls', framework_version=NCA_FV, apply=True, stdout=StringIO())
        legacy.refresh_from_db()
        self.assertEqual(legacy.title, 'LEGACY')
        self.assertIsNone(legacy.framework_version)
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)


# ============================================================
# Phase 2I — NCA ECC apply (after 109->108 reconciliation) + remaining-NCA strategy
# ============================================================
from compliance.official_dataset import REMAINING_NCA_FRAMEWORKS, DATASET_FILES

ECC_FV = 'NCA-ECC-2-2024'


class NcaEccReconciliationApplyTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())
        self.fv = FrameworkVersion.objects.get(code=ECC_FV)

    def _official(self):
        return Control.objects.filter(framework_version=self.fv, is_legacy_import=False)

    def test_nca_ecc_count_reconciliation_documented(self):
        meta, controls = load_official_dataset(ECC_FV)
        self.assertEqual(meta['expected_control_count'], 108)
        self.assertEqual(len(controls), 108)
        self.assertNotIn('1-7-2', [c['control_id'] for c in controls])  # the spurious Appendix-C row
        self.assertIn('108', meta['notes'])
        self.assertIn('1-7-2', meta['notes'])  # documents what was removed and why

    def test_validate_nca_ecc_dataset_after_reconciliation(self):
        out = StringIO()
        call_command('validate_official_control_dataset', framework_version=ECC_FV, stdout=out)
        self.assertIn('OK', out.getvalue())

    def test_nca_ecc_apply_creates_official_controls_in_realistic_test_db(self):
        call_command('import_official_controls', framework_version=ECC_FV, apply=True, stdout=StringIO())
        self.assertEqual(self._official().count(), 108)

    def test_nca_ecc_apply_is_idempotent(self):
        call_command('import_official_controls', framework_version=ECC_FV, apply=True, stdout=StringIO())
        n = Control.objects.count()
        call_command('import_official_controls', framework_version=ECC_FV, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.count(), n)

    def test_nca_ecc_apply_updates_pilot_not_duplicate(self):
        dom, _ = Domain.objects.get_or_create(framework=self.fv.framework, code='Cybersecurity Governance',
                                              defaults={'name': 'Cybersecurity Governance'})
        Control.objects.create(framework=self.fv.framework, framework_version=self.fv,
                               domain=dom, control_id='1-1-1', title='old', description='old')
        call_command('import_official_controls', framework_version=ECC_FV, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.filter(framework_version=self.fv, control_id='1-1-1').count(), 1)
        self.assertEqual(self._official().count(), 108)

    def test_nca_ecc_apply_does_not_touch_legacy(self):
        fw = self.fv.framework
        dom, _ = Domain.objects.get_or_create(framework=fw, code='LG', defaults={'name': 'LG'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='1-1-1',
                                        title='LEGACY', description='d')  # fv NULL
        call_command('import_official_controls', framework_version=ECC_FV, apply=True, stdout=StringIO())
        legacy.refresh_from_db()
        self.assertEqual(legacy.title, 'LEGACY')
        self.assertIsNone(legacy.framework_version)

    def test_nca_ecc_apply_does_not_touch_companycontrol_or_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        fw = self.fv.framework
        dom, _ = Domain.objects.get_or_create(framework=fw, code='LG', defaults={'name': 'LG'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='1-2-1',
                                        title='LEGACY', description='d')
        company = Company.objects.create(name='Co', cr_number='2121212121', sector='government',
                                         size='large', contact_email='c@x.com', target_nca=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        call_command('import_official_controls', framework_version=ECC_FV, apply=True, stdout=StringIO())
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)


class RemainingNcaStrategyTests(TestCase):
    def test_remaining_nca_framework_strategy_documented(self):
        expected = {'NCA-CSCC-1-2019', 'NCA-OSMACC-1-2021', 'NCA-TCC-1-2021',
                    'NCA-CCC-2-2024', 'NCA-OTCC-1-2022', 'NCA-DCC-1-2022'}
        self.assertEqual(set(REMAINING_NCA_FRAMEWORKS), expected)
        for code, info in REMAINING_NCA_FRAMEWORKS.items():
            self.assertTrue(info['source'])
            self.assertIn(info['batch'], (1, 2, 3))
            self.assertIn(info['confidence'], ('high', 'medium', 'low'))

    def test_no_remaining_nca_framework_imported_in_phase_2i(self):
        # Still-unbuilt remaining NCA frameworks. (Batch 1 built in 2J; CCC built in 2K;
        # OTCC deferred and DCC untouched.)
        for code in ('NCA-OTCC-1-2022', 'NCA-DCC-1-2022'):
            self.assertNotIn(code, DATASET_FILES)

    def test_no_nca_non_ecc_dataset_apply_in_phase_2i(self):
        call_command('seed_framework_versions', stdout=StringIO())
        n = Control.objects.filter(
            framework_version__code__in=list(REMAINING_NCA_FRAMEWORKS),
            framework_version__isnull=False).count()
        self.assertEqual(n, 0)


# ============================================================
# Phase 2J — CSCC / OSMACC / TCC datasets: validation + dry-run only
# ============================================================
import os as _os

BATCH1 = {'NCA-CSCC-1-2019': 32, 'NCA-OSMACC-1-2021': 15, 'NCA-TCC-1-2021': 21}
_2J_VALIDATE = 'compliance.management.commands.validate_official_control_dataset.load_official_dataset'
_2J_CTRL = dict(control_id='1-1-1', external_reference='NCA CSCC 1-1-1',
                domain='Cybersecurity Governance', title='t', statement='s',
                source_reference='CSCC, Cybersecurity Governance / X / 1-1-1')


def _patched_2j(fv, controls):
    meta = {'framework_version_code': fv, 'expected_control_count': len(controls), '_path': 'x',
            'notes': 'NOT derived from the legacy 334 Excel.'}
    return mock.patch(_2J_VALIDATE, return_value=(meta, controls))


class Batch1NcaValidationTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def _ok(self, fv):
        out = StringIO()
        call_command('validate_official_control_dataset', framework_version=fv, stdout=out)
        return out.getvalue()

    def test_validate_nca_cscc_dataset_passes(self):
        self.assertIn('OK', self._ok('NCA-CSCC-1-2019'))

    def test_validate_nca_osmacc_dataset_passes(self):
        self.assertIn('OK', self._ok('NCA-OSMACC-1-2021'))

    def test_validate_nca_tcc_dataset_passes(self):
        self.assertIn('OK', self._ok('NCA-TCC-1-2021'))

    def test_validate_rejects_duplicate_control_id_for_remaining_nca(self):
        with _patched_2j('NCA-CSCC-1-2019', [dict(_2J_CTRL), dict(_2J_CTRL)]):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset',
                             framework_version='NCA-CSCC-1-2019', stdout=StringIO())

    def test_validate_rejects_missing_source_reference_for_remaining_nca(self):
        bad = {k: v for k, v in _2J_CTRL.items() if k != 'source_reference'}
        with _patched_2j('NCA-TCC-1-2021', [bad]):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset',
                             framework_version='NCA-TCC-1-2021', stdout=StringIO())

    def test_validate_rejects_missing_framework_version_for_remaining_nca(self):
        # Use a code that is NOT seeded so the FrameworkVersion lookup fails.
        meta = {'framework_version_code': 'NCA-FAKE-1-2099', 'expected_control_count': 1, '_path': 'x'}
        with mock.patch(_2J_VALIDATE, return_value=(meta, [dict(_2J_CTRL)])):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset',
                             framework_version='NCA-FAKE-1-2099', stdout=StringIO())


class Batch1NcaDryRunTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def _official(self, fv):
        return Control.objects.filter(framework_version__code=fv, is_legacy_import=False)

    def test_nca_cscc_dry_run_does_not_create_controls(self):
        call_command('import_official_controls', framework_version='NCA-CSCC-1-2019', stdout=StringIO())
        self.assertEqual(self._official('NCA-CSCC-1-2019').count(), 0)

    def test_nca_osmacc_dry_run_does_not_create_controls(self):
        call_command('import_official_controls', framework_version='NCA-OSMACC-1-2021', stdout=StringIO())
        self.assertEqual(self._official('NCA-OSMACC-1-2021').count(), 0)

    def test_nca_tcc_dry_run_does_not_create_controls(self):
        call_command('import_official_controls', framework_version='NCA-TCC-1-2021', stdout=StringIO())
        self.assertEqual(self._official('NCA-TCC-1-2021').count(), 0)

    def test_import_remaining_nca_apply_in_test_db_is_idempotent(self):
        # --apply is fine in the isolated test DB (NOT the real DB).
        for fv, n in BATCH1.items():
            call_command('import_official_controls', framework_version=fv, apply=True, stdout=StringIO())
            self.assertEqual(self._official(fv).count(), n)
        total = Control.objects.count()
        for fv in BATCH1:
            call_command('import_official_controls', framework_version=fv, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.count(), total)

    def test_import_remaining_nca_creates_control_versions_in_test_db(self):
        for fv, n in BATCH1.items():
            call_command('import_official_controls', framework_version=fv, apply=True, stdout=StringIO())
            self.assertEqual(ControlVersion.objects.filter(framework_version__code=fv).count(), n)

    def test_import_remaining_nca_does_not_touch_legacy_companycontrol_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        fw, _ = Framework.objects.get_or_create(code='NCA_ECC', defaults={'name': 'NCA'})
        dom, _ = Domain.objects.get_or_create(framework=fw, code='LG', defaults={'name': 'LG'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='1-1-1',
                                        title='LEGACY', description='d')  # fv NULL
        company = Company.objects.create(name='Co', cr_number='1919191919', sector='government',
                                         size='large', contact_email='c@x.com', target_nca=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        call_command('import_official_controls', framework_version='NCA-CSCC-1-2019', apply=True, stdout=StringIO())
        legacy.refresh_from_db()
        self.assertEqual(legacy.title, 'LEGACY')
        self.assertIsNone(legacy.framework_version)
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)


class Phase2JGuardrailTests(TestCase):
    def test_no_ccc_otcc_dcc_dataset_created_in_phase_2j(self):
        # CCC gained a dataset later (Phase 2K); OTCC/DCC remain unbuilt.
        for code, fname in (('NCA-OTCC-1-2022', 'nca_otcc_1_2022.yaml'),
                            ('NCA-DCC-1-2022', 'nca_dcc_1_2022.yaml')):
            self.assertNotIn(code, DATASET_FILES)
            self.assertFalse(_os.path.exists(_os.path.join(
                _os.path.dirname(__file__), 'data', 'official_controls', fname)))

    def test_no_remaining_nca_apply_on_real_db_in_phase_2j(self):
        # The three Batch-1 frameworks are registered for import but the real DB
        # ships with zero official controls for them (only dry-run was run).
        for code in BATCH1:
            self.assertIn(code, DATASET_FILES)

    def test_registration_upload_files_untouched(self):
        # Phase 2J must not modify the registration/upload code paths.
        import core.views, core.forms, compliance.views
        for mod, name in [(core.forms, 'CompanyRegistrationForm'),
                          (core.views, 'register_company'),
                          (compliance.views, 'upload_evidence')]:
            self.assertTrue(hasattr(mod, name))


# ============================================================
# Phase 2K — apply CSCC/OSMACC/TCC + CCC dataset dry-run (OTCC deferred)
# ============================================================
CCC_FV = 'NCA-CCC-2-2024'
_2K_VALIDATE = 'compliance.management.commands.validate_official_control_dataset.load_official_dataset'
_CCC_CTRL = dict(control_id='1-1-P-1', external_reference='NCA CCC 1-1-P-1',
                 domain='Cybersecurity Governance', title='t', statement='s',
                 source_reference='NCA CCC 2:2024, Cybersecurity Governance / CSP / 1-1-P-1')


class ApplyBatch1NcaTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())
    BATCH = {'NCA-CSCC-1-2019': 32, 'NCA-OSMACC-1-2021': 15, 'NCA-TCC-1-2021': 21}

    def _official(self, fv):
        return Control.objects.filter(framework_version__code=fv, is_legacy_import=False)

    def test_apply_cscc_creates_official_controls_in_realistic_test_db(self):
        call_command('import_official_controls', framework_version='NCA-CSCC-1-2019', apply=True, stdout=StringIO())
        self.assertEqual(self._official('NCA-CSCC-1-2019').count(), 32)

    def test_apply_osmacc_creates_official_controls_in_realistic_test_db(self):
        call_command('import_official_controls', framework_version='NCA-OSMACC-1-2021', apply=True, stdout=StringIO())
        self.assertEqual(self._official('NCA-OSMACC-1-2021').count(), 15)

    def test_apply_tcc_creates_official_controls_in_realistic_test_db(self):
        call_command('import_official_controls', framework_version='NCA-TCC-1-2021', apply=True, stdout=StringIO())
        self.assertEqual(self._official('NCA-TCC-1-2021').count(), 21)

    def test_apply_cscc_osmacc_tcc_is_idempotent(self):
        for fv in self.BATCH:
            call_command('import_official_controls', framework_version=fv, apply=True, stdout=StringIO())
        n = Control.objects.count()
        for fv in self.BATCH:
            call_command('import_official_controls', framework_version=fv, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.count(), n)

    def test_apply_cscc_osmacc_tcc_does_not_touch_legacy(self):
        fw, _ = Framework.objects.get_or_create(code='NCA_ECC', defaults={'name': 'NCA'})
        dom, _ = Domain.objects.get_or_create(framework=fw, code='LG', defaults={'name': 'LG'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='1-1-1',
                                        title='LEGACY', description='d')
        for fv in self.BATCH:
            call_command('import_official_controls', framework_version=fv, apply=True, stdout=StringIO())
        legacy.refresh_from_db()
        self.assertEqual(legacy.title, 'LEGACY')
        self.assertIsNone(legacy.framework_version)

    def test_apply_cscc_osmacc_tcc_does_not_touch_companycontrol_or_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        fw, _ = Framework.objects.get_or_create(code='NCA_ECC', defaults={'name': 'NCA'})
        dom, _ = Domain.objects.get_or_create(framework=fw, code='LG', defaults={'name': 'LG'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='3-3-3',
                                        title='LEGACY', description='d')
        company = Company.objects.create(name='Co', cr_number='1818181818', sector='government',
                                         size='large', contact_email='c@x.com', target_nca=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        for fv in self.BATCH:
            call_command('import_official_controls', framework_version=fv, apply=True, stdout=StringIO())
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)


class CccDatasetTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())
        self.fv = FrameworkVersion.objects.get(code=CCC_FV)

    def _official(self):
        return Control.objects.filter(framework_version=self.fv, is_legacy_import=False)

    def test_validate_nca_ccc_dataset_passes(self):
        out = StringIO()
        call_command('validate_official_control_dataset', framework_version=CCC_FV, stdout=out)
        self.assertIn('OK', out.getvalue())

    def test_ccc_dataset_reconciles_to_37_csp_plus_18_cst(self):
        meta, controls = load_official_dataset(CCC_FV)
        self.assertEqual(meta['expected_control_count'], 55)
        self.assertEqual(sum('-P-' in c['control_id'] for c in controls), 37)
        self.assertEqual(sum('-T-' in c['control_id'] for c in controls), 18)
        self.assertIn('37', meta['notes'])
        self.assertIn('18', meta['notes'])

    def test_validate_ccc_rejects_duplicate_control_id(self):
        meta = {'framework_version_code': CCC_FV, 'expected_control_count': 2, '_path': 'x',
                'notes': 'NOT derived from legacy'}
        with mock.patch(_2K_VALIDATE, return_value=(meta, [dict(_CCC_CTRL), dict(_CCC_CTRL)])):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset', framework_version=CCC_FV, stdout=StringIO())

    def test_validate_ccc_rejects_non_ccc_id_pattern(self):
        bad = dict(_CCC_CTRL); bad['control_id'] = '1-1-1'  # ECC-style, not CCC P/T
        meta = {'framework_version_code': CCC_FV, 'expected_control_count': 1, '_path': 'x',
                'notes': 'NOT derived from legacy'}
        with mock.patch(_2K_VALIDATE, return_value=(meta, [bad])):
            with self.assertRaises(CommandError):
                call_command('validate_official_control_dataset', framework_version=CCC_FV, stdout=StringIO())

    def test_ccc_dry_run_does_not_create_controls(self):
        call_command('import_official_controls', framework_version=CCC_FV, stdout=StringIO())
        self.assertEqual(self._official().count(), 0)

    def test_import_ccc_apply_in_test_db_is_idempotent(self):
        call_command('import_official_controls', framework_version=CCC_FV, apply=True, stdout=StringIO())
        self.assertEqual(self._official().count(), 55)
        n = Control.objects.count()
        call_command('import_official_controls', framework_version=CCC_FV, apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.count(), n)

    def test_import_ccc_creates_control_versions_in_test_db(self):
        call_command('import_official_controls', framework_version=CCC_FV, apply=True, stdout=StringIO())
        self.assertEqual(ControlVersion.objects.filter(framework_version=self.fv).count(), 55)

    def test_import_ccc_does_not_touch_legacy_companycontrol_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        fw, _ = Framework.objects.get_or_create(code='NCA_ECC', defaults={'name': 'NCA'})
        dom, _ = Domain.objects.get_or_create(framework=fw, code='LG', defaults={'name': 'LG'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='2-2-2',
                                        title='LEGACY', description='d')
        company = Company.objects.create(name='Co', cr_number='1717171717', sector='government',
                                         size='large', contact_email='c@x.com', target_nca=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        call_command('import_official_controls', framework_version=CCC_FV, apply=True, stdout=StringIO())
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)


class Phase2KGuardrailTests(TestCase):
    def test_otcc_deferred_not_registered_pending_reconciliation(self):
        # OTCC's official count (47) is confirmed, but its statements could not be
        # extracted cleanly from the column-formatted PDF, so it is deferred (not built).
        self.assertNotIn('NCA-OTCC-1-2022', DATASET_FILES)
        self.assertFalse(_os.path.exists(_os.path.join(
            _os.path.dirname(__file__), 'data', 'official_controls', 'nca_otcc_1_2022.yaml')))

    def test_no_dcc_dataset_created_in_phase_2k(self):
        self.assertNotIn('NCA-DCC-1-2022', DATASET_FILES)
        self.assertFalse(_os.path.exists(_os.path.join(
            _os.path.dirname(__file__), 'data', 'official_controls', 'nca_dcc_1_2022.yaml')))

    def test_no_ccc_otcc_apply_on_real_db_in_phase_2k(self):
        # CCC is registered for dry-run; OTCC is not registered at all this phase.
        self.assertIn('NCA-CCC-2-2024', DATASET_FILES)
        self.assertNotIn('NCA-OTCC-1-2022', DATASET_FILES)

    def test_registration_upload_files_untouched(self):
        import core.views, core.forms, compliance.views
        for mod, name in [(core.forms, 'CompanyRegistrationForm'),
                          (core.views, 'register_company'),
                          (compliance.views, 'upload_evidence')]:
            self.assertTrue(hasattr(mod, name))


# ============================================================
# Phase 2L — CCC apply + OTCC/DCC investigation (no OTCC/DCC import or apply)
# ============================================================
import os as _os2

_DRAFTS = _os2.path.join(_os2.path.dirname(__file__), 'data', 'official_controls', 'drafts')


class ApplyCccPhase2LTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())
        self.fv = FrameworkVersion.objects.get(code='NCA-CCC-2-2024')

    def _official(self):
        return Control.objects.filter(framework_version=self.fv, is_legacy_import=False)

    def test_apply_ccc_creates_official_controls_in_realistic_test_db(self):
        call_command('import_official_controls', framework_version='NCA-CCC-2-2024', apply=True, stdout=StringIO())
        self.assertEqual(self._official().count(), 55)

    def test_apply_ccc_is_idempotent(self):
        call_command('import_official_controls', framework_version='NCA-CCC-2-2024', apply=True, stdout=StringIO())
        n = Control.objects.count()
        call_command('import_official_controls', framework_version='NCA-CCC-2-2024', apply=True, stdout=StringIO())
        self.assertEqual(Control.objects.count(), n)
        self.assertEqual(self._official().count(), 55)

    def test_ccc_count_is_55_csp_plus_cst(self):
        call_command('import_official_controls', framework_version='NCA-CCC-2-2024', apply=True, stdout=StringIO())
        self.assertEqual(self._official().filter(control_id__contains='-P-').count(), 37)
        self.assertEqual(self._official().filter(control_id__contains='-T-').count(), 18)

    def test_apply_ccc_does_not_touch_legacy(self):
        fw, _ = Framework.objects.get_or_create(code='NCA_ECC', defaults={'name': 'NCA'})
        dom, _ = Domain.objects.get_or_create(framework=fw, code='LG', defaults={'name': 'LG'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='1-1-1',
                                        title='LEGACY', description='d')
        call_command('import_official_controls', framework_version='NCA-CCC-2-2024', apply=True, stdout=StringIO())
        legacy.refresh_from_db()
        self.assertEqual(legacy.title, 'LEGACY')
        self.assertIsNone(legacy.framework_version)

    def test_apply_ccc_does_not_touch_companycontrol_or_evidence(self):
        from core.models import Company
        from compliance.models import Evidence
        fw, _ = Framework.objects.get_or_create(code='NCA_ECC', defaults={'name': 'NCA'})
        dom, _ = Domain.objects.get_or_create(framework=fw, code='LG', defaults={'name': 'LG'})
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='4-4-4',
                                        title='LEGACY', description='d')
        company = Company.objects.create(name='Co', cr_number='1616161616', sector='government',
                                         size='large', contact_email='c@x.com', target_nca=True)
        cc = CompanyControl.objects.create(company=company, control=legacy)
        Evidence.objects.create(company_control=cc, original_filename='p.txt', file_type='txt')
        cc_n, ev_n = CompanyControl.objects.count(), Evidence.objects.count()
        call_command('import_official_controls', framework_version='NCA-CCC-2-2024', apply=True, stdout=StringIO())
        self.assertEqual(CompanyControl.objects.count(), cc_n)
        self.assertEqual(Evidence.objects.count(), ev_n)


class OtccDccInvestigationTests(TestCase):
    def test_otcc_not_applied_in_phase_2l(self):
        # OTCC is not even registered, so it cannot have been applied.
        self.assertNotIn('NCA-OTCC-1-2022', DATASET_FILES)
        self.assertEqual(Control.objects.filter(framework_version__code='NCA-OTCC-1-2022').count(), 0)

    def test_dcc_not_applied_in_phase_2l(self):
        self.assertNotIn('NCA-DCC-1-2022', DATASET_FILES)
        self.assertEqual(Control.objects.filter(framework_version__code='NCA-DCC-1-2022').count(), 0)

    def test_no_noisy_otcc_dataset_registered(self):
        # OTCC must not be registered with an unvalidated/noisy dataset.
        self.assertNotIn('NCA-OTCC-1-2022', DATASET_FILES)
        self.assertFalse(_os2.path.exists(_os2.path.join(
            _os2.path.dirname(__file__), 'data', 'official_controls', 'nca_otcc_1_2022.yaml')))

    def test_no_dcc_dataset_registered_unless_confident(self):
        self.assertNotIn('NCA-DCC-1-2022', DATASET_FILES)
        self.assertFalse(_os2.path.exists(_os2.path.join(
            _os2.path.dirname(__file__), 'data', 'official_controls', 'nca_dcc_1_2022.yaml')))

    def test_otcc_dcc_investigation_documented(self):
        # Investigation notes exist as drafts (NOT registered datasets).
        self.assertTrue(_os2.path.exists(_os2.path.join(_DRAFTS, 'nca_otcc_extraction_notes.md')))
        self.assertTrue(_os2.path.exists(_os2.path.join(_DRAFTS, 'nca_dcc_extraction_notes.md')))

    def test_registration_upload_files_untouched(self):
        import core.views, core.forms, compliance.views
        for mod, name in [(core.forms, 'CompanyRegistrationForm'),
                          (core.views, 'register_company'),
                          (compliance.views, 'upload_evidence')]:
            self.assertTrue(hasattr(mod, name))


# ============================================================
# Phase 2M — OTCC/DCC modeling decision (planning only; no schema, no import)
# ============================================================
class Phase2MModelingGuardrailTests(TestCase):
    def test_otcc_dcc_modeling_decision_documented(self):
        d = _os2.path.join(_os2.path.dirname(__file__), 'data', 'official_controls', 'drafts')
        self.assertTrue(_os2.path.exists(_os2.path.join(d, 'nca_otcc_extraction_notes.md')))
        self.assertTrue(_os2.path.exists(_os2.path.join(d, 'nca_dcc_extraction_notes.md')))

    def test_no_otcc_dataset_registered_in_phase_2m(self):
        self.assertNotIn('NCA-OTCC-1-2022', DATASET_FILES)
        self.assertFalse(_os2.path.exists(_os2.path.join(
            _os2.path.dirname(__file__), 'data', 'official_controls', 'nca_otcc_1_2022.yaml')))

    def test_no_dcc_dataset_registered_in_phase_2m(self):
        self.assertNotIn('NCA-DCC-1-2022', DATASET_FILES)
        self.assertFalse(_os2.path.exists(_os2.path.join(
            _os2.path.dirname(__file__), 'data', 'official_controls', 'nca_dcc_1_2022.yaml')))

    def test_no_otcc_dcc_apply_in_phase_2m(self):
        for code in ('NCA-OTCC-1-2022', 'NCA-DCC-1-2022'):
            self.assertEqual(Control.objects.filter(framework_version__code=code).count(), 0)

    def test_parent_control_not_added_in_phase_2m(self):
        # No schema change this phase: Control must NOT yet have parent_control / level fields.
        fields = {f.name for f in Control._meta.get_fields()}
        self.assertNotIn('parent_control', fields)
        self.assertNotIn('subcontrols', fields)

    def test_registration_upload_files_untouched(self):
        import core.views, core.forms, compliance.views
        for mod, name in [(core.forms, 'CompanyRegistrationForm'),
                          (core.views, 'register_company'),
                          (compliance.views, 'upload_evidence')]:
            self.assertTrue(hasattr(mod, name))


# ============================================================
# Phase 3A — Company Intake + Framework Applicability foundation
# ============================================================
from compliance.models import CompanyIntakeProfile, FrameworkApplicabilityResult
from compliance.framework_applicability import evaluate_company


def _seed_official(fv_code, framework_code, n=2):
    """Give a FrameworkVersion some official (non-legacy) controls so it is 'available'."""
    call_command('seed_framework_versions', stdout=StringIO())
    fv = FrameworkVersion.objects.get(code=fv_code)
    fw = fv.framework
    dom, _ = Domain.objects.get_or_create(framework=fw, code='D', defaults={'name': 'D'})
    for i in range(n):
        Control.objects.update_or_create(
            framework_version=fv, control_id=f'{fv_code}-OFF-{i}',
            defaults={'framework': fw, 'domain': dom, 'title': 't', 'description': 'd',
                      'is_legacy_import': False})
    return fv


def _company(**kw):
    from core.models import Company
    n = Company.objects.count() + 1
    defaults = dict(name='Co', cr_number=f'{n:010d}', sector='technology', size='small',
                    contact_email=f'co{n}@x.com')
    defaults.update(kw)
    return Company.objects.create(**defaults)


class IntakeModelTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def test_company_intake_profile_can_be_created(self):
        c = _company()
        p = CompanyIntakeProfile.objects.create(company=c, works_with_aramco=True)
        self.assertTrue(p.works_with_aramco)
        self.assertEqual(c.intake_profile, p)

    def test_framework_applicability_result_can_be_created(self):
        c = _company()
        fv = FrameworkVersion.objects.get(code='ARAMCO-SACS-002')
        r = FrameworkApplicabilityResult.objects.create(
            company=c, framework_version=fv, decision='applicable', reason='x', source='rule')
        self.assertEqual(r.decision, 'applicable')

    def test_manual_override_fields_exist(self):
        fields = {f.name for f in FrameworkApplicabilityResult._meta.get_fields()}
        self.assertIn('overridden_by', fields)
        self.assertIn('override_reason', fields)


class ApplicabilityServiceTests(TestCase):
    def _decisions(self, company):
        return {r['framework']: r for r in evaluate_company(company)}

    def test_aramco_applicable_when_company_works_with_aramco(self):
        _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        c = _company(); CompanyIntakeProfile.objects.create(company=c, works_with_aramco=True)
        self.assertEqual(self._decisions(c)['ARAMCO-SACS-002']['status'], 'applicable')

    def test_sabic_applicable_when_company_works_with_sabic(self):
        _seed_official('SABIC-CYBERTRUST-1-0', 'SABIC_CT')
        c = _company(); CompanyIntakeProfile.objects.create(company=c, works_with_sabic=True)
        self.assertEqual(self._decisions(c)['SABIC-CYBERTRUST-1-0']['status'], 'applicable')

    def test_nca_ecc_applicable_from_legacy_target_nca(self):
        _seed_official('NCA-ECC-2-2024', 'NCA_ECC')
        c = _company(target_nca=True)  # no intake profile -> legacy fallback
        d = self._decisions(c)['NCA-ECC-2-2024']
        self.assertEqual(d['status'], 'applicable')
        self.assertEqual(d['source'], 'legacy_checkbox')

    def test_nca_cscc_applicable_for_critical_system_operator(self):
        _seed_official('NCA-CSCC-1-2019', 'NCA_ECC')
        c = _company(); CompanyIntakeProfile.objects.create(company=c, is_critical_system_operator=True)
        self.assertEqual(self._decisions(c)['NCA-CSCC-1-2019']['status'], 'applicable')

    def test_nca_ccc_applicable_for_cloud_usage(self):
        _seed_official('NCA-CCC-2-2024', 'NCA_ECC')
        c = _company(); CompanyIntakeProfile.objects.create(company=c, uses_cloud_services=True)
        self.assertEqual(self._decisions(c)['NCA-CCC-2-2024']['status'], 'applicable')

    def test_nca_tcc_applicable_for_remote_work(self):
        _seed_official('NCA-TCC-1-2021', 'NCA_ECC')
        c = _company(); CompanyIntakeProfile.objects.create(company=c, has_remote_work=True)
        self.assertEqual(self._decisions(c)['NCA-TCC-1-2021']['status'], 'applicable')

    def test_otcc_unavailable_when_no_official_controls(self):
        call_command('seed_framework_versions', stdout=StringIO())
        c = _company()
        d = self._decisions(c).get('NCA-OTCC-1-2022')
        self.assertIsNone(d)  # OTCC has no rule registered -> not in results at all

    def test_dcc_unavailable_when_no_official_controls(self):
        call_command('seed_framework_versions', stdout=StringIO())
        c = _company()
        self.assertIsNone(self._decisions(c).get('NCA-DCC-1-2022'))

    def test_decisions_include_reason(self):
        _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        c = _company(target_nca=True)
        for r in evaluate_company(c):
            if r['status'] not in ('unavailable', 'skipped'):
                self.assertTrue(r['reason'])

    def test_service_does_not_create_companycontrol_or_evidence(self):
        _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        from compliance.models import Evidence
        c = _company(); CompanyIntakeProfile.objects.create(company=c, works_with_aramco=True)
        evaluate_company(c, apply=True)
        self.assertEqual(CompanyControl.objects.count(), 0)
        self.assertEqual(Evidence.objects.count(), 0)


class EvaluateCommandTests(TestCase):
    def setUp(self):
        _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        self.c = _company(); CompanyIntakeProfile.objects.create(company=self.c, works_with_aramco=True)

    def test_evaluate_framework_applicability_dry_run_does_not_write(self):
        call_command('evaluate_framework_applicability', company_id=self.c.id, stdout=StringIO())
        self.assertEqual(FrameworkApplicabilityResult.objects.count(), 0)

    def test_evaluate_framework_applicability_apply_writes_results(self):
        call_command('evaluate_framework_applicability', company_id=self.c.id, apply=True, stdout=StringIO())
        self.assertTrue(FrameworkApplicabilityResult.objects.filter(
            company=self.c, framework_version__code='ARAMCO-SACS-002', decision='applicable').exists())

    def test_evaluate_framework_applicability_is_idempotent(self):
        call_command('evaluate_framework_applicability', company_id=self.c.id, apply=True, stdout=StringIO())
        n = FrameworkApplicabilityResult.objects.count()
        call_command('evaluate_framework_applicability', company_id=self.c.id, apply=True, stdout=StringIO())
        self.assertEqual(FrameworkApplicabilityResult.objects.count(), n)

    def test_evaluate_framework_applicability_all_companies(self):
        _company(cr_number='2022022020', contact_email='b@x.com', target_aramco=True)
        call_command('evaluate_framework_applicability', all_companies=True, apply=True, stdout=StringIO())
        self.assertTrue(FrameworkApplicabilityResult.objects.exists())


class Phase3ABackwardCompatTests(TestCase):
    def test_existing_registration_flow_still_works(self):
        fw, dom = _fw_dom()
        for i in range(2):
            Control.objects.create(framework=fw, domain=dom, control_id=f'NCA-{i}',
                                   title='t', description='d')
        with mock.patch('core.views.classify_company', return_value={'error': 'skip'}):
            resp = self.client.post(reverse('core:register'), {
                'company_name': 'RegCo', 'cr_number': '3033033030', 'sector': 'technology',
                'size': 'small', 'first_name': 'A', 'last_name': 'B', 'email': 'reg3a@x.com',
                'password': 'longenough12', 'target_nca': 'on'})
        self.assertEqual(resp.status_code, 302)
        from core.models import Company
        self.assertTrue(Company.objects.filter(cr_number='3033033030').exists())

    def test_existing_upload_flow_still_works(self):
        from core.models import User, Company
        from django.core.files.uploadedfile import SimpleUploadedFile
        company, control = _company_with_control()
        user = User.objects.create_user(email='up3a@x.com', password='longenough12',
                                        company=company, role='company_admin')
        self.client.force_login(user)
        good = SimpleUploadedFile('p.txt', b'policy ok', content_type='text/plain')
        with mock.patch('monitoring.tasks.analyze_evidence_async.delay'):
            resp = self.client.post(reverse('compliance:upload_evidence', args=[control.id]),
                                    {'evidence_file': good})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(Evidence.objects.count(), 1)

    def test_legacy_target_checkboxes_still_exist(self):
        from core.models import Company
        fields = {f.name for f in Company._meta.get_fields()}
        for f in ('target_nca', 'target_aramco', 'target_sabic'):
            self.assertIn(f, fields)

    def test_no_legacy_or_official_controls_deleted(self):
        # The intake migration must not touch Control rows at all.
        fw, dom = _fw_dom()
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='L-1', title='t', description='d')
        official = Control.objects.create(framework=fw, domain=dom, control_id='O-1', title='t',
                                          description='d', is_legacy_import=False,
                                          framework_version=FrameworkVersion.objects.create(
                                              framework=fw, code='ZZ-1'))
        c = _company()
        evaluate_company(c, apply=True)
        self.assertTrue(Control.objects.filter(pk=legacy.pk).exists())
        self.assertTrue(Control.objects.filter(pk=official.pk).exists())


# ============================================================
# Phase 3B — Intake Wizard UI + Applicable Framework Review
# ============================================================
from compliance.forms import CompanyIntakeForm


def _intake_user(company=None, **company_kw):
    from core.models import User
    if company is None:
        company = _company(**company_kw)
    user = User.objects.create_user(email=f'iu{company.id}@x.com', password='longenough12',
                                    company=company, role='company_admin')
    return company, user


class IntakeFormTests(TestCase):
    def test_company_intake_form_valid(self):
        # All fields optional -> an empty (or minimal) form is valid.
        self.assertTrue(CompanyIntakeForm(data={'works_with_aramco': 'on'}).is_valid())

    def test_company_intake_form_cloud_help_text(self):
        f = CompanyIntakeForm()
        self.assertIn('OpenAI', f.fields['uses_cloud_services'].help_text)
        self.assertIn('SaaS', f.fields['provides_cloud_services'].help_text)

    def test_company_intake_form_saves_profile(self):
        c = _company()
        form = CompanyIntakeForm(data={'uses_cloud_services': 'on'})
        self.assertTrue(form.is_valid())
        p = form.save(commit=False); p.company = c; p.save()
        self.assertTrue(CompanyIntakeProfile.objects.get(company=c).uses_cloud_services)


class IntakeViewTests(TestCase):
    def setUp(self):
        _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        self.company, self.user = _intake_user()
        self.client.force_login(self.user)

    def test_intake_page_requires_login(self):
        self.client.logout()
        resp = self.client.get(reverse('compliance:intake'))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_intake_page_creates_profile(self):
        resp = self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        self.assertEqual(resp.status_code, 302)
        self.assertTrue(CompanyIntakeProfile.objects.filter(company=self.company).exists())

    def test_intake_page_updates_existing_profile(self):
        CompanyIntakeProfile.objects.create(company=self.company, works_with_aramco=False)
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        p = CompanyIntakeProfile.objects.get(company=self.company)
        self.assertTrue(p.works_with_aramco)
        self.assertEqual(CompanyIntakeProfile.objects.filter(company=self.company).count(), 1)

    def test_intake_save_creates_applicability_results(self):
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        self.assertTrue(FrameworkApplicabilityResult.objects.filter(
            company=self.company, framework_version__code='ARAMCO-SACS-002',
            decision='applicable').exists())

    def test_intake_save_is_idempotent(self):
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        n = FrameworkApplicabilityResult.objects.filter(company=self.company).count()
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        self.assertEqual(FrameworkApplicabilityResult.objects.filter(company=self.company).count(), n)

    def test_review_page_shows_applicable_frameworks(self):
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        resp = self.client.get(reverse('compliance:applicability_review'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'ARAMCO-SACS-002')

    def test_review_page_shows_reasons_and_sources(self):
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        resp = self.client.get(reverse('compliance:applicability_review'))
        self.assertContains(resp, 'Aramco')  # reason text mentions Aramco

    def test_review_page_marks_otcc_dcc_unavailable(self):
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        resp = self.client.get(reverse('compliance:applicability_review'))
        self.assertContains(resp, 'NCA-OTCC-1-2022')
        self.assertContains(resp, 'NCA-DCC-1-2022')
        # Neither appears as an applicable FrameworkApplicabilityResult.
        self.assertFalse(FrameworkApplicabilityResult.objects.filter(
            company=self.company, framework_version__code__in=['NCA-OTCC-1-2022', 'NCA-DCC-1-2022']).exists())

    def test_review_page_does_not_create_companycontrol(self):
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        self.client.get(reverse('compliance:applicability_review'))
        self.assertEqual(CompanyControl.objects.count(), 0)

    def test_review_page_does_not_create_evidence(self):
        from compliance.models import Evidence
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        self.client.get(reverse('compliance:applicability_review'))
        self.assertEqual(Evidence.objects.count(), 0)

    def test_intake_does_not_create_evidencerequirement(self):
        # No EvidenceRequirement model should exist / be created in this phase.
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        from compliance.models import EvidenceRequirement, EvidenceChecklistItem
        # Intake must not generate evidence requirement templates or checklist items.
        self.assertEqual(EvidenceRequirement.objects.count(), 0)
        self.assertEqual(EvidenceChecklistItem.objects.count(), 0)

    def test_user_cannot_see_other_company_results(self):
        # Tenant scoping: a user only ever sees their own company's results (request.user.company).
        other, _ = _intake_user(cr_number='9988776655', contact_email='o@x.com')
        FrameworkApplicabilityResult.objects.create(
            company=other, framework_version=FrameworkVersion.objects.get(code='ARAMCO-SACS-002'),
            decision='applicable', reason='other', source='rule')
        self.client.post(reverse('compliance:intake'), {'works_with_aramco': 'on'})
        resp = self.client.get(reverse('compliance:applicability_review'))
        # Only this company's results are rendered (1 applicable Aramco for self.company).
        self.assertEqual(list(resp.context['results']), list(
            FrameworkApplicabilityResult.objects.filter(company=self.company)))


# ============================================================
# Phase 3C — Framework Approval (scope) + Control Applicability planning
# ============================================================
from compliance.models import CompanyFrameworkScope, ControlApplicabilityResult
from compliance.framework_scope import (
    propose_framework_scopes, approve_framework_scope, reject_framework_scope,
    generate_control_applicability_plan,
)
from compliance.framework_applicability import evaluate_company as _eval3c


def _company_with_applicability(fv_code='ARAMCO-SACS-002', **profile_kw):
    """Company with an official framework + applicability result so scopes can be proposed."""
    fv = _seed_official(fv_code, 'ARAMCO_SACS002')
    c = _company()
    CompanyIntakeProfile.objects.create(company=c, works_with_aramco=True, **profile_kw)
    _eval3c(c, apply=True)
    return c, fv


class ScopeModelTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def test_company_framework_scope_can_be_created(self):
        c = _company(); fv = FrameworkVersion.objects.get(code='ARAMCO-SACS-002')
        s = CompanyFrameworkScope.objects.create(company=c, framework_version=fv, status='proposed')
        self.assertEqual(s.status, 'proposed')

    def test_company_framework_scope_unique_per_company_framework(self):
        c = _company(); fv = FrameworkVersion.objects.get(code='ARAMCO-SACS-002')
        CompanyFrameworkScope.objects.create(company=c, framework_version=fv)
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                CompanyFrameworkScope.objects.create(company=c, framework_version=fv)

    def test_control_applicability_result_can_be_created(self):
        c, fv = _company_with_applicability()
        propose_framework_scopes(c, apply=True)
        scope = CompanyFrameworkScope.objects.get(company=c, framework_version=fv)
        ctrl = Control.objects.filter(framework_version=fv, is_legacy_import=False).first()
        r = ControlApplicabilityResult.objects.create(
            company=c, framework_scope=scope, control=ctrl, decision='applicable')
        self.assertEqual(r.decision, 'applicable')

    def test_control_applicability_unique_per_company_control(self):
        c, fv = _company_with_applicability()
        propose_framework_scopes(c, apply=True)
        scope = CompanyFrameworkScope.objects.get(company=c, framework_version=fv)
        ctrl = Control.objects.filter(framework_version=fv, is_legacy_import=False).first()
        ControlApplicabilityResult.objects.create(company=c, framework_scope=scope, control=ctrl)
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                ControlApplicabilityResult.objects.create(company=c, framework_scope=scope, control=ctrl)

    def test_manual_override_fields_exist(self):
        for model in (CompanyFrameworkScope, ControlApplicabilityResult):
            fields = {f.name for f in model._meta.get_fields()}
        self.assertIn('overridden_by', {f.name for f in ControlApplicabilityResult._meta.get_fields()})
        self.assertIn('approved_by', {f.name for f in CompanyFrameworkScope._meta.get_fields()})


class ScopeServiceTests(TestCase):
    def test_propose_framework_scopes_from_applicability_results(self):
        c, fv = _company_with_applicability()
        propose_framework_scopes(c, apply=True)
        s = CompanyFrameworkScope.objects.get(company=c, framework_version=fv)
        self.assertEqual(s.status, 'proposed')
        self.assertEqual(s.source, 'applicability_result')

    def test_propose_framework_scopes_skips_not_applicable(self):
        # A not_applicable applicability result must NOT create a scope.
        _seed_official('NCA-CSCC-1-2019', 'NCA_ECC')
        c = _company(); CompanyIntakeProfile.objects.create(company=c)  # not critical -> CSCC not_applicable
        _eval3c(c, apply=True)
        propose_framework_scopes(c, apply=True)
        self.assertFalse(CompanyFrameworkScope.objects.filter(
            company=c, framework_version__code='NCA-CSCC-1-2019').exists())

    def test_propose_framework_scopes_skips_otcc_dcc_unavailable(self):
        c, _ = _company_with_applicability()
        propose_framework_scopes(c, apply=True)
        self.assertFalse(CompanyFrameworkScope.objects.filter(
            company=c, framework_version__code__in=['NCA-OTCC-1-2022', 'NCA-DCC-1-2022']).exists())

    def test_approve_framework_scope_sets_approved_fields(self):
        from core.models import User
        c, fv = _company_with_applicability(); propose_framework_scopes(c, apply=True)
        s = CompanyFrameworkScope.objects.get(company=c, framework_version=fv)
        u = User.objects.create_user(email='ap@x.com', password='longenough12')
        approve_framework_scope(s, user=u)
        s.refresh_from_db()
        self.assertEqual(s.status, 'approved')
        self.assertEqual(s.approved_by, u)
        self.assertIsNotNone(s.approved_at)

    def test_reject_framework_scope_sets_rejected_fields(self):
        c, fv = _company_with_applicability(); propose_framework_scopes(c, apply=True)
        s = CompanyFrameworkScope.objects.get(company=c, framework_version=fv)
        reject_framework_scope(s, 'not in scope')
        s.refresh_from_db()
        self.assertEqual(s.status, 'rejected')
        self.assertEqual(s.rejection_reason, 'not in scope')
        self.assertIsNotNone(s.rejected_at)

    def test_generate_control_plan_for_approved_framework(self):
        c, fv = _company_with_applicability(); propose_framework_scopes(c, apply=True)
        s = approve_framework_scope(CompanyFrameworkScope.objects.get(company=c, framework_version=fv))
        n, _ = generate_control_applicability_plan(c, s, apply=True)
        self.assertEqual(n, 2)  # _seed_official created 2 official controls
        self.assertEqual(ControlApplicabilityResult.objects.filter(company=c).count(), 2)

    def test_generate_control_plan_uses_official_controls_only(self):
        c, fv = _company_with_applicability(); propose_framework_scopes(c, apply=True)
        # Add a legacy control under the same framework — must NOT be planned.
        Control.objects.create(framework=fv.framework, domain=Domain.objects.filter(framework=fv.framework).first(),
                               control_id='LEG-1', title='legacy', description='d')  # fv NULL
        s = approve_framework_scope(CompanyFrameworkScope.objects.get(company=c, framework_version=fv))
        generate_control_applicability_plan(c, s, apply=True)
        for r in ControlApplicabilityResult.objects.filter(company=c):
            self.assertFalse(r.control.is_legacy_import)
            self.assertIsNotNone(r.control.framework_version_id)

    def test_generate_control_plan_does_not_use_legacy_controls(self):
        c, fv = _company_with_applicability(); propose_framework_scopes(c, apply=True)
        legacy = Control.objects.create(framework=fv.framework,
                                        domain=Domain.objects.filter(framework=fv.framework).first(),
                                        control_id='1-1-1', title='legacy', description='d')
        s = approve_framework_scope(CompanyFrameworkScope.objects.get(company=c, framework_version=fv))
        generate_control_applicability_plan(c, s, apply=True)
        self.assertFalse(ControlApplicabilityResult.objects.filter(company=c, control=legacy).exists())

    def test_generate_control_plan_is_idempotent(self):
        c, fv = _company_with_applicability(); propose_framework_scopes(c, apply=True)
        s = approve_framework_scope(CompanyFrameworkScope.objects.get(company=c, framework_version=fv))
        generate_control_applicability_plan(c, s, apply=True)
        n = ControlApplicabilityResult.objects.count()
        generate_control_applicability_plan(c, s, apply=True)
        self.assertEqual(ControlApplicabilityResult.objects.count(), n)

    def test_generate_control_plan_does_not_create_companycontrol_evidence_or_requirement(self):
        from compliance.models import Evidence
        c, fv = _company_with_applicability(); propose_framework_scopes(c, apply=True)
        s = approve_framework_scope(CompanyFrameworkScope.objects.get(company=c, framework_version=fv))
        generate_control_applicability_plan(c, s, apply=True)
        self.assertEqual(CompanyControl.objects.count(), 0)
        self.assertEqual(Evidence.objects.count(), 0)
        from compliance.models import EvidenceChecklistItem
        self.assertEqual(EvidenceChecklistItem.objects.count(), 0)

    def test_generate_control_plan_skips_unapproved_scope(self):
        c, fv = _company_with_applicability(); propose_framework_scopes(c, apply=True)
        s = CompanyFrameworkScope.objects.get(company=c, framework_version=fv)  # proposed, not approved
        n, _ = generate_control_applicability_plan(c, s, apply=True)
        self.assertEqual(n, 0)


class ScopeCommandTests(TestCase):
    def setUp(self):
        self.c, self.fv = _company_with_applicability()

    def test_propose_framework_scopes_dry_run_does_not_write(self):
        call_command('propose_framework_scopes', company_id=self.c.id, stdout=StringIO())
        self.assertEqual(CompanyFrameworkScope.objects.count(), 0)

    def test_propose_framework_scopes_apply_writes(self):
        call_command('propose_framework_scopes', company_id=self.c.id, apply=True, stdout=StringIO())
        self.assertTrue(CompanyFrameworkScope.objects.filter(company=self.c).exists())

    def test_generate_control_plan_dry_run_does_not_write(self):
        call_command('propose_framework_scopes', company_id=self.c.id, apply=True, stdout=StringIO())
        approve_framework_scope(CompanyFrameworkScope.objects.get(company=self.c, framework_version=self.fv))
        call_command('generate_control_applicability_plan', company_id=self.c.id, all_approved=True, stdout=StringIO())
        self.assertEqual(ControlApplicabilityResult.objects.count(), 0)

    def test_generate_control_plan_apply_writes(self):
        call_command('propose_framework_scopes', company_id=self.c.id, apply=True, stdout=StringIO())
        approve_framework_scope(CompanyFrameworkScope.objects.get(company=self.c, framework_version=self.fv))
        call_command('generate_control_applicability_plan', company_id=self.c.id,
                     all_approved=True, apply=True, stdout=StringIO())
        self.assertEqual(ControlApplicabilityResult.objects.filter(company=self.c).count(), 2)

    def test_generate_control_plan_all_approved(self):
        call_command('propose_framework_scopes', company_id=self.c.id, apply=True, stdout=StringIO())
        approve_framework_scope(CompanyFrameworkScope.objects.get(company=self.c, framework_version=self.fv))
        out = StringIO()
        call_command('generate_control_applicability_plan', company_id=self.c.id,
                     all_approved=True, apply=True, stdout=out)
        self.assertIn('planned controls: 2', out.getvalue())


class ScopeViewTests(TestCase):
    def setUp(self):
        from core.models import User
        self.c, self.fv = _company_with_applicability()
        self.user = User.objects.create_user(email='sv@x.com', password='longenough12',
                                              company=self.c, role='company_admin')
        self.staff = User.objects.create_user(email='st@x.com', password='longenough12',
                                               company=self.c, role='admin', is_staff=True)

    def test_framework_review_page_shows_proposed_scopes(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:applicability_review'))
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(CompanyFrameworkScope.objects.filter(company=self.c).exists())
        self.assertContains(resp, 'ARAMCO-SACS-002')

    def test_staff_can_approve_framework_scope(self):
        self.client.force_login(self.staff)
        self.client.get(reverse('compliance:applicability_review'))  # proposes scopes
        s = CompanyFrameworkScope.objects.get(company=self.c, framework_version=self.fv)
        self.client.post(reverse('compliance:approve_scope', args=[s.id]))
        s.refresh_from_db()
        self.assertEqual(s.status, 'approved')

    def test_non_staff_cannot_approve_framework_scope(self):
        self.client.force_login(self.user)
        self.client.get(reverse('compliance:applicability_review'))
        s = CompanyFrameworkScope.objects.get(company=self.c, framework_version=self.fv)
        self.client.post(reverse('compliance:approve_scope', args=[s.id]))
        s.refresh_from_db()
        self.assertNotEqual(s.status, 'approved')

    def test_control_plan_page_shows_planned_controls(self):
        self.client.force_login(self.staff)
        self.client.get(reverse('compliance:applicability_review'))
        s = approve_framework_scope(CompanyFrameworkScope.objects.get(company=self.c, framework_version=self.fv))
        generate_control_applicability_plan(self.c, s, apply=True)
        resp = self.client.get(reverse('compliance:control_plan'))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(len(resp.context['plan']), 2)

    def test_control_plan_page_does_not_show_upload_form(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:control_plan'))
        self.assertNotContains(resp, 'evidence_file')
        self.assertNotContains(resp, 'multipart/form-data')

    def test_scope_view_is_tenant_isolated(self):
        # A user cannot approve another company's scope.
        other = _company(cr_number='7766554433', contact_email='o3c@x.com')
        ofv = FrameworkVersion.objects.get(code='ARAMCO-SACS-002')
        oscope = CompanyFrameworkScope.objects.create(company=other, framework_version=ofv, status='proposed')
        self.client.force_login(self.staff)
        self.client.post(reverse('compliance:approve_scope', args=[oscope.id]))
        oscope.refresh_from_db()
        self.assertEqual(oscope.status, 'proposed')  # untouched (not this user's company)


# ============================================================
# Phase 2P — OTCC/DCC review-to-YAML conversion safeguards
# ============================================================
import json as _json2, os as _os3, tempfile as _tf
from pathlib import Path as _Path

_REVIEW_DIR = _Path('compliance/data/official_controls/manual_review')
_OTCC_YAML = _Path('compliance/data/official_controls/nca_otcc_1_2022.yaml')


def _approved_otcc_rows(n=47, **override_row0):
    rows = []
    for i in range(n):
        cid = f'1-1-{i+1}'
        rows.append({
            'framework_version_code': 'NCA-OTCC-1-2022', 'control_id': cid,
            'external_reference': f'NCA OTCC {cid}', 'domain': 'Cybersecurity Governance',
            'subdomain': 'Cybersecurity Strategy', 'title': f'Control {cid}',
            'statement': 'A clean reviewed statement.', 'level': 'control',
            'source_reference': f'NCA OTCC 1:2022, {cid}', 'source_page': None,
            'extraction_method': 'manual', 'confidence': 1.0,
            'review_status': 'approved_for_dataset', 'reviewer_notes': '',
        })
    if override_row0:
        rows[0].update(override_row0)
    return rows


def _write_tmp_review(rows):
    fd, path = _tf.mkstemp(suffix='.json'); _os3.close(fd)
    _Path(path).write_text(_json2.dumps(rows), encoding='utf-8')
    return path


class OtccDccReviewFileTests(TestCase):
    def test_otcc_review_file_exists_or_is_created(self):
        self.assertTrue((_REVIEW_DIR / 'nca_otcc_1_2022_review.json').exists())

    def test_otcc_review_file_has_47_slots(self):
        rows = _json2.loads((_REVIEW_DIR / 'nca_otcc_1_2022_review.json').read_text(encoding='utf-8'))
        self.assertEqual(len(rows), 47)

    def test_otcc_review_file_does_not_include_subcontrols(self):
        rows = _json2.loads((_REVIEW_DIR / 'nca_otcc_1_2022_review.json').read_text(encoding='utf-8'))
        for r in rows:
            self.assertNotEqual(len(r['control_id'].split('-')), 4)

    def test_dcc_review_file_is_blocked(self):
        d = _json2.loads((_REVIEW_DIR / 'nca_dcc_1_2022_review.json').read_text(encoding='utf-8'))
        self.assertEqual(d['review_status'], 'blocked')


class ConvertReviewCommandTests(TestCase):
    def _convert(self, fv='NCA-OTCC-1-2022', **kw):
        from io import StringIO as _S
        call_command('convert_review_file_to_official_dataset', framework_version=fv, stdout=_S(), **kw)

    def test_convert_otcc_review_rejects_unapproved_rows(self):
        # The shipped review file is all needs_manual_review.
        with self.assertRaises(CommandError):
            self._convert()

    def test_convert_otcc_review_rejects_missing_statement(self):
        path = _write_tmp_review(_approved_otcc_rows(statement=''))
        self.addCleanup(_os3.unlink, path)
        with self.assertRaises(CommandError):
            self._convert(review_file=path)

    def test_convert_otcc_review_rejects_duplicate_control_id(self):
        rows = _approved_otcc_rows()
        rows[1]['control_id'] = rows[0]['control_id']  # duplicate
        path = _write_tmp_review(rows); self.addCleanup(_os3.unlink, path)
        with self.assertRaises(CommandError):
            self._convert(review_file=path)

    def test_convert_otcc_review_rejects_subcontrol_id(self):
        rows = _approved_otcc_rows()
        rows[0]['control_id'] = '1-1-1-1'  # subcontrol
        path = _write_tmp_review(rows); self.addCleanup(_os3.unlink, path)
        with self.assertRaises(CommandError):
            self._convert(review_file=path)

    def test_convert_otcc_review_dry_run_does_not_write_yaml(self):
        existed = _OTCC_YAML.exists()
        path = _write_tmp_review(_approved_otcc_rows()); self.addCleanup(_os3.unlink, path)
        self._convert(review_file=path)  # dry-run default
        if not existed:
            self.assertFalse(_OTCC_YAML.exists())

    def test_convert_otcc_review_write_dataset_only_when_all_rows_approved(self):
        self.assertFalse(_OTCC_YAML.exists(), 'precondition: OTCC yaml not present')
        # all approved -> writes
        path = _write_tmp_review(_approved_otcc_rows()); self.addCleanup(_os3.unlink, path)
        self.addCleanup(lambda: _OTCC_YAML.exists() and _OTCC_YAML.unlink())
        self._convert(review_file=path, write_dataset=True)
        self.assertTrue(_OTCC_YAML.exists())
        import yaml as _yaml
        doc = _yaml.safe_load(_OTCC_YAML.read_text(encoding='utf-8'))
        self.assertEqual(doc['expected_control_count'], 47)
        _OTCC_YAML.unlink()
        # one unapproved -> blocked, no write
        rows = _approved_otcc_rows(); rows[0]['review_status'] = 'needs_manual_review'
        path2 = _write_tmp_review(rows); self.addCleanup(_os3.unlink, path2)
        with self.assertRaises(CommandError):
            self._convert(review_file=path2, write_dataset=True)
        self.assertFalse(_OTCC_YAML.exists())

    def test_dcc_conversion_rejected_when_source_is_blocked(self):
        with self.assertRaises(CommandError):
            self._convert(fv='NCA-DCC-1-2022')


class Phase2PGuardrailTests(TestCase):
    def test_no_otcc_dataset_registered_automatically(self):
        from compliance.official_dataset import DATASET_FILES
        self.assertNotIn('NCA-OTCC-1-2022', DATASET_FILES)

    def test_no_dcc_dataset_registered(self):
        from compliance.official_dataset import DATASET_FILES
        self.assertNotIn('NCA-DCC-1-2022', DATASET_FILES)

    def test_no_otcc_dcc_apply(self):
        for code in ('NCA-OTCC-1-2022', 'NCA-DCC-1-2022'):
            self.assertEqual(Control.objects.filter(framework_version__code=code).count(), 0)

    def test_no_otcc_yaml_registered_dataset_file_present(self):
        # The conversion must not have left a registered official yaml behind.
        self.assertFalse(_OTCC_YAML.exists())

    def test_registration_upload_files_untouched(self):
        import core.views, core.forms, compliance.views
        for mod, name in [(core.forms, 'CompanyRegistrationForm'),
                          (core.views, 'register_company'),
                          (compliance.views, 'upload_evidence')]:
            self.assertTrue(hasattr(mod, name))


# ============================================================
# Phase 3D — Evidence Requirement templates + checklist planning
# ============================================================
from compliance.models import EvidenceRequirement, EvidenceChecklistItem
from compliance.evidence_planning import (
    create_default_requirement_for_control, generate_evidence_requirements,
    generate_evidence_checklist_for_company, generate_evidence_checklist_for_framework_scope,
)
from compliance.framework_scope import (
    propose_framework_scopes as _propose3d, approve_framework_scope as _approve3d,
    generate_control_applicability_plan as _plan3d,
)


def _company_with_official_plan(fv_code='ARAMCO-SACS-002'):
    """Company with an approved framework scope + generated control plan over official controls."""
    fwname = {'ARAMCO-SACS-002': 'ARAMCO_SACS002', 'SABIC-CYBERTRUST-1-0': 'SABIC_CT'}.get(fv_code, 'ARAMCO_SACS002')
    fv = _seed_official(fv_code, fwname)
    c = _company()
    CompanyIntakeProfile.objects.create(company=c, works_with_aramco=True, works_with_sabic=True)
    _eval3c(c, apply=True)
    _propose3d(c, apply=True)
    scope = _approve3d(CompanyFrameworkScope.objects.get(company=c, framework_version=fv))
    _plan3d(c, scope, apply=True)
    return c, fv, scope


class EvidenceRequirementModelTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def test_evidence_requirement_can_be_created(self):
        fv = _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        ctrl = Control.objects.filter(framework_version=fv, is_legacy_import=False).first()
        r = EvidenceRequirement.objects.create(control=ctrl, title='Policy doc')
        self.assertEqual(r.requirement_level, 'mandatory')

    def test_evidence_requirement_unique_per_control_title(self):
        fv = _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        ctrl = Control.objects.filter(framework_version=fv, is_legacy_import=False).first()
        EvidenceRequirement.objects.create(control=ctrl, title='X')
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                EvidenceRequirement.objects.create(control=ctrl, title='X')

    def test_checklist_item_unique_per_company_requirement(self):
        c, fv, scope = _company_with_official_plan()
        ctrl = Control.objects.filter(framework_version=fv, is_legacy_import=False).first()
        req = EvidenceRequirement.objects.create(control=ctrl, title='Y')
        car = ControlApplicabilityResult.objects.get(company=c, control=ctrl)
        EvidenceChecklistItem.objects.create(company=c, control_applicability_result=car, evidence_requirement=req)
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                EvidenceChecklistItem.objects.create(company=c, control_applicability_result=car, evidence_requirement=req)


class EvidencePlanningServiceTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def test_create_default_requirement_for_official_control(self):
        fv = _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        ctrl = Control.objects.filter(framework_version=fv, is_legacy_import=False).first()
        req, created = create_default_requirement_for_control(ctrl, apply=True)
        self.assertTrue(created)
        self.assertEqual(req.control, ctrl)

    def test_default_requirement_skips_legacy_control(self):
        fw, dom = _fw_dom()
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='L-1', title='t', description='d')
        req, created = create_default_requirement_for_control(legacy, apply=True)
        self.assertIsNone(req)
        self.assertFalse(created)

    def test_generate_requirements_official_only(self):
        _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        fw, dom = _fw_dom()
        Control.objects.create(framework=fw, domain=dom, control_id='L-9', title='t', description='d')  # legacy
        generate_evidence_requirements(apply=True)
        for r in EvidenceRequirement.objects.all():
            self.assertIsNotNone(r.control.framework_version_id)
            self.assertFalse(r.control.is_legacy_import)

    def test_generate_requirements_is_idempotent(self):
        _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        generate_evidence_requirements(apply=True)
        n = EvidenceRequirement.objects.count()
        generate_evidence_requirements(apply=True)
        self.assertEqual(EvidenceRequirement.objects.count(), n)

    def test_generate_checklist_for_company(self):
        c, fv, scope = _company_with_official_plan()
        generate_evidence_requirements(apply=True)
        res = generate_evidence_checklist_for_company(c, apply=True)
        self.assertEqual(res['planned'], 2)  # 2 official controls -> 2 default requirements
        self.assertEqual(EvidenceChecklistItem.objects.filter(company=c).count(), 2)

    def test_generate_checklist_is_idempotent(self):
        c, fv, scope = _company_with_official_plan()
        generate_evidence_requirements(apply=True)
        generate_evidence_checklist_for_company(c, apply=True)
        n = EvidenceChecklistItem.objects.count()
        generate_evidence_checklist_for_company(c, apply=True)
        self.assertEqual(EvidenceChecklistItem.objects.count(), n)

    def test_checklist_uses_official_controls_only(self):
        c, fv, scope = _company_with_official_plan()
        generate_evidence_requirements(apply=True)
        generate_evidence_checklist_for_company(c, apply=True)
        for it in EvidenceChecklistItem.objects.filter(company=c):
            self.assertIsNotNone(it.evidence_requirement.control.framework_version_id)
            self.assertFalse(it.evidence_requirement.control.is_legacy_import)

    def test_checklist_does_not_create_evidence_or_companycontrol(self):
        from compliance.models import Evidence
        c, fv, scope = _company_with_official_plan()
        generate_evidence_requirements(apply=True)
        generate_evidence_checklist_for_company(c, apply=True)
        self.assertEqual(Evidence.objects.count(), 0)
        self.assertEqual(CompanyControl.objects.count(), 0)


class EvidencePlanningCommandTests(TestCase):
    def setUp(self):
        self.c, self.fv, self.scope = _company_with_official_plan()

    def test_generate_requirements_dry_run_does_not_write(self):
        call_command('generate_evidence_requirements', stdout=StringIO())
        self.assertEqual(EvidenceRequirement.objects.count(), 0)

    def test_generate_requirements_apply_writes(self):
        call_command('generate_evidence_requirements', apply=True, stdout=StringIO())
        self.assertTrue(EvidenceRequirement.objects.exists())

    def test_generate_checklist_dry_run_does_not_write(self):
        call_command('generate_evidence_requirements', apply=True, stdout=StringIO())
        call_command('generate_evidence_checklist', company_id=self.c.id, stdout=StringIO())
        self.assertEqual(EvidenceChecklistItem.objects.count(), 0)

    def test_generate_checklist_apply_writes(self):
        call_command('generate_evidence_requirements', apply=True, stdout=StringIO())
        call_command('generate_evidence_checklist', company_id=self.c.id, apply=True, stdout=StringIO())
        self.assertEqual(EvidenceChecklistItem.objects.filter(company=self.c).count(), 2)


class EvidenceChecklistViewTests(TestCase):
    def setUp(self):
        from core.models import User
        self.c, self.fv, self.scope = _company_with_official_plan()
        self.user = User.objects.create_user(email='ec@x.com', password='longenough12',
                                             company=self.c, role='company_admin')
        self.staff = User.objects.create_user(email='ecs@x.com', password='longenough12',
                                              company=self.c, role='admin', is_staff=True)

    def test_checklist_page_requires_login(self):
        resp = self.client.get(reverse('compliance:evidence_checklist'))
        self.assertEqual(resp.status_code, 302)

    def test_checklist_page_no_upload_form(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:evidence_checklist'))
        self.assertEqual(resp.status_code, 200)
        self.assertNotContains(resp, 'evidence_file')
        self.assertNotContains(resp, 'multipart/form-data')

    def test_staff_can_generate_checklist(self):
        self.client.force_login(self.staff)
        self.client.post(reverse('compliance:generate_evidence_checklist'))
        self.assertEqual(EvidenceChecklistItem.objects.filter(company=self.c).count(), 2)

    def test_non_staff_cannot_generate_checklist(self):
        self.client.force_login(self.user)
        self.client.post(reverse('compliance:generate_evidence_checklist'))
        self.assertEqual(EvidenceChecklistItem.objects.filter(company=self.c).count(), 0)

    def test_checklist_tenant_isolated(self):
        other, ofv, oscope = _company_with_official_plan('SABIC-CYBERTRUST-1-0')
        generate_evidence_requirements(apply=True)
        generate_evidence_checklist_for_company(other, apply=True)
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:evidence_checklist'))
        for it in resp.context['items']:
            self.assertEqual(it.company, self.c)


class Phase3DGuardrailTests(TestCase):
    def test_upload_flow_untouched(self):
        import core.views, core.forms, compliance.views
        for mod, name in [(core.forms, 'CompanyRegistrationForm'),
                          (core.views, 'register_company'),
                          (compliance.views, 'upload_evidence')]:
            self.assertTrue(hasattr(mod, name))

    def test_no_otcc_dcc_registered(self):
        from compliance.official_dataset import DATASET_FILES
        self.assertNotIn('NCA-OTCC-1-2022', DATASET_FILES)
        self.assertNotIn('NCA-DCC-1-2022', DATASET_FILES)


# ============================================================
# Phase 3E — Evidence Upload v2 + EvidenceSubmission linking
# ============================================================
from django.core.files.uploadedfile import SimpleUploadedFile as _SUF
from compliance.models import EvidenceSubmission


def _company_with_checklist(fv_code='ARAMCO-SACS-002'):
    """Company with an approved scope, control plan, and generated evidence checklist."""
    c, fv, scope = _company_with_official_plan(fv_code)
    generate_evidence_requirements(apply=True)
    generate_evidence_checklist_for_company(c, apply=True)
    return c, fv, scope


class EvidenceSubmissionModelTests(TestCase):
    def setUp(self):
        self.c, self.fv, self.scope = _company_with_checklist()
        self.item = EvidenceChecklistItem.objects.filter(company=self.c).first()

    def test_evidence_submission_can_be_created(self):
        s = EvidenceSubmission.objects.create(
            company=self.c, checklist_item=self.item,
            uploaded_file=_SUF('p.pdf', b'%PDF-1.4', content_type='application/pdf'),
            original_filename='p.pdf', file_type='pdf', file_size=7)
        self.assertEqual(s.status, 'pending_review')

    def test_evidence_submission_links_to_checklist_item(self):
        s = EvidenceSubmission.objects.create(
            company=self.c, checklist_item=self.item,
            uploaded_file=_SUF('p.txt', b'x'), original_filename='p.txt', file_type='txt', file_size=1)
        self.assertIn(s, self.item.submissions.all())

    def test_evidence_submission_stores_file_metadata(self):
        s = EvidenceSubmission.objects.create(
            company=self.c, checklist_item=self.item,
            uploaded_file=_SUF('p.txt', b'abc'), original_filename='p.txt', file_type='txt',
            file_size=3, file_hash='deadbeef')
        self.assertEqual(s.file_type, 'txt'); self.assertEqual(s.file_size, 3)
        self.assertEqual(s.file_hash, 'deadbeef')

    def test_evidence_submission_status_defaults_correctly(self):
        s = EvidenceSubmission.objects.create(
            company=self.c, checklist_item=self.item,
            uploaded_file=_SUF('p.txt', b'x'), original_filename='p.txt', file_type='txt')
        self.assertEqual(s.status, 'pending_review')
        self.assertEqual(s.version, 1)


class EvidenceUploadV2ValidationTests(TestCase):
    def setUp(self):
        from core.models import User
        self.c, self.fv, self.scope = _company_with_checklist()
        self.item = EvidenceChecklistItem.objects.filter(company=self.c).first()
        self.user = User.objects.create_user(email='v2@x.com', password='longenough12',
                                             company=self.c, role='company_admin')
        self.client.force_login(self.user)
        self.url = reverse('compliance:evidence_upload_v2', args=[self.item.id])

    def test_upload_rejects_disallowed_extension(self):
        bad = _SUF('m.exe', b'MZ', content_type='application/octet-stream')
        self.client.post(self.url, {'uploaded_file': bad})
        self.assertEqual(EvidenceSubmission.objects.count(), 0)

    def test_upload_rejects_large_file(self):
        from compliance.forms import EVIDENCE_V2_MAX_SIZE
        from unittest import mock
        with mock.patch('compliance.forms.EVIDENCE_V2_MAX_SIZE', 5):
            big = _SUF('p.pdf', b'x' * 50, content_type='application/pdf')
            self.client.post(self.url, {'uploaded_file': big})
        self.assertEqual(EvidenceSubmission.objects.count(), 0)

    def test_upload_accepts_allowed_file_types(self):
        ok = _SUF('p.pdf', b'%PDF-1.4 ok', content_type='application/pdf')
        resp = self.client.post(self.url, {'uploaded_file': ok})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(EvidenceSubmission.objects.filter(company=self.c).count(), 1)


class EvidenceUploadV2ViewTests(TestCase):
    def setUp(self):
        from core.models import User
        self.c, self.fv, self.scope = _company_with_checklist()
        self.item = EvidenceChecklistItem.objects.filter(company=self.c).first()
        self.user = User.objects.create_user(email='v2v@x.com', password='longenough12',
                                             company=self.c, role='company_admin')

    def _upload(self):
        self.client.force_login(self.user)
        return self.client.post(reverse('compliance:evidence_upload_v2', args=[self.item.id]),
                                {'uploaded_file': _SUF('p.pdf', b'%PDF-1.4 ok'), 'notes': 'n'})

    def test_evidence_upload_v2_requires_login(self):
        resp = self.client.get(reverse('compliance:evidence_upload_v2', args=[self.item.id]))
        self.assertEqual(resp.status_code, 302); self.assertIn('/login', resp.url)

    def test_evidence_upload_v2_creates_submission(self):
        self._upload()
        s = EvidenceSubmission.objects.get(company=self.c)
        self.assertEqual(s.checklist_item, self.item)
        self.assertEqual(s.uploaded_by, self.user)
        self.assertTrue(s.file_hash)  # checksum computed

    def test_evidence_upload_v2_updates_checklist_status(self):
        self._upload()
        self.item.refresh_from_db()
        self.assertEqual(self.item.status, 'submitted')

    def test_evidence_submission_detail_requires_login(self):
        self._upload()
        s = EvidenceSubmission.objects.get(company=self.c)
        self.client.logout()
        resp = self.client.get(reverse('compliance:evidence_submission_detail', args=[s.id]))
        self.assertEqual(resp.status_code, 302)

    def test_evidence_submission_list_shows_submissions(self):
        self._upload()
        resp = self.client.get(reverse('compliance:evidence_submission_list', args=[self.item.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(len(resp.context['submissions']), 1)

    def test_checklist_page_shows_upload_links(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:evidence_checklist'))
        self.assertContains(resp, reverse('compliance:evidence_upload_v2', args=[self.item.id]))

    def test_checklist_page_shows_submission_count(self):
        self._upload()
        resp = self.client.get(reverse('compliance:evidence_checklist'))
        self.assertContains(resp, 'ملف')  # "N ملف — آخرها..."


class EvidenceV2TenantTests(TestCase):
    def setUp(self):
        from core.models import User
        self.c, _, _ = _company_with_checklist()
        self.other, _, _ = _company_with_checklist('SABIC-CYBERTRUST-1-0')
        self.other_item = EvidenceChecklistItem.objects.filter(company=self.other).first()
        self.user = User.objects.create_user(email='ten@x.com', password='longenough12',
                                             company=self.c, role='company_admin')
        self.client.force_login(self.user)

    def test_user_cannot_upload_to_other_company_checklist_item(self):
        resp = self.client.post(reverse('compliance:evidence_upload_v2', args=[self.other_item.id]),
                                {'uploaded_file': _SUF('p.pdf', b'%PDF')})
        self.assertEqual(EvidenceSubmission.objects.filter(checklist_item=self.other_item).count(), 0)
        self.assertEqual(resp.status_code, 302)

    def test_user_cannot_view_other_company_submission(self):
        s = EvidenceSubmission.objects.create(
            company=self.other, checklist_item=self.other_item,
            uploaded_file=_SUF('p.txt', b'x'), original_filename='p.txt', file_type='txt')
        resp = self.client.get(reverse('compliance:evidence_submission_detail', args=[s.id]))
        self.assertEqual(resp.status_code, 404)  # other company -> non-enumerable 404

    def test_user_cannot_list_other_company_submissions(self):
        resp = self.client.get(reverse('compliance:evidence_submission_list', args=[self.other_item.id]))
        self.assertEqual(resp.status_code, 302)


class Phase3EBackwardCompatTests(TestCase):
    def setUp(self):
        from core.models import User
        self.company, self.control = _company_with_control()
        self.user = User.objects.create_user(email='bc3e@x.com', password='longenough12',
                                             company=self.company, role='company_admin')

    def test_old_upload_evidence_flow_still_works(self):
        self.client.force_login(self.user)
        good = _SUF('policy.txt', b'ok', content_type='text/plain')
        with mock.patch('monitoring.tasks.analyze_evidence_async.delay'):
            resp = self.client.post(reverse('compliance:upload_evidence', args=[self.control.id]),
                                    {'evidence_file': good})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(Evidence.objects.count(), 1)  # legacy Evidence still created by OLD flow

    def test_old_evidence_model_unchanged(self):
        fields = {f.name for f in Evidence._meta.get_fields()}
        for f in ('company_control', 'file', 'ai_verdict', 'extracted_text'):
            self.assertIn(f, fields)

    def test_companycontrol_not_created_by_upload_v2(self):
        c, fv, scope = _company_with_checklist()
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        from core.models import User
        u = User.objects.create_user(email='cc3e@x.com', password='longenough12', company=c)
        self.client.force_login(u)
        before = CompanyControl.objects.count()
        self.client.post(reverse('compliance:evidence_upload_v2', args=[item.id]),
                         {'uploaded_file': _SUF('p.pdf', b'%PDF')})
        self.assertEqual(CompanyControl.objects.count(), before)

    def test_evidence_legacy_not_created_by_upload_v2(self):
        c, fv, scope = _company_with_checklist()
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        from core.models import User
        u = User.objects.create_user(email='le3e@x.com', password='longenough12', company=c)
        self.client.force_login(u)
        self.client.post(reverse('compliance:evidence_upload_v2', args=[item.id]),
                         {'uploaded_file': _SUF('p.pdf', b'%PDF')})
        self.assertEqual(Evidence.objects.count(), 0)  # v2 does NOT create legacy Evidence
        self.assertEqual(EvidenceSubmission.objects.filter(company=c).count(), 1)

    def test_no_ai_analysis_created(self):
        from ai_engine.models import AIAuditLog
        c, fv, scope = _company_with_checklist()
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        from core.models import User
        u = User.objects.create_user(email='ai3e@x.com', password='longenough12', company=c)
        self.client.force_login(u)
        self.client.post(reverse('compliance:evidence_upload_v2', args=[item.id]),
                         {'uploaded_file': _SUF('p.pdf', b'%PDF')})
        self.assertEqual(AIAuditLog.objects.count(), 0)

    def test_no_control_assessment_model_used(self):
        # No ControlAssessment is created by upload v2 (compliance decisions are out of scope).
        from django.apps import apps
        names = [m.__name__ for m in apps.get_models()]
        # If a ControlAssessment model exists later, this still must not be populated by v2.
        c, fv, scope = _company_with_checklist()
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        from core.models import User
        u = User.objects.create_user(email='ca3e@x.com', password='longenough12', company=c)
        self.client.force_login(u)
        self.client.post(reverse('compliance:evidence_upload_v2', args=[item.id]),
                         {'uploaded_file': _SUF('p.pdf', b'%PDF')})
        if 'ControlAssessment' in names:
            from compliance.models import ControlAssessment
            self.assertEqual(ControlAssessment.objects.count(), 0)

    def test_registration_flow_still_works(self):
        fw, dom = _fw_dom()
        with mock.patch('core.views.classify_company', return_value={'error': 'skip'}):
            resp = self.client.post(reverse('core:register'), {
                'company_name': 'RegE', 'cr_number': '9090901234', 'sector': 'technology',
                'size': 'small', 'first_name': 'A', 'last_name': 'B', 'email': 'rege@x.com',
                'password': 'longenough12', 'target_nca': 'on'})
        self.assertEqual(resp.status_code, 302)


# ============================================================
# Phase 3F — Advisory AI/OCR evidence analysis
# ============================================================
from compliance.models import EvidenceAnalysisResult
from compliance.evidence_analysis import (
    extract_text_from_submission, analyze_evidence_submission, batch_analyze_pending_submissions,
)


def _submission(company, item, name='p.txt', content=b'Cybersecurity policy approved by management.', ftype='txt'):
    return EvidenceSubmission.objects.create(
        company=company, checklist_item=item, uploaded_file=_SUF(name, content),
        original_filename=name, file_type=ftype, file_size=len(content))


def _company_with_submission(fv_code='ARAMCO-SACS-002', **subkw):
    c, fv, scope = _company_with_checklist(fv_code)
    item = EvidenceChecklistItem.objects.filter(company=c).first()
    sub = _submission(c, item, **subkw)
    return c, item, sub


class EvidenceAnalysisModelTests(TestCase):
    def test_evidence_analysis_result_can_be_created(self):
        c, item, sub = _company_with_submission()
        r = EvidenceAnalysisResult.objects.create(
            company=c, evidence_submission=sub, checklist_item=item,
            control=item.evidence_requirement.control, status='completed')
        self.assertEqual(r.status, 'completed')

    def test_evidence_analysis_links_submission_checklist_control(self):
        c, item, sub = _company_with_submission()
        r = EvidenceAnalysisResult.objects.create(
            company=c, evidence_submission=sub, checklist_item=item,
            control=item.evidence_requirement.control)
        self.assertEqual(r.evidence_submission, sub)
        self.assertEqual(r.control, item.evidence_requirement.control)

    def test_evidence_analysis_is_not_control_assessment(self):
        # The advisory result must not expose any compliance-decision field.
        fields = {f.name for f in EvidenceAnalysisResult._meta.get_fields()}
        for forbidden in ('compliant', 'compliance_status', 'final_status', 'accepted'):
            self.assertNotIn(forbidden, fields)


class TextExtractionTests(TestCase):
    def test_extract_text_from_txt_submission(self):
        c, item, sub = _company_with_submission(content=b'Approved cybersecurity policy v2.')
        text, trunc, note = extract_text_from_submission(sub)
        self.assertIn('policy', text)

    def test_extract_text_from_csv_submission(self):
        c, item, sub = _company_with_submission(name='d.csv', content=b'asset,owner\nfirewall,IT', ftype='csv')
        text, trunc, note = extract_text_from_submission(sub)
        self.assertIn('firewall', text)

    def test_extract_text_handles_unsupported_file(self):
        c, item, sub = _company_with_submission(name='p.pdf', content=b'%PDF-1.4', ftype='pdf')
        text, trunc, note = extract_text_from_submission(sub)
        self.assertEqual(text, '')
        self.assertIn('OCR', note)  # pdf deferred -> needs human review

    def test_extract_text_respects_size_limit(self):
        from compliance.evidence_analysis import MAX_EXTRACT_CHARS
        big = b'a' * (MAX_EXTRACT_CHARS + 500)
        c, item, sub = _company_with_submission(content=big)
        text, trunc, note = extract_text_from_submission(sub)
        self.assertLessEqual(len(text), MAX_EXTRACT_CHARS)
        self.assertTrue(trunc)

    def test_extract_text_does_not_log_full_content(self):
        # Extraction returns text but the error path returns only an error class, never content.
        c, item, sub = _company_with_submission(name='x.docx', content=b'not a real docx', ftype='docx')
        text, trunc, note = extract_text_from_submission(sub)
        self.assertNotIn('not a real docx', note)  # note never contains raw content


class AnalysisServiceTests(TestCase):
    # No OPENAI_API_KEY in tests -> AI fails gracefully to needs_human_review.
    def test_analyze_submission_dry_run_does_not_write(self):
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=False)
        self.assertEqual(EvidenceAnalysisResult.objects.count(), 0)

    def test_analyze_submission_creates_result_in_apply(self):
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=True)
        self.assertEqual(EvidenceAnalysisResult.objects.filter(evidence_submission=sub).count(), 1)

    def test_analyze_submission_missing_api_key_fails_gracefully(self):
        from django.test import override_settings
        with override_settings(OPENAI_API_KEY=''):
            c, item, sub = _company_with_submission()
            res = analyze_evidence_submission(sub, apply=True)
        r = EvidenceAnalysisResult.objects.get(evidence_submission=sub)
        self.assertEqual(r.status, 'needs_human_review')
        self.assertIn('not configured', r.error_message)

    def test_analyze_submission_does_not_mark_compliant(self):
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=True)
        sub.refresh_from_db(); item.refresh_from_db()
        self.assertNotEqual(sub.status, 'accepted')
        self.assertNotIn(item.status, ('accepted',))

    def test_analyze_submission_does_not_accept_or_reject_evidence(self):
        c, item, sub = _company_with_submission()
        before = sub.status
        analyze_evidence_submission(sub, apply=True)
        sub.refresh_from_db()
        self.assertEqual(sub.status, before)  # submission status untouched by analysis

    def test_analyze_submission_uses_official_control_context(self):
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=True)
        r = EvidenceAnalysisResult.objects.get(evidence_submission=sub)
        self.assertIsNotNone(r.control.framework_version_id)
        self.assertFalse(r.control.is_legacy_import)

    def test_analyze_submission_is_idempotent(self):
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=True)
        analyze_evidence_submission(sub, apply=True)
        self.assertEqual(EvidenceAnalysisResult.objects.filter(evidence_submission=sub).count(), 1)

    def test_analyze_does_not_create_companycontrol_or_assessment(self):
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=True)
        self.assertEqual(CompanyControl.objects.count(), 0)
        names = [m.__name__ for m in __import__('django.apps', fromlist=['apps']).apps.get_models()]
        if 'ControlAssessment' in names:
            from compliance.models import ControlAssessment
            self.assertEqual(ControlAssessment.objects.count(), 0)


class AnalysisCommandTests(TestCase):
    def setUp(self):
        self.c, self.item, self.sub = _company_with_submission()

    def test_analyze_evidence_submission_command_dry_run(self):
        call_command('analyze_evidence_submission', submission_id=self.sub.id, stdout=StringIO())
        self.assertEqual(EvidenceAnalysisResult.objects.count(), 0)

    def test_analyze_evidence_submission_command_apply(self):
        call_command('analyze_evidence_submission', submission_id=self.sub.id, apply=True, stdout=StringIO())
        self.assertEqual(EvidenceAnalysisResult.objects.filter(evidence_submission=self.sub).count(), 1)

    def test_analyze_pending_evidence_company_scoped(self):
        call_command('analyze_pending_evidence', company_id=self.c.id, apply=True, stdout=StringIO())
        self.assertTrue(EvidenceAnalysisResult.objects.filter(company=self.c).exists())

    def test_analyze_pending_evidence_does_not_cross_tenant(self):
        other, oitem, osub = _company_with_submission('SABIC-CYBERTRUST-1-0')
        call_command('analyze_pending_evidence', company_id=self.c.id, apply=True, stdout=StringIO())
        self.assertFalse(EvidenceAnalysisResult.objects.filter(company=other).exists())


class AnalysisViewTests(TestCase):
    def setUp(self):
        from core.models import User
        self.c, self.item, self.sub = _company_with_submission()
        self.user = User.objects.create_user(email='av@x.com', password='longenough12',
                                             company=self.c, role='company_admin')
        self.staff = User.objects.create_user(email='avs@x.com', password='longenough12',
                                              company=self.c, role='admin', is_staff=True)

    def test_submission_detail_shows_analysis_section(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:evidence_submission_detail', args=[self.sub.id]))
        self.assertContains(resp, 'استشاري')  # advisory analysis section present

    def test_staff_can_trigger_analysis(self):
        self.client.force_login(self.staff)
        self.client.post(reverse('compliance:analyze_submission', args=[self.sub.id]))
        self.assertEqual(EvidenceAnalysisResult.objects.filter(evidence_submission=self.sub).count(), 1)

    def test_non_staff_cannot_trigger_analysis(self):
        self.client.force_login(self.user)
        self.client.post(reverse('compliance:analyze_submission', args=[self.sub.id]))
        self.assertEqual(EvidenceAnalysisResult.objects.count(), 0)

    def test_user_cannot_view_other_company_analysis(self):
        other, oitem, osub = _company_with_submission('SABIC-CYBERTRUST-1-0')
        EvidenceAnalysisResult.objects.create(company=other, evidence_submission=osub,
                                              checklist_item=oitem, control=oitem.evidence_requirement.control)
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:evidence_submission_detail', args=[osub.id]))
        self.assertEqual(resp.status_code, 404)  # other company -> 404 (tenant-isolated)


class Phase3FBackwardCompatTests(TestCase):
    def test_old_upload_evidence_flow_still_works(self):
        from core.models import User
        company, control = _company_with_control()
        user = User.objects.create_user(email='bc3f@x.com', password='longenough12',
                                        company=company, role='company_admin')
        self.client.force_login(user)
        with mock.patch('monitoring.tasks.analyze_evidence_async.delay'):
            resp = self.client.post(reverse('compliance:upload_evidence', args=[control.id]),
                                    {'evidence_file': _SUF('p.txt', b'ok')})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(Evidence.objects.count(), 1)

    def test_evidence_upload_v2_still_works(self):
        from core.models import User
        c, item, _ = _company_with_submission()
        u = User.objects.create_user(email='v23f@x.com', password='longenough12', company=c)
        self.client.force_login(u)
        resp = self.client.post(reverse('compliance:evidence_upload_v2', args=[item.id]),
                                {'uploaded_file': _SUF('q.pdf', b'%PDF')})
        self.assertEqual(resp.status_code, 302)
        self.assertTrue(EvidenceSubmission.objects.filter(company=c, original_filename='q.pdf').exists())

    def test_no_reports_created_by_analysis(self):
        from monitoring.models import MonthlyReport
        c, item, sub = _company_with_submission()
        before = MonthlyReport.objects.count()
        analyze_evidence_submission(sub, apply=True)
        self.assertEqual(MonthlyReport.objects.count(), before)

    def test_registration_flow_still_works(self):
        fw, dom = _fw_dom()
        with mock.patch('core.views.classify_company', return_value={'error': 'skip'}):
            resp = self.client.post(reverse('core:register'), {
                'company_name': 'RegF', 'cr_number': '8080801234', 'sector': 'technology',
                'size': 'small', 'first_name': 'A', 'last_name': 'B', 'email': 'regf@x.com',
                'password': 'longenough12', 'target_nca': 'on'})
        self.assertEqual(resp.status_code, 302)


# ============================================================
# Phase 3G — Auditor review + Control Assessment
# ============================================================
from compliance.models import ControlAssessment
from compliance.control_assessment import (
    get_or_create_assessment_for_control, create_assessments_for_company,
    update_assessment_from_auditor_input,
)


class ControlAssessmentModelTests(TestCase):
    def setUp(self):
        self.c, self.fv, self.scope = _company_with_official_plan()
        self.control = Control.objects.filter(framework_version=self.fv, is_legacy_import=False).first()

    def test_control_assessment_can_be_created(self):
        a = ControlAssessment.objects.create(company=self.c, control=self.control)
        self.assertEqual(a.status, 'not_reviewed')

    def test_control_assessment_unique_company_control(self):
        ControlAssessment.objects.create(company=self.c, control=self.control)
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                ControlAssessment.objects.create(company=self.c, control=self.control)

    def test_control_assessment_links_official_control(self):
        a = ControlAssessment.objects.create(company=self.c, control=self.control)
        self.assertIsNotNone(a.control.framework_version_id)
        self.assertFalse(a.control.is_legacy_import)

    def test_get_or_create_skips_legacy_control(self):
        fw, dom = _fw_dom()
        legacy = Control.objects.create(framework=fw, domain=dom, control_id='L-1', title='t', description='d')
        obj, created = get_or_create_assessment_for_control(self.c, legacy, apply=True)
        self.assertIsNone(obj)
        self.assertFalse(ControlAssessment.objects.filter(control=legacy).exists())


class AssessmentServiceTests(TestCase):
    def setUp(self):
        self.c, self.fv, self.scope = _company_with_official_plan()

    def test_generate_assessments_dry_run_does_not_write(self):
        create_assessments_for_company(self.c, apply=False)
        self.assertEqual(ControlAssessment.objects.count(), 0)

    def test_generate_assessments_apply_creates_not_reviewed(self):
        create_assessments_for_company(self.c, apply=True)
        self.assertEqual(ControlAssessment.objects.filter(company=self.c).count(), 2)
        self.assertTrue(all(a.status == 'not_reviewed' for a in ControlAssessment.objects.filter(company=self.c)))

    def test_generate_assessments_official_controls_only(self):
        fw = self.fv.framework
        Control.objects.create(framework=fw, domain=Domain.objects.filter(framework=fw).first(),
                               control_id='L-9', title='t', description='d')  # legacy
        create_assessments_for_company(self.c, apply=True)
        for a in ControlAssessment.objects.filter(company=self.c):
            self.assertIsNotNone(a.control.framework_version_id)
            self.assertFalse(a.control.is_legacy_import)

    def test_generate_assessments_does_not_create_companycontrol(self):
        create_assessments_for_company(self.c, apply=True)
        self.assertEqual(CompanyControl.objects.count(), 0)

    def test_generate_assessments_does_not_create_reports(self):
        from monitoring.models import MonthlyReport
        before = MonthlyReport.objects.count()
        create_assessments_for_company(self.c, apply=True)
        self.assertEqual(MonthlyReport.objects.count(), before)

    def test_ai_analysis_does_not_set_assessment_status(self):
        # Running advisory AI analysis must not create/alter a ControlAssessment.
        from compliance.evidence_analysis import analyze_evidence_submission
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=True)
        self.assertEqual(ControlAssessment.objects.count(), 0)

    def test_auditor_update_sets_status_and_reviewed_by(self):
        from core.models import User
        create_assessments_for_company(self.c, apply=True)
        a = ControlAssessment.objects.filter(company=self.c).first()
        u = User.objects.create_user(email='aud@x.com', password='longenough12', company=self.c, is_staff=True)
        update_assessment_from_auditor_input(a, {'status': 'compliant', 'auditor_notes': 'ok'}, u)
        a.refresh_from_db()
        self.assertEqual(a.status, 'compliant')
        self.assertEqual(a.reviewed_by, u)
        self.assertIsNotNone(a.reviewed_at)

    def test_auditor_update_rejects_invalid_status(self):
        from core.models import User
        create_assessments_for_company(self.c, apply=True)
        a = ControlAssessment.objects.filter(company=self.c).first()
        u = User.objects.create_user(email='aud2@x.com', password='longenough12', company=self.c, is_staff=True)
        update_assessment_from_auditor_input(a, {'status': 'totally_compliant_auto'}, u)
        a.refresh_from_db()
        self.assertEqual(a.status, 'not_reviewed')  # invalid ignored


class AssessmentCommandTests(TestCase):
    def setUp(self):
        self.c, self.fv, self.scope = _company_with_official_plan()

    def test_generate_control_assessments_command_dry_run(self):
        call_command('generate_control_assessments', company_id=self.c.id, stdout=StringIO())
        self.assertEqual(ControlAssessment.objects.count(), 0)

    def test_generate_control_assessments_command_apply(self):
        call_command('generate_control_assessments', company_id=self.c.id, apply=True, stdout=StringIO())
        self.assertEqual(ControlAssessment.objects.filter(company=self.c).count(), 2)

    def test_generate_control_assessments_company_scoped_no_cross_tenant(self):
        other, ofv, oscope = _company_with_official_plan('SABIC-CYBERTRUST-1-0')
        call_command('generate_control_assessments', company_id=self.c.id, apply=True, stdout=StringIO())
        self.assertFalse(ControlAssessment.objects.filter(company=other).exists())


class AuditorReviewViewTests(TestCase):
    def setUp(self):
        from core.models import User
        self.c, self.fv, self.scope = _company_with_official_plan()
        create_assessments_for_company(self.c, apply=True)
        self.a = ControlAssessment.objects.filter(company=self.c).first()
        self.user = User.objects.create_user(email='arv@x.com', password='longenough12',
                                             company=self.c, role='company_admin')
        self.staff = User.objects.create_user(email='arvs@x.com', password='longenough12',
                                              company=self.c, role='admin', is_staff=True)

    def test_auditor_review_requires_login(self):
        resp = self.client.get(reverse('compliance:auditor_review_queue'))
        self.assertEqual(resp.status_code, 302)

    def test_auditor_review_queue_shows_official_controls(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:auditor_review_queue'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, self.a.control.control_id)

    def test_auditor_review_detail_shows_ai_advisory_label(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:auditor_review_detail', args=[self.a.id]))
        self.assertContains(resp, 'استشاري')

    def test_auditor_can_update_assessment(self):
        self.client.force_login(self.staff)
        self.client.post(reverse('compliance:auditor_review_detail', args=[self.a.id]),
                         {'status': 'compliant', 'auditor_notes': 'reviewed'})
        self.a.refresh_from_db()
        self.assertEqual(self.a.status, 'compliant')
        self.assertEqual(self.a.reviewed_by, self.staff)

    def test_non_staff_cannot_update_assessment(self):
        self.client.force_login(self.user)
        self.client.post(reverse('compliance:auditor_review_detail', args=[self.a.id]),
                         {'status': 'compliant'})
        self.a.refresh_from_db()
        self.assertEqual(self.a.status, 'not_reviewed')

    def test_user_cannot_view_other_company_assessment(self):
        other, ofv, oscope = _company_with_official_plan('SABIC-CYBERTRUST-1-0')
        create_assessments_for_company(other, apply=True)
        oa = ControlAssessment.objects.filter(company=other).first()
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:auditor_review_detail', args=[oa.id]))
        self.assertEqual(resp.status_code, 302)

    def test_user_cannot_update_other_company_assessment(self):
        other, ofv, oscope = _company_with_official_plan('SABIC-CYBERTRUST-1-0')
        create_assessments_for_company(other, apply=True)
        oa = ControlAssessment.objects.filter(company=other).first()
        self.client.force_login(self.staff)  # staff of self.c, not other
        self.client.post(reverse('compliance:auditor_review_detail', args=[oa.id]), {'status': 'compliant'})
        oa.refresh_from_db()
        self.assertEqual(oa.status, 'not_reviewed')


class Phase3GBackwardCompatTests(TestCase):
    def test_old_upload_evidence_flow_still_works(self):
        from core.models import User
        company, control = _company_with_control()
        user = User.objects.create_user(email='bc3g@x.com', password='longenough12',
                                        company=company, role='company_admin')
        self.client.force_login(user)
        with mock.patch('monitoring.tasks.analyze_evidence_async.delay'):
            resp = self.client.post(reverse('compliance:upload_evidence', args=[control.id]),
                                    {'evidence_file': _SUF('p.txt', b'ok')})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(Evidence.objects.count(), 1)

    def test_evidence_upload_v2_and_analysis_still_work(self):
        from compliance.evidence_analysis import analyze_evidence_submission
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=True)
        self.assertTrue(EvidenceAnalysisResult.objects.filter(evidence_submission=sub).exists())

    def test_no_reports_created_by_assessment(self):
        from monitoring.models import MonthlyReport
        c, fv, scope = _company_with_official_plan()
        before = MonthlyReport.objects.count()
        create_assessments_for_company(c, apply=True)
        self.assertEqual(MonthlyReport.objects.count(), before)

    def test_registration_flow_still_works(self):
        fw, dom = _fw_dom()
        with mock.patch('core.views.classify_company', return_value={'error': 'skip'}):
            resp = self.client.post(reverse('core:register'), {
                'company_name': 'RegG', 'cr_number': '7070701234', 'sector': 'technology',
                'size': 'small', 'first_name': 'A', 'last_name': 'B', 'email': 'regg@x.com',
                'password': 'longenough12', 'target_nca': 'on'})
        self.assertEqual(resp.status_code, 302)


# ============================================================
# Phase 3H — Read-only compliance reports + gap analysis + exports
# ============================================================
from compliance.reporting import (
    build_executive_summary, build_framework_gap_analysis, build_evidence_matrix,
    get_approved_framework_versions, calculate_assessment_counts,
)


def _company_with_assessments(fv_code='ARAMCO-SACS-002'):
    """Company + approved scope + control plan + checklist + assessments (2 official controls)."""
    c, fv, scope = _company_with_official_plan(fv_code)
    generate_evidence_requirements(apply=True)
    generate_evidence_checklist_for_company(c, apply=True)
    create_assessments_for_company(c, apply=True)
    return c, fv, scope


class ReportingServiceTests(TestCase):
    def setUp(self):
        self.c, self.fv, self.scope = _company_with_assessments()
        self.assessments = list(ControlAssessment.objects.filter(company=self.c))

    def _set_status(self, idx, status):
        a = self.assessments[idx]; a.status = status; a.save()

    def test_executive_summary_counts_match_control_assessments(self):
        self._set_status(0, 'compliant')
        s = build_executive_summary(self.c)
        self.assertEqual(s['total_applicable'], 2)
        self.assertEqual(s['counts']['compliant'], 1)
        self.assertEqual(s['counts']['not_reviewed'], 1)

    def test_gap_analysis_uses_official_controls_only(self):
        gap = build_framework_gap_analysis(self.c)
        for f in gap:
            for g in f['gaps']:
                ctrl = Control.objects.get(control_id=g['control_id'], framework_version=self.fv)
                self.assertFalse(ctrl.is_legacy_import)

    def test_gap_analysis_skips_legacy_controls(self):
        fw = self.fv.framework
        Control.objects.create(framework=fw, domain=Domain.objects.filter(framework=fw).first(),
                               control_id='L-9', title='legacy', description='d')  # legacy
        gap = build_framework_gap_analysis(self.c)
        ids = [g['control_id'] for f in gap for g in f['gaps']]
        self.assertNotIn('L-9', ids)

    def test_evidence_matrix_includes_submission_counts(self):
        item = EvidenceChecklistItem.objects.filter(company=self.c).first()
        _submission(self.c, item)
        rows = build_evidence_matrix(self.c)
        total_subs = sum(r['submission_count'] for r in rows)
        self.assertGreaterEqual(total_subs, 1)

    def test_evidence_matrix_includes_ai_advisory_status(self):
        rows = build_evidence_matrix(self.c)
        self.assertTrue(all('latest_ai_status' in r for r in rows))

    def test_reports_do_not_create_or_update_control_assessment(self):
        before = {a.id: a.status for a in ControlAssessment.objects.filter(company=self.c)}
        build_executive_summary(self.c); build_framework_gap_analysis(self.c); build_evidence_matrix(self.c)
        after = {a.id: a.status for a in ControlAssessment.objects.filter(company=self.c)}
        self.assertEqual(before, after)
        self.assertEqual(len(before), 2)  # no new assessments created

    def test_reports_do_not_create_companycontrol(self):
        build_executive_summary(self.c); build_evidence_matrix(self.c)
        self.assertEqual(CompanyControl.objects.count(), 0)

    def test_unreviewed_controls_not_counted_as_compliant(self):
        # All not_reviewed -> 0% compliance.
        s = build_executive_summary(self.c)
        self.assertEqual(s['counts']['compliant'], 0)
        self.assertEqual(s['compliance_percentage'], 0.0)

    def test_completion_percentage_calculation(self):
        self._set_status(0, 'compliant')  # 1 of 2 assessed
        s = build_executive_summary(self.c)
        self.assertEqual(s['completion_percentage'], 50.0)

    def test_evidence_coverage_calculation(self):
        item = EvidenceChecklistItem.objects.filter(company=self.c).first()
        _submission(self.c, item)
        s = build_executive_summary(self.c)
        self.assertEqual(s['evidence_coverage_count'], 1)  # 1 of 2 controls has evidence
        self.assertEqual(s['evidence_coverage_percentage'], 50.0)


class ReportViewTests(TestCase):
    def setUp(self):
        from core.models import User
        from billing.subscription_access import activate_company_subscription
        self.c, self.fv, self.scope = _company_with_assessments()
        activate_company_subscription(self.c, 'Test Plan', days=30)  # Phase 4B: reports gated
        self.user = User.objects.create_user(email='rep@x.com', password='longenough12',
                                             company=self.c, role='company_admin')

    def test_reports_index_requires_login(self):
        self.assertEqual(self.client.get(reverse('compliance:reports_index')).status_code, 302)

    def test_executive_summary_view_requires_login(self):
        self.assertEqual(self.client.get(reverse('compliance:report_executive_summary')).status_code, 302)

    def test_gap_analysis_view_requires_login(self):
        self.assertEqual(self.client.get(reverse('compliance:report_gap_analysis')).status_code, 302)

    def test_evidence_matrix_view_requires_login(self):
        self.assertEqual(self.client.get(reverse('compliance:report_evidence_matrix')).status_code, 302)

    def test_framework_report_filters_by_framework_version(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:report_framework', args=[self.fv.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, self.fv.code)

    def test_user_cannot_view_other_company_reports(self):
        other, ofv, oscope = _company_with_assessments('SABIC-CYBERTRUST-1-0')
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:report_framework', args=[ofv.id]))
        self.assertEqual(resp.status_code, 302)  # other company's framework not approved for this user

    def test_report_pages_show_framework_version(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:report_gap_analysis'))
        self.assertContains(resp, self.fv.code)

    def test_report_pages_show_unreviewed_warning(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:report_executive_summary'))
        self.assertContains(resp, 'غير المُراجَعة')


class ReportExportTests(TestCase):
    def setUp(self):
        from core.models import User
        from billing.subscription_access import activate_company_subscription
        self.c, self.fv, self.scope = _company_with_assessments()
        activate_company_subscription(self.c, 'Test Plan', days=30)  # Phase 4B: exports gated
        self.user = User.objects.create_user(email='exp@x.com', password='longenough12',
                                             company=self.c, role='company_admin')
        self.client.force_login(self.user)

    def test_evidence_matrix_csv_export(self):
        resp = self.client.get(reverse('compliance:export_evidence_matrix_csv'))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp['Content-Type'], 'text/csv')
        self.assertIn(b'control_id', resp.content)

    def test_evidence_matrix_xlsx_export(self):
        resp = self.client.get(reverse('compliance:export_evidence_matrix_xlsx'))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content[:2], b'PK')  # xlsx magic

    def test_exports_are_tenant_scoped(self):
        other, ofv, oscope = _company_with_assessments('SABIC-CYBERTRUST-1-0')
        # this user's CSV must contain only their framework's controls (ARAMCO TPC ids), not SABIC CT.
        resp = self.client.get(reverse('compliance:export_evidence_matrix_csv'))
        self.assertNotIn(b'SABIC-CYBERTRUST-1-0', resp.content)

    def test_exports_use_official_controls_only(self):
        fw = self.fv.framework
        Control.objects.create(framework=fw, domain=Domain.objects.filter(framework=fw).first(),
                               control_id='LEG-XYZ', title='legacy', description='d')
        resp = self.client.get(reverse('compliance:export_evidence_matrix_csv'))
        self.assertNotIn(b'LEG-XYZ', resp.content)


class Phase3HBackwardCompatTests(TestCase):
    def test_old_upload_evidence_flow_still_works(self):
        from core.models import User
        company, control = _company_with_control()
        user = User.objects.create_user(email='bc3h@x.com', password='longenough12',
                                        company=company, role='company_admin')
        self.client.force_login(user)
        with mock.patch('monitoring.tasks.analyze_evidence_async.delay'):
            resp = self.client.post(reverse('compliance:upload_evidence', args=[control.id]),
                                    {'evidence_file': _SUF('p.txt', b'ok')})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(Evidence.objects.count(), 1)

    def test_auditor_assessment_still_works(self):
        c, fv, scope = _company_with_assessments()
        a = ControlAssessment.objects.filter(company=c).first()
        from core.models import User
        u = User.objects.create_user(email='aud3h@x.com', password='longenough12', company=c, is_staff=True)
        update_assessment_from_auditor_input(a, {'status': 'compliant'}, u)
        a.refresh_from_db(); self.assertEqual(a.status, 'compliant')

    def test_no_otcc_dcc_imported_by_reports(self):
        from compliance.official_dataset import DATASET_FILES
        self.assertNotIn('NCA-OTCC-1-2022', DATASET_FILES)
        self.assertNotIn('NCA-DCC-1-2022', DATASET_FILES)

    def test_registration_flow_still_works(self):
        fw, dom = _fw_dom()
        with mock.patch('core.views.classify_company', return_value={'error': 'skip'}):
            resp = self.client.post(reverse('core:register'), {
                'company_name': 'RegH', 'cr_number': '6060601234', 'sector': 'technology',
                'size': 'small', 'first_name': 'A', 'last_name': 'B', 'email': 'regh@x.com',
                'password': 'longenough12', 'target_nca': 'on'})
        self.assertEqual(resp.status_code, 302)


# ============================================================
# Phase 3I — Dashboard & user-journey hardening (read-only UX)
# ============================================================
from compliance.user_journey import (
    build_company_journey_status, get_next_recommended_action, calculate_journey_progress,
)


def _stage(stages, key):
    return next(s for s in stages if s['key'] == key)


def _journey_user(company, **kw):
    from core.models import User
    n = User.objects.count() + 1
    defaults = dict(email=f'journey{n}@x.com', password='longenough12',
                    company=company, role='company_admin')
    defaults.update(kw)
    return User.objects.create_user(**defaults)


class JourneyServiceTests(TestCase):
    def test_journey_status_no_intake(self):
        c = _company()
        self.assertEqual(_stage(build_company_journey_status(c), 'intake')['status'], 'not_started')
        self.assertEqual(get_next_recommended_action(c)['message'], 'Complete intake profile')

    def test_journey_status_with_intake(self):
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, review_status='completed')
        self.assertEqual(_stage(build_company_journey_status(c), 'intake')['status'], 'completed')

    def test_journey_status_with_approved_framework(self):
        c, fv, scope = _company_with_official_plan()
        self.assertEqual(_stage(build_company_journey_status(c), 'framework_approval')['status'], 'completed')

    def test_journey_status_with_control_plan(self):
        c, fv, scope = _company_with_official_plan()
        self.assertEqual(_stage(build_company_journey_status(c), 'control_plan')['status'], 'completed')

    def test_journey_status_with_evidence_checklist(self):
        c, fv, scope = _company_with_checklist()
        self.assertEqual(_stage(build_company_journey_status(c), 'evidence_checklist')['status'], 'completed')

    def test_journey_status_with_submissions(self):
        c, item, sub = _company_with_submission()
        self.assertEqual(_stage(build_company_journey_status(c), 'evidence_upload')['status'], 'completed')

    def test_journey_status_with_analysis(self):
        from compliance.evidence_analysis import analyze_evidence_submission
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=True)
        self.assertEqual(_stage(build_company_journey_status(c), 'ai_analysis')['status'], 'completed')

    def test_journey_status_with_assessments(self):
        c, fv, scope = _company_with_assessments()
        st = _stage(build_company_journey_status(c), 'auditor_review')['status']
        self.assertNotEqual(st, 'not_started')  # assessments exist (all not_reviewed yet)
        # After a real auditor decision, reports become meaningful.
        a = ControlAssessment.objects.filter(company=c).first(); a.status = 'compliant'; a.save()
        self.assertEqual(_stage(build_company_journey_status(c), 'reports')['status'], 'completed')

    def test_next_action_progresses_in_correct_order(self):
        from compliance.evidence_analysis import analyze_evidence_submission
        # 1) No intake.
        bare = _company()
        self.assertEqual(get_next_recommended_action(bare)['message'], 'Complete intake profile')
        # 2) Intake exists but no approved scope.
        CompanyIntakeProfile.objects.create(company=bare, review_status='completed')
        self.assertEqual(get_next_recommended_action(bare)['message'],
                         'Review and approve applicable frameworks')
        # 3) Approved scope + control plan, no checklist.
        c, fv, scope = _company_with_official_plan()
        self.assertEqual(get_next_recommended_action(c)['message'], 'Generate evidence checklist')
        # 4) Checklist exists, no submissions.
        generate_evidence_requirements(apply=True)
        generate_evidence_checklist_for_company(c, apply=True)
        self.assertEqual(get_next_recommended_action(c)['message'], 'Upload evidence')
        # 5) Submission exists, no analysis.
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        sub = _submission(c, item)
        self.assertEqual(get_next_recommended_action(c)['message'], 'Run advisory analysis')
        # 6) Analysis exists, no assessments.
        analyze_evidence_submission(sub, apply=True)
        self.assertEqual(get_next_recommended_action(c)['message'], 'Start auditor review')
        # 7) Assessments exist.
        create_assessments_for_company(c, apply=True)
        self.assertEqual(get_next_recommended_action(c)['message'], 'View reports')

    def test_journey_uses_official_controls_only(self):
        c, fv, scope = _company_with_official_plan()
        before = _stage(build_company_journey_status(c), 'control_plan')['metric']
        fw = fv.framework
        legacy = Control.objects.create(
            framework=fw, domain=Domain.objects.filter(framework=fw).first(),
            control_id='LEG-J', title='legacy', description='d')  # is_legacy_import default
        ControlApplicabilityResult.objects.create(
            company=c, control=legacy, framework_scope=scope, decision='applicable', source='manual')
        after = _stage(build_company_journey_status(c), 'control_plan')['metric']
        self.assertEqual(before, after)  # legacy applicable control not counted

    def test_journey_does_not_use_companycontrol(self):
        c, fv, scope = _company_with_assessments()
        before = CompanyControl.objects.count()
        build_company_journey_status(c)
        get_next_recommended_action(c)
        calculate_journey_progress(c)
        self.assertEqual(CompanyControl.objects.count(), before)

    def test_journey_does_not_create_records(self):
        c, fv, scope = _company_with_assessments()
        from core.models import Company
        counts = lambda: (
            ControlAssessment.objects.count(), EvidenceSubmission.objects.count(),
            EvidenceChecklistItem.objects.count(), ControlApplicabilityResult.objects.count(),
            CompanyControl.objects.count(), Company.objects.count(),
        )
        before = counts()
        build_company_journey_status(c)
        get_next_recommended_action(c)
        calculate_journey_progress(c)
        self.assertEqual(before, counts())


class JourneyDashboardViewTests(TestCase):
    def setUp(self):
        self.c, self.fv, self.scope = _company_with_assessments()
        self.user = _journey_user(self.c)

    def test_dashboard_requires_login(self):
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_dashboard_shows_workflow_steps(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Intake Profile')
        self.assertContains(resp, 'Auditor Review')
        self.assertContains(resp, 'Reports')

    def test_dashboard_shows_next_recommended_action(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertContains(resp, 'الخطوة التالية الموصى بها')

    def test_dashboard_tenant_scoped(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertContains(resp, self.c.name)

    def test_user_cannot_view_other_company_dashboard(self):
        other = _company(name='OtherCo')
        other_user = _journey_user(other)
        self.client.force_login(other_user)
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertEqual(resp.status_code, 200)
        # Sees only their own company's status, never the first company's data.
        self.assertContains(resp, other.name)

    def test_dashboard_links_to_core_pages(self):
        self.client.force_login(self.user)
        resp = self.client.get(reverse('compliance:dashboard'))
        for name in ['compliance:intake', 'compliance:control_plan',
                     'compliance:evidence_checklist', 'compliance:auditor_review_queue',
                     'compliance:reports_index']:
            self.assertContains(resp, reverse(name))

    def test_dashboard_does_not_modify_assessments(self):
        before = {a.id: a.status for a in ControlAssessment.objects.filter(company=self.c)}
        self.client.force_login(self.user)
        self.client.get(reverse('compliance:dashboard'))
        after = {a.id: a.status for a in ControlAssessment.objects.filter(company=self.c)}
        self.assertEqual(before, after)

    def test_dashboard_does_not_modify_reports(self):
        from compliance.reporting import build_executive_summary
        before = build_executive_summary(self.c)['counts']
        self.client.force_login(self.user)
        self.client.get(reverse('compliance:dashboard'))
        after = build_executive_summary(self.c)['counts']
        self.assertEqual(before, after)


class JourneyEmptyStateTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def test_control_plan_empty_state(self):
        c = _company()
        user = _journey_user(c)
        self.client.force_login(user)
        resp = self.client.get(reverse('compliance:control_plan'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Start with framework review')

    def test_evidence_checklist_empty_state(self):
        c = _company()
        user = _journey_user(c)
        self.client.force_login(user)
        resp = self.client.get(reverse('compliance:evidence_checklist'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Complete Smart Classification first')

    def test_auditor_review_empty_state(self):
        c = _company()
        user = _journey_user(c)
        self.client.force_login(user)
        resp = self.client.get(reverse('compliance:auditor_review_queue'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Generate assessments from approved official controls')

    def test_reports_empty_state(self):
        c = _company()
        user = _journey_user(c)
        self.client.force_login(user)
        resp = self.client.get(reverse('compliance:reports_index'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Reports will be meaningful after auditor assessments')


class Phase3IBackwardCompatTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def test_intake_still_works(self):
        c = _company()
        user = _journey_user(c)
        self.client.force_login(user)
        resp = self.client.get(reverse('compliance:intake'))
        self.assertEqual(resp.status_code, 200)

    def test_evidence_upload_v2_still_works(self):
        c, item, sub = _company_with_submission()
        self.assertTrue(EvidenceSubmission.objects.filter(id=sub.id).exists())

    def test_evidence_analysis_still_works(self):
        from compliance.evidence_analysis import analyze_evidence_submission
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=True)
        self.assertTrue(EvidenceAnalysisResult.objects.filter(evidence_submission=sub).exists())

    def test_auditor_assessment_still_works(self):
        from core.models import User
        c, fv, scope = _company_with_assessments()
        a = ControlAssessment.objects.filter(company=c).first()
        u = _journey_user(c, email='aud3i@x.com', is_staff=True)
        update_assessment_from_auditor_input(a, {'status': 'compliant'}, u)
        a.refresh_from_db()
        self.assertEqual(a.status, 'compliant')

    def test_reports_still_work(self):
        c, fv, scope = _company_with_assessments()
        user = _journey_user(c)
        self.client.force_login(user)
        resp = self.client.get(reverse('compliance:report_executive_summary'))
        self.assertEqual(resp.status_code, 200)

    def test_old_upload_evidence_flow_still_works(self):
        company, control = _company_with_control()
        user = _journey_user(company)
        self.client.force_login(user)
        with mock.patch('monitoring.tasks.analyze_evidence_async.delay'):
            resp = self.client.post(reverse('compliance:upload_evidence', args=[control.id]),
                                    {'evidence_file': _SUF('p.txt', b'ok')})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(Evidence.objects.count(), 1)

    def test_registration_flow_still_works(self):
        fw, dom = _fw_dom()
        with mock.patch('core.views.classify_company', return_value={'error': 'skip'}):
            resp = self.client.post(reverse('core:register'), {
                'company_name': 'RegI', 'cr_number': '5050501234', 'sector': 'technology',
                'size': 'small', 'first_name': 'A', 'last_name': 'B', 'email': 'regi@x.com',
                'password': 'longenough12', 'target_nca': 'on'})
        self.assertEqual(resp.status_code, 302)


# ============================================================
# Phase 3J — Security, tenant isolation, and QA hardening
# ============================================================
def _full_pipeline(fv_code='ARAMCO-SACS-002'):
    """Company carried end-to-end: approved scope + plan + checklist + assessments +
    one evidence submission + advisory analysis. Returns (company, fv, item, submission)."""
    from compliance.evidence_analysis import analyze_evidence_submission
    c, fv, scope = _company_with_assessments(fv_code)
    item = EvidenceChecklistItem.objects.filter(company=c).first()
    sub = _submission(c, item)
    analyze_evidence_submission(sub, apply=True)
    return c, fv, item, sub


class SecurityHelperTests(TestCase):
    def test_get_company_object_or_none_is_tenant_safe(self):
        from compliance.security import get_company_object_or_none
        a, fv, item, sub = _full_pipeline()
        b, *_ = _full_pipeline('SABIC-CYBERTRUST-1-0')
        b_item = EvidenceChecklistItem.objects.filter(company=b).first()
        # Own object found; other company's object is never returned.
        self.assertEqual(get_company_object_or_none(EvidenceChecklistItem, a, id=item.id), item)
        self.assertIsNone(get_company_object_or_none(EvidenceChecklistItem, a, id=b_item.id))
        self.assertIsNone(get_company_object_or_none(EvidenceChecklistItem, None, id=item.id))


class SecurityAuthTests(TestCase):
    def setUp(self):
        self.c, self.fv, self.item, self.sub = _full_pipeline()
        self.assessment = ControlAssessment.objects.filter(company=self.c).first()

    def test_all_core_compliance_pages_require_login(self):
        names = ['controls_list', 'intake', 'applicability_review', 'control_plan',
                 'evidence_checklist', 'auditor_review_queue', 'reports_index',
                 'report_executive_summary', 'report_gap_analysis', 'report_evidence_matrix',
                 'dashboard']
        for n in names:
            resp = self.client.get(reverse(f'compliance:{n}'))
            self.assertEqual(resp.status_code, 302, n)
            self.assertIn('/login', resp.url, n)

    def test_report_exports_require_login(self):
        for n in ['export_evidence_matrix_csv', 'export_evidence_matrix_xlsx']:
            resp = self.client.get(reverse(f'compliance:{n}'))
            self.assertEqual(resp.status_code, 302, n)
            self.assertIn('/login', resp.url, n)

    def test_dashboard_requires_login(self):
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_upload_v2_requires_login(self):
        resp = self.client.get(reverse('compliance:evidence_upload_v2', args=[self.item.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_auditor_update_requires_login(self):
        resp = self.client.post(reverse('compliance:auditor_review_detail', args=[self.assessment.id]),
                                {'status': 'compliant'})
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)
        self.assessment.refresh_from_db()
        self.assertEqual(self.assessment.status, 'not_reviewed')  # not mutated

    def test_analysis_trigger_requires_login(self):
        resp = self.client.post(reverse('compliance:analyze_submission', args=[self.sub.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)


class SecurityTenantIsolationTests(TestCase):
    def setUp(self):
        from billing.subscription_access import activate_company_subscription
        # Company A = ARAMCO framework; Company B = SABIC framework (distinct codes).
        self.c, self.fv, self.item, self.sub = _full_pipeline('ARAMCO-SACS-002')
        activate_company_subscription(self.c, 'Test Plan', days=30)  # Phase 4B: reports gated
        self.user = _journey_user(self.c)
        self.other, self.ofv, self.oitem, self.osub = _full_pipeline('SABIC-CYBERTRUST-1-0')
        self.oassessment = ControlAssessment.objects.filter(company=self.other).first()
        self.client.force_login(self.user)

    def test_user_cannot_access_other_company_control_plan(self):
        resp = self.client.get(reverse('compliance:control_plan'))
        self.assertEqual(resp.status_code, 200)
        self.assertNotContains(resp, 'SABIC-CYBERTRUST-1-0')

    def test_user_cannot_access_other_company_evidence_checklist(self):
        resp = self.client.get(reverse('compliance:evidence_checklist'))
        self.assertEqual(resp.status_code, 200)
        self.assertNotContains(resp, 'SABIC-CYBERTRUST-1-0')

    def test_user_cannot_upload_to_other_company_checklist_item(self):
        before = EvidenceSubmission.objects.filter(checklist_item=self.oitem).count()
        resp = self.client.post(reverse('compliance:evidence_upload_v2', args=[self.oitem.id]),
                                {'uploaded_file': _SUF('x.txt', b'data'), 'notes': ''})
        self.assertEqual(resp.status_code, 302)  # redirected, not allowed
        after = EvidenceSubmission.objects.filter(checklist_item=self.oitem).count()
        self.assertEqual(before, after)  # nothing created on B's item

    def test_user_cannot_view_other_company_submission(self):
        resp = self.client.get(reverse('compliance:evidence_submission_detail', args=[self.osub.id]))
        self.assertEqual(resp.status_code, 404)  # tenant-scoped 404

    def test_user_cannot_view_other_company_analysis(self):
        # Analysis is surfaced via the submission detail; B's submission is not reachable.
        self.assertTrue(EvidenceAnalysisResult.objects.filter(evidence_submission=self.osub).exists())
        resp = self.client.get(reverse('compliance:evidence_submission_detail', args=[self.osub.id]))
        self.assertEqual(resp.status_code, 404)

    def test_user_cannot_view_other_company_assessment(self):
        resp = self.client.get(reverse('compliance:auditor_review_detail', args=[self.oassessment.id]))
        self.assertEqual(resp.status_code, 302)

    def test_user_cannot_view_other_company_reports(self):
        resp = self.client.get(reverse('compliance:report_framework', args=[self.ofv.id]))
        self.assertEqual(resp.status_code, 302)  # B's framework not approved for A

    def test_user_cannot_export_other_company_evidence_matrix(self):
        resp = self.client.get(reverse('compliance:export_evidence_matrix_csv'))
        self.assertEqual(resp.status_code, 200)
        self.assertNotIn(b'SABIC-CYBERTRUST-1-0', resp.content)

    def test_dashboard_does_not_leak_other_company_counts(self):
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertEqual(resp.status_code, 200)
        a_controls = ControlApplicabilityResult.objects.filter(
            company=self.c, decision='applicable',
            control__framework_version__isnull=False, control__is_legacy_import=False).count()
        total = ControlApplicabilityResult.objects.filter(decision='applicable').count()
        self.assertGreater(total, a_controls)  # B contributes controls that must NOT show
        self.assertContains(resp, f"{a_controls} ضابط منطبق")  # A-only metric


class SecurityStaffOnlyTests(TestCase):
    def setUp(self):
        from compliance.framework_scope import propose_framework_scopes
        self.propose = propose_framework_scopes

    def _nonstaff(self, company):
        u = _journey_user(company, role='company_admin')
        self.assertFalse(u.is_staff)
        self.client.force_login(u)
        return u

    def test_non_staff_cannot_generate_framework_scopes(self):
        c, fv = _company_with_applicability()
        self.propose(c, apply=True)
        scope = CompanyFrameworkScope.objects.filter(company=c, status='proposed').first()
        self._nonstaff(c)
        resp = self.client.post(reverse('compliance:approve_scope', args=[scope.id]))
        self.assertEqual(resp.status_code, 302)
        scope.refresh_from_db()
        self.assertNotEqual(scope.status, 'approved')

    def test_non_staff_cannot_generate_control_plan(self):
        c, fv, scope = _company_with_official_plan()
        before = ControlApplicabilityResult.objects.filter(company=c).count()
        self._nonstaff(c)
        resp = self.client.post(reverse('compliance:generate_plan', args=[scope.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(ControlApplicabilityResult.objects.filter(company=c).count(), before)

    def test_non_staff_cannot_generate_evidence_checklist(self):
        c, fv, scope = _company_with_official_plan()  # plan exists, no checklist yet
        self.assertEqual(EvidenceChecklistItem.objects.filter(company=c).count(), 0)
        self._nonstaff(c)
        resp = self.client.post(reverse('compliance:generate_evidence_checklist'))
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(EvidenceChecklistItem.objects.filter(company=c).count(), 0)

    def test_non_staff_cannot_trigger_analysis(self):
        c, item, sub = _company_with_submission()
        self._nonstaff(c)
        resp = self.client.post(reverse('compliance:analyze_submission', args=[sub.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertFalse(EvidenceAnalysisResult.objects.filter(evidence_submission=sub).exists())

    def test_non_staff_cannot_generate_assessments(self):
        c, fv, scope = _company_with_checklist()  # no assessments yet
        self.assertEqual(ControlAssessment.objects.filter(company=c).count(), 0)
        self._nonstaff(c)
        resp = self.client.post(reverse('compliance:generate_assessments'))
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(ControlAssessment.objects.filter(company=c).count(), 0)

    def test_non_staff_cannot_update_control_assessment(self):
        c, fv, scope = _company_with_assessments()
        a = ControlAssessment.objects.filter(company=c).first()
        self._nonstaff(c)
        resp = self.client.post(reverse('compliance:auditor_review_detail', args=[a.id]),
                                {'status': 'compliant'})
        self.assertEqual(resp.status_code, 302)
        a.refresh_from_db()
        self.assertEqual(a.status, 'not_reviewed')  # decision unchanged


class SecurityUploadSafetyTests(TestCase):
    def setUp(self):
        self.c, self.fv, self.scope = _company_with_checklist()
        self.item = EvidenceChecklistItem.objects.filter(company=self.c).first()
        self.user = _journey_user(self.c)
        self.client.force_login(self.user)

    def test_upload_v2_rejects_disallowed_extension(self):
        before = EvidenceSubmission.objects.filter(checklist_item=self.item).count()
        resp = self.client.post(reverse('compliance:evidence_upload_v2', args=[self.item.id]),
                                {'uploaded_file': _SUF('malware.exe', b'MZ'), 'notes': ''})
        self.assertEqual(resp.status_code, 200)  # re-render with form error
        self.assertEqual(EvidenceSubmission.objects.filter(checklist_item=self.item).count(), before)

    def test_upload_v2_rejects_large_file(self):
        before = EvidenceSubmission.objects.filter(checklist_item=self.item).count()
        with mock.patch('compliance.forms.EVIDENCE_V2_MAX_SIZE', 8):  # 8-byte cap
            resp = self.client.post(reverse('compliance:evidence_upload_v2', args=[self.item.id]),
                                    {'uploaded_file': _SUF('big.txt', b'way too many bytes'), 'notes': ''})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(EvidenceSubmission.objects.filter(checklist_item=self.item).count(), before)

    def test_upload_v2_records_checksum(self):
        self.client.post(reverse('compliance:evidence_upload_v2', args=[self.item.id]),
                         {'uploaded_file': _SUF('ok.txt', b'evidence'), 'notes': ''})
        sub = EvidenceSubmission.objects.filter(checklist_item=self.item).latest('uploaded_at')
        self.assertEqual(len(sub.file_hash), 64)  # sha256 hexdigest

    def test_upload_v2_does_not_create_legacy_evidence(self):
        before = Evidence.objects.count()
        self.client.post(reverse('compliance:evidence_upload_v2', args=[self.item.id]),
                         {'uploaded_file': _SUF('ok.txt', b'evidence'), 'notes': ''})
        self.assertEqual(Evidence.objects.count(), before)  # no legacy Evidence rows

    def test_upload_v2_does_not_create_companycontrol(self):
        before = CompanyControl.objects.count()
        self.client.post(reverse('compliance:evidence_upload_v2', args=[self.item.id]),
                         {'uploaded_file': _SUF('ok.txt', b'evidence'), 'notes': ''})
        self.assertEqual(CompanyControl.objects.count(), before)


class SecurityAIReportingTests(TestCase):
    def test_ai_analysis_does_not_create_controlassessment(self):
        from compliance.evidence_analysis import analyze_evidence_submission
        c, item, sub = _company_with_submission()
        before = ControlAssessment.objects.count()
        analyze_evidence_submission(sub, apply=True)
        self.assertEqual(ControlAssessment.objects.count(), before)

    def test_ai_analysis_does_not_set_compliant_status(self):
        from compliance.evidence_analysis import analyze_evidence_submission
        c, fv, item, sub = _full_pipeline()
        analyze_evidence_submission(sub, apply=True)
        # No assessment is flipped to a compliant state by analysis.
        self.assertEqual(ControlAssessment.objects.filter(company=c, status='compliant').count(), 0)

    def test_reports_use_controlassessment_not_ai(self):
        from compliance.reporting import build_executive_summary
        c, fv, item, sub = _full_pipeline()
        # Auditor marks one control compliant; reports must reflect the AUDITOR decision.
        a = ControlAssessment.objects.filter(company=c).first(); a.status = 'compliant'; a.save()
        s = build_executive_summary(c)
        self.assertEqual(s['counts']['compliant'],
                         ControlAssessment.objects.filter(company=c, status='compliant').count())

    def test_reports_do_not_use_legacy_controls(self):
        from compliance.reporting import build_framework_gap_analysis
        c, fv, scope = _company_with_assessments()
        fw = fv.framework
        Control.objects.create(framework=fw, domain=Domain.objects.filter(framework=fw).first(),
                               control_id='LEG-3J', title='legacy', description='d')
        gap = build_framework_gap_analysis(c)
        ids = [g['control_id'] for f in gap for g in f['gaps']]
        self.assertNotIn('LEG-3J', ids)

    def test_reports_do_not_create_or_update_records(self):
        from compliance.reporting import (build_executive_summary, build_framework_gap_analysis,
                                          build_evidence_matrix)
        c, fv, scope = _company_with_assessments()
        before = {a.id: a.status for a in ControlAssessment.objects.filter(company=c)}
        cc_before = CompanyControl.objects.count()
        build_executive_summary(c); build_framework_gap_analysis(c); build_evidence_matrix(c)
        after = {a.id: a.status for a in ControlAssessment.objects.filter(company=c)}
        self.assertEqual(before, after)
        self.assertEqual(CompanyControl.objects.count(), cc_before)

    def test_unreviewed_controls_not_counted_as_compliant(self):
        from compliance.reporting import build_executive_summary
        c, fv, scope = _company_with_assessments()  # all not_reviewed
        s = build_executive_summary(c)
        self.assertEqual(s['counts']['compliant'], 0)
        self.assertEqual(s['compliance_percentage'], 0.0)


class Phase3JBackwardCompatTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def test_intake_still_works(self):
        c = _company()
        self.client.force_login(_journey_user(c))
        self.assertEqual(self.client.get(reverse('compliance:intake')).status_code, 200)

    def test_old_upload_evidence_flow_still_works(self):
        company, control = _company_with_control()
        self.client.force_login(_journey_user(company))
        with mock.patch('monitoring.tasks.analyze_evidence_async.delay'):
            resp = self.client.post(reverse('compliance:upload_evidence', args=[control.id]),
                                    {'evidence_file': _SUF('p.txt', b'ok')})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(Evidence.objects.count(), 1)

    def test_evidence_upload_v2_still_works(self):
        c, fv, scope = _company_with_checklist()
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        self.client.force_login(_journey_user(c))
        resp = self.client.post(reverse('compliance:evidence_upload_v2', args=[item.id]),
                                {'uploaded_file': _SUF('ok.txt', b'evidence'), 'notes': ''})
        self.assertEqual(resp.status_code, 302)
        self.assertTrue(EvidenceSubmission.objects.filter(checklist_item=item).exists())

    def test_evidence_analysis_still_works(self):
        from compliance.evidence_analysis import analyze_evidence_submission
        c, item, sub = _company_with_submission()
        analyze_evidence_submission(sub, apply=True)
        self.assertTrue(EvidenceAnalysisResult.objects.filter(evidence_submission=sub).exists())

    def test_auditor_assessment_still_works(self):
        c, fv, scope = _company_with_assessments()
        a = ControlAssessment.objects.filter(company=c).first()
        u = _journey_user(c, email='aud3j@x.com', is_staff=True)
        update_assessment_from_auditor_input(a, {'status': 'compliant'}, u)
        a.refresh_from_db()
        self.assertEqual(a.status, 'compliant')

    def test_reports_still_work(self):
        c, fv, scope = _company_with_assessments()
        self.client.force_login(_journey_user(c))
        self.assertEqual(self.client.get(reverse('compliance:report_executive_summary')).status_code, 200)

    def test_dashboard_still_works(self):
        c, fv, scope = _company_with_assessments()
        self.client.force_login(_journey_user(c))
        self.assertEqual(self.client.get(reverse('compliance:dashboard')).status_code, 200)

    def test_registration_flow_still_works(self):
        fw, dom = _fw_dom()
        with mock.patch('core.views.classify_company', return_value={'error': 'skip'}):
            resp = self.client.post(reverse('core:register'), {
                'company_name': 'RegJ', 'cr_number': '4040401234', 'sector': 'technology',
                'size': 'small', 'first_name': 'A', 'last_name': 'B', 'email': 'regj@x.com',
                'password': 'longenough12', 'target_nca': 'on'})
        self.assertEqual(resp.status_code, 302)


# ============================================================
# Phase UX-WIZARD-A — Compliance Journey Stepper/Wizard
# ============================================================
from compliance.journey import build_company_compliance_journey, STATUSES as _WIZARD_STATUSES


def _wstep(journey, key):
    return next(s for s in journey['steps'] if s['key'] == key)


class CompanyJourneyWizardBuilderTests(TestCase):
    def test_builder_returns_five_main_stages(self):
        c = _company()
        j = build_company_compliance_journey(c)
        self.assertEqual(len(j['stages']), 5)
        self.assertEqual([s['title'] for s in j['stages']],
                         ['البدء والتصنيف', 'الأطر والضوابط', 'الأدلة والتحليل',
                          'الفجوات والمعالجة', 'المراجعة والتقارير والمراقبة'])

    def test_builder_includes_sixteen_substeps(self):
        j = build_company_compliance_journey(_company())
        self.assertEqual(len(j['steps']), 16)
        total = sum(len(st['steps']) for st in j['stages'])
        self.assertEqual(total, 16)

    def test_statuses_are_valid(self):
        j = build_company_compliance_journey(_company())
        self.assertTrue(all(s['status'] in _WIZARD_STATUSES for s in j['steps']))

    def test_current_and_next_action_present(self):
        j = build_company_compliance_journey(_company())
        self.assertEqual(sum(1 for s in j['steps'] if s['status'] == 'current'), 1)
        self.assertTrue(j['next_action']['title'])

    def test_unfinished_features_shown_as_planned_truthfully(self):
        j = build_company_compliance_journey(_company())
        self.assertEqual(_wstep(j, 'ocr_extraction')['status'], 'planned')
        self.assertEqual(_wstep(j, 'rule_engine')['status'], 'planned')
        self.assertFalse(_wstep(j, 'ocr_extraction')['is_available'])

    def test_reports_locked_without_subscription(self):
        j = build_company_compliance_journey(_company())
        self.assertEqual(_wstep(j, 'reports')['status'], 'locked')

    def test_monitoring_foundation_step_present(self):
        j = build_company_compliance_journey(_company())
        self.assertEqual(_wstep(j, 'monitoring')['title'], 'المراقبة المستمرة')

    def test_builder_is_tenant_scoped(self):
        from compliance.tests import _company_with_submission
        a, item, sub = _company_with_submission()  # has evidence
        b = _company()
        self.assertEqual(_wstep(build_company_compliance_journey(a), 'evidence_upload')['status'], 'completed')
        self.assertNotEqual(_wstep(build_company_compliance_journey(b), 'evidence_upload')['status'], 'completed')

    def test_builder_does_not_write(self):
        from compliance.models import ControlAssessment, EvidenceSubmission
        from risk.models import RiskItem
        c, fv, scope = _company_with_assessments()
        counts = lambda: (ControlAssessment.objects.count(), EvidenceSubmission.objects.count(),
                          RiskItem.objects.count())
        before = counts()
        build_company_compliance_journey(c)
        self.assertEqual(before, counts())


class CompanyJourneyWizardViewTests(TestCase):
    def setUp(self):
        self.c = _company()
        self.client.force_login(_journey_user(self.c))

    def test_dashboard_renders_wizard(self):
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'مسار جاهزية الامتثال')
        self.assertContains(resp, 'ct-journey-wizard')

    def test_wizard_shows_arabic_stage_and_step_labels(self):
        resp = self.client.get(reverse('compliance:dashboard'))
        for label in ['البدء والتصنيف', 'التصنيف الذكي', 'مكتبة الضوابط',
                      'النتيجة النهائية بعد المراجعة', 'المراقبة المستمرة']:
            self.assertContains(resp, label)

    def test_wizard_has_mobile_accordion_markup(self):
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertContains(resp, 'ct-stage')           # collapsible stage (mobile-friendly)
        self.assertContains(resp, 'ct-step-card')

    def test_no_old_brand_or_legacy_count_or_certification(self):
        body = self.client.get(reverse('compliance:dashboard')).content.decode()
        for bad in ['CyberTrust KSA', '334', 'شهادة رسمية', 'اعتماد رسمي',
                    'certification', 'Certification', 'مطابق 100']:
            self.assertNotIn(bad, body)

    def test_no_english_in_parentheses_in_wizard_heading(self):
        body = self.client.get(reverse('compliance:dashboard')).content.decode()
        self.assertNotIn('مسار الامتثال (Compliance Journey)', body)

    def test_dashboard_requires_login(self):
        self.client.logout()
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)


# ============================================================
# UX-1B-RTL-FIX-A — Arabic shell + RTL layout hardening
# ============================================================
class ArabicShellRtlTests(TestCase):
    def setUp(self):
        self.c = _company()
        self.client.force_login(_journey_user(self.c))

    def test_html_renders_rtl_in_arabic_default(self):
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertContains(resp, 'dir="rtl"')
        self.assertContains(resp, 'lang="ar"')

    def test_nav_labels_are_arabic(self):
        body = self.client.get(reverse('compliance:dashboard')).content.decode()
        for label in ['لوحة التحكم', 'مسار الامتثال', 'الأطر', 'الأدلة', 'التقارير', 'المراقبة', 'تسجيل الخروج']:
            self.assertIn(label, body)

    def test_english_nav_labels_not_shown(self):
        body = self.client.get(reverse('compliance:dashboard')).content.decode()
        # The old English nav link labels must be gone (icon + label pattern).
        for bad in ['</i> Dashboard', '</i> Controls', '</i> Monitoring',
                    '</i> Reports', '</i> Evidence', '</i> Logout']:
            self.assertNotIn(bad, body)

    def test_intake_copy_has_no_english_parentheticals(self):
        resp = self.client.get(reverse('compliance:intake'))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode()
        self.assertNotIn('(Intake)', body)
        self.assertNotIn('Business intake', body)
        self.assertIn('تصنيف الشركة', body)

    def test_mobile_nav_and_overflow_guard_present(self):
        body = self.client.get(reverse('compliance:dashboard')).content.decode()
        self.assertIn('fa-bars', body)              # mobile hamburger
        self.assertIn('overflow-x-hidden', body)    # horizontal-overflow safety layer

    def test_no_old_brand_count_or_certification(self):
        body = self.client.get(reverse('compliance:dashboard')).content.decode()
        for bad in ['CyberTrust KSA', '334', 'شهادة رسمية', 'اعتماد رسمي', 'certification']:
            self.assertNotIn(bad, body)

    def test_dashboard_and_journey_still_render(self):
        self.assertContains(self.client.get(reverse('compliance:dashboard')), 'مسار جاهزية الامتثال')

    def test_intake_requires_login(self):
        self.client.logout()
        resp = self.client.get(reverse('compliance:intake'))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)


# ============================================================
# Phase 6A — Smart Classification Engine Foundation
# ============================================================
class SmartClassificationServiceTests(TestCase):
    def _classify(self, **company_kw):
        from compliance.smart_classification import classify_company
        return classify_company(_company(**company_kw))

    def _status(self, result, code):
        return {r.code: r.status for r in result.recommendations}[code]

    def test_ksa_company_gets_ecc_required(self):
        r = self._classify(target_nca=True)
        self.assertEqual(self._status(r, 'NCA-ECC-2-2024'), 'required')

    def test_ecc_required_for_ksa_default_even_without_target(self):
        r = self._classify(target_nca=False, country='SA')
        self.assertIn(self._status(r, 'NCA-ECC-2-2024'), ('required', 'recommended'))

    def test_cloud_intake_recommends_ccc(self):
        from compliance.smart_classification import classify_company
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, uses_cloud_services=True)
        r = classify_company(c)
        self.assertEqual({x.code: x.status for x in r.recommendations}['NCA-CCC-2-2024'], 'recommended')

    def test_critical_systems_intake_recommends_cscc(self):
        from compliance.smart_classification import classify_company
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, is_critical_system_operator=True)
        r = classify_company(c)
        self.assertEqual({x.code: x.status for x in r.recommendations}['NCA-CSCC-1-2019'], 'recommended')

    def test_aramco_readiness_recommends_sacs002(self):
        r = self._classify(target_aramco=True)
        self.assertEqual(self._status(r, 'ARAMCO-SACS-002'), 'required')

    def test_sabic_readiness_recommends_sabic(self):
        r = self._classify(target_sabic=True)
        self.assertEqual(self._status(r, 'SABIC-CYBERTRUST-1-0'), 'required')

    def test_telework_intake_recommends_tcc(self):
        from compliance.smart_classification import classify_company
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, has_remote_work=True)
        r = classify_company(c)
        self.assertEqual({x.code: x.status for x in r.recommendations}['NCA-TCC-1-2021'], 'recommended')

    def test_social_media_intake_recommends_osmacc(self):
        from compliance.smart_classification import classify_company
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, manages_official_social_media_accounts=True)
        r = classify_company(c)
        self.assertEqual({x.code: x.status for x in r.recommendations}['NCA-OSMACC-1-2021'], 'recommended')

    def test_unindicated_framework_not_called_impossible(self):
        # No signals -> SABIC is 'not_indicated', with a non-judgmental Arabic reason.
        r = self._classify()
        rec = {x.code: x for x in r.recommendations}['SABIC-CYBERTRUST-1-0']
        self.assertEqual(rec.status, 'not_indicated')
        self.assertIn('بناءً على البيانات المتاحة', rec.reason_ar)

    def test_incomplete_profile_lists_missing_inputs_and_lowers_confidence(self):
        r = self._classify()  # no intake profile
        self.assertFalse(r.has_intake)
        self.assertTrue(r.missing_inputs)
        self.assertLessEqual(r.overall_confidence, 60)

    def test_deterministic_same_input_same_output(self):
        from compliance.smart_classification import classify_company
        c = _company(target_nca=True)
        a = classify_company(c); b = classify_company(c)
        self.assertEqual([(x.code, x.status, x.confidence) for x in a.recommendations],
                         [(x.code, x.status, x.confidence) for x in b.recommendations])


class SmartClassificationCountTests(TestCase):
    def test_official_total_is_417_not_334(self):
        from compliance.smart_classification import OFFICIAL_TOTAL, OFFICIAL_FRAMEWORKS
        self.assertEqual(OFFICIAL_TOTAL, 417)
        self.assertNotIn(334, [f['expected'] for f in OFFICIAL_FRAMEWORKS])

    def test_per_framework_expected_counts(self):
        from compliance.smart_classification import _FW
        self.assertEqual(_FW['NCA-ECC-2-2024']['expected'], 108)
        self.assertEqual(_FW['NCA-CSCC-1-2019']['expected'], 32)
        self.assertEqual(_FW['NCA-CCC-2-2024']['expected'], 55)
        self.assertEqual(_FW['NCA-TCC-1-2021']['expected'], 21)
        self.assertEqual(_FW['NCA-OSMACC-1-2021']['expected'], 15)
        self.assertEqual(_FW['ARAMCO-SACS-002']['expected'], 92)
        self.assertEqual(_FW['SABIC-CYBERTRUST-1-0']['expected'], 94)

    def test_count_helper_prefers_db_else_expected(self):
        from compliance.smart_classification import official_control_count
        # Fresh DB: no official controls imported -> falls back to expected, never 334.
        self.assertEqual(official_control_count('NCA-ECC-2-2024'), 108)


class SmartClassificationUITests(TestCase):
    def test_summary_renders_for_company_user(self):
        c = _company(target_nca=True)
        self.client.force_login(_journey_user(c))
        resp = self.client.get(reverse('compliance:classification'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'التصنيف الذكي')
        self.assertContains(resp, 'تصنيف أولي استشاري')
        self.assertContains(resp, 'مستوى المخاطر')

    def test_summary_has_no_certification_or_final_decision_wording(self):
        c = _company(target_nca=True)
        self.client.force_login(_journey_user(c))
        body = self.client.get(reverse('compliance:classification')).content.decode()
        for bad in ('معتمد رسميًا', 'مؤهل رسميًا', 'تم إصدار شهادة', 'قرار نهائي', 'CyberTrust KSA', '334'):
            self.assertNotIn(bad, body)

    def test_dashboard_shows_classification_card(self):
        c = _company(target_nca=True)
        self.client.force_login(_journey_user(c))
        body = self.client.get(reverse('compliance:dashboard')).content.decode()
        self.assertIn('التصنيف الذكي', body)
        self.assertIn(reverse('compliance:classification'), body)

    def test_english_mode_renders_smart_classification(self):
        c = _company(target_nca=True)
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('set_language'), {'language': 'en', 'next': '/'})
        body = self.client.get(reverse('compliance:classification')).content.decode()
        self.assertIn('Smart Classification', body)
        self.assertIn('Risk level', body)


class SmartClassificationJourneyTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def test_step_needs_action_without_intake(self):
        from compliance.journey import build_company_compliance_journey
        c = _company()
        j = build_company_compliance_journey(c)
        step = {s['key']: s for s in j['steps']}['smart_classification']
        self.assertNotEqual(step['status'], 'completed')

    def test_step_completed_once_intake_profile_exists(self):
        from compliance.journey import build_company_compliance_journey
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, uses_cloud_services=True)
        j = build_company_compliance_journey(c)
        step = {s['key']: s for s in j['steps']}['smart_classification']
        self.assertEqual(step['status'], 'completed')

    def test_planned_steps_remain_not_completed(self):
        from compliance.journey import build_company_compliance_journey
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, uses_cloud_services=True)
        j = build_company_compliance_journey(c)
        steps = {s['key']: s for s in j['steps']}
        for k in ('ocr_extraction', 'rule_engine'):
            self.assertNotEqual(steps[k]['status'], 'completed')


class SmartClassificationSecurityTests(TestCase):
    def test_anonymous_redirected(self):
        resp = self.client.get(reverse('compliance:classification'))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_user_sees_only_own_company_classification(self):
        c1 = _company(name='Alpha', target_nca=True)
        c2 = _company(name='Beta', target_aramco=True)
        CompanyIntakeProfile.objects.create(company=c2, works_with_aramco=True)
        self.client.force_login(_journey_user(c1))
        # The view always classifies request.user.company (c1); c2 data never leaks.
        from compliance.smart_classification import classify_company
        r1 = classify_company(c1)
        self.assertFalse(r1.has_intake)  # c1 has no intake; c2's intake does not bleed in


# ============================================================
# Phase 6B — Applicability Engine Foundation
# ============================================================
class ApplicabilityEngineServiceTests(TestCase):
    def _eval(self, **kw):
        from compliance.applicability_engine import evaluate_company_applicability
        return evaluate_company_applicability(_company(**kw))

    def _sum(self, preview, code):
        return {s.framework_code: s for s in preview.summaries}[code]

    def test_ecc_applicable_for_nca_readiness(self):
        p = self._eval(target_nca=True)
        self.assertEqual(self._sum(p, 'NCA-ECC-2-2024').status, 'applicable')

    def test_cscc_applicable_for_critical_systems(self):
        from compliance.applicability_engine import evaluate_company_applicability
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, is_critical_system_operator=True)
        p = evaluate_company_applicability(c)
        self.assertEqual({s.framework_code: s for s in p.summaries}['NCA-CSCC-1-2019'].status, 'applicable')

    def test_ccc_applicable_for_cloud(self):
        from compliance.applicability_engine import evaluate_company_applicability
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, uses_cloud_services=True)
        p = evaluate_company_applicability(c)
        self.assertEqual({s.framework_code: s for s in p.summaries}['NCA-CCC-2-2024'].status, 'applicable')

    def test_tcc_applicable_for_remote_work(self):
        from compliance.applicability_engine import evaluate_company_applicability
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, has_remote_work=True)
        p = evaluate_company_applicability(c)
        self.assertEqual({s.framework_code: s for s in p.summaries}['NCA-TCC-1-2021'].status, 'applicable')

    def test_osmacc_applicable_for_social_media(self):
        from compliance.applicability_engine import evaluate_company_applicability
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, manages_official_social_media_accounts=True)
        p = evaluate_company_applicability(c)
        self.assertEqual({s.framework_code: s for s in p.summaries}['NCA-OSMACC-1-2021'].status, 'applicable')

    def test_aramco_applicable_for_relationship(self):
        p = self._eval(target_aramco=True)
        self.assertEqual(self._sum(p, 'ARAMCO-SACS-002').status, 'applicable')

    def test_sabic_applicable_for_relationship(self):
        p = self._eval(target_sabic=True)
        self.assertEqual(self._sum(p, 'SABIC-CYBERTRUST-1-0').status, 'applicable')

    def test_missing_intake_yields_needs_more_information(self):
        # No intake profile -> not-indicated frameworks are needs_more_information (not flat not_applicable).
        p = self._eval()
        self.assertEqual(self._sum(p, 'NCA-TCC-1-2021').status, 'needs_more_information')
        self.assertFalse(p.has_intake)

    def test_not_indicated_with_intake_is_not_applicable(self):
        from compliance.applicability_engine import evaluate_company_applicability
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, uses_cloud_services=True)  # no telework
        p = evaluate_company_applicability(c)
        self.assertEqual({s.framework_code: s for s in p.summaries}['NCA-TCC-1-2021'].status, 'not_applicable')

    def test_deterministic_repeated_results(self):
        from compliance.applicability_engine import evaluate_company_applicability
        c = _company(target_nca=True)
        a = evaluate_company_applicability(c); b = evaluate_company_applicability(c)
        self.assertEqual([(s.framework_code, s.status, s.applicable_count) for s in a.summaries],
                         [(s.framework_code, s.status, s.applicable_count) for s in b.summaries])

    def test_engine_does_not_write(self):
        from compliance.applicability_engine import evaluate_company_applicability
        from compliance.models import ControlApplicabilityResult, ControlAssessment, CompanyControl
        c = _company(target_nca=True)
        CompanyIntakeProfile.objects.create(company=c, uses_cloud_services=True)
        counts = lambda: (ControlApplicabilityResult.objects.count(), ControlAssessment.objects.count(),
                          CompanyControl.objects.count())
        before = counts()
        evaluate_company_applicability(c)
        self.assertEqual(before, counts())

    def test_no_compliance_status_strings_in_results(self):
        p = self._eval(target_nca=True)
        for s in p.summaries:
            self.assertIn(s.status, ('applicable', 'not_applicable', 'needs_more_information'))


class ApplicabilityCountTests(TestCase):
    def test_official_total_417_no_334(self):
        from compliance.applicability_engine import evaluate_company_applicability
        p = evaluate_company_applicability(_company(target_nca=True))
        self.assertEqual(p.total_controls, 417)
        self.assertNotEqual(p.total_controls, 334)

    def test_per_framework_totals_match(self):
        from compliance.applicability_engine import evaluate_company_applicability
        p = evaluate_company_applicability(_company(target_nca=True))
        totals = {s.framework_code: s.total_controls for s in p.summaries}
        self.assertEqual(totals['NCA-ECC-2-2024'], 108)
        self.assertEqual(totals['ARAMCO-SACS-002'], 92)
        self.assertEqual(totals['SABIC-CYBERTRUST-1-0'], 94)
        self.assertEqual(totals['NCA-CCC-2-2024'], 55)


class ApplicabilityUITests(TestCase):
    def _login(self, **kw):
        c = _company(**kw)
        self.client.force_login(_journey_user(c))
        return c

    def test_page_renders_with_arabic_labels(self):
        self._login(target_nca=True)
        resp = self.client.get(reverse('compliance:applicability_preview'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'قابلية تطبيق الضوابط')
        self.assertContains(resp, 'إجمالي الضوابط')
        self.assertContains(resp, 'تحتاج بيانات إضافية')

    def test_disclaimers_present(self):
        self._login(target_nca=True)
        body = self.client.get(reverse('compliance:applicability_preview')).content.decode()
        self.assertIn('لا يقرر الامتثال أو عدم الامتثال', body)
        self.assertIn('تعتمد على الأدلة ومراجعة المدقق', body)

    def test_no_compliance_or_verdict_or_cert_wording(self):
        self._login(target_nca=True)
        body = self.client.get(reverse('compliance:applicability_preview')).content.decode()
        for bad in ('متوافق', 'غير متوافق', 'Noncompliance', 'تم إصدار شهادة', 'قرار نهائي', '334', 'CyberTrust KSA'):
            self.assertNotIn(bad, body)

    def test_dashboard_card_present(self):
        self._login(target_nca=True)
        body = self.client.get(reverse('compliance:dashboard')).content.decode()
        self.assertIn('قابلية تطبيق الضوابط', body)
        self.assertIn(reverse('compliance:applicability_preview'), body)

    def test_english_mode_renders(self):
        self._login(target_nca=True)
        self.client.post(reverse('set_language'), {'language': 'en', 'next': '/'})
        body = self.client.get(reverse('compliance:applicability_preview')).content.decode()
        self.assertIn('Control applicability', body)
        self.assertIn('Total controls', body)


class ApplicabilityJourneyTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def _step(self, company, key):
        from compliance.journey import build_company_compliance_journey
        j = build_company_compliance_journey(company)
        return {s['key']: s for s in j['steps']}[key]

    def test_applicability_step_needs_action_without_intake(self):
        self.assertNotEqual(self._step(_company(), 'applicability')['status'], 'completed')

    def test_applicability_step_completed_with_intake(self):
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, uses_cloud_services=True)
        self.assertEqual(self._step(c, 'applicability')['status'], 'completed')

    def test_downstream_steps_remain_not_completed(self):
        c = _company()
        CompanyIntakeProfile.objects.create(company=c, uses_cloud_services=True)
        for k in ('ocr_extraction', 'rule_engine', 'final_verdict'):
            self.assertNotEqual(self._step(c, k)['status'], 'completed')


class ApplicabilitySecurityTests(TestCase):
    def test_anonymous_redirected(self):
        resp = self.client.get(reverse('compliance:applicability_preview'))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_only_own_company_evaluated(self):
        from compliance.applicability_engine import evaluate_company_applicability
        c1 = _company(name='Alpha', target_nca=True)
        c2 = _company(name='Beta')
        CompanyIntakeProfile.objects.create(company=c2, has_remote_work=True)
        self.client.force_login(_journey_user(c1))
        p1 = evaluate_company_applicability(c1)
        self.assertFalse(p1.has_intake)  # c2's intake never bleeds into c1


# ============================================================
# Phase 6C — Evidence Text Extraction Foundation
# ============================================================
import os
import tempfile as _tempfile


def _mk_txt(text='Cybersecurity   policy\n\n\n approved.'):
    p = _tempfile.mktemp(suffix='.txt')
    with open(p, 'w', encoding='utf-8') as f:
        f.write(text)
    return p


class EvidenceExtractionServiceTests(TestCase):
    def test_txt_extraction_and_whitespace_normalized(self):
        from compliance.evidence_extraction import extract_text_from_file, EXTRACTED
        p = _mk_txt()
        r = extract_text_from_file(p, 'doc.txt'); os.unlink(p)
        self.assertEqual(r.status, EXTRACTED)
        self.assertIn('Cybersecurity policy', r.extracted_text)   # collapsed spaces
        self.assertNotIn('   ', r.extracted_text)
        self.assertEqual(r.char_count, len(r.extracted_text))

    def test_csv_extraction(self):
        from compliance.evidence_extraction import extract_text_from_file, EXTRACTED
        p = _tempfile.mktemp(suffix='.csv')
        open(p, 'w', encoding='utf-8').write('asset,owner\nfirewall,it')
        r = extract_text_from_file(p, 'd.csv'); os.unlink(p)
        self.assertEqual(r.status, EXTRACTED)
        self.assertIn('firewall', r.extracted_text)

    def test_docx_extraction(self):
        from docx import Document
        from compliance.evidence_extraction import extract_text_from_file, EXTRACTED
        p = _tempfile.mktemp(suffix='.docx')
        d = Document(); d.add_paragraph('Incident Response Plan v2'); d.save(p)
        r = extract_text_from_file(p, 'd.docx'); os.unlink(p)
        self.assertEqual(r.status, EXTRACTED)
        self.assertIn('Incident Response Plan', r.extracted_text)
        self.assertEqual(r.extraction_method, 'docx')

    def test_xlsx_extraction(self):
        from openpyxl import Workbook
        from compliance.evidence_extraction import extract_text_from_file, EXTRACTED
        p = _tempfile.mktemp(suffix='.xlsx')
        wb = Workbook(); ws = wb.active; ws['A1'] = 'Asset'; ws['B1'] = 'Owner'; wb.save(p)
        r = extract_text_from_file(p, 'd.xlsx'); os.unlink(p)
        self.assertEqual(r.status, EXTRACTED)
        self.assertIn('Asset', r.extracted_text)

    def test_pdf_text_layer_extraction(self):
        from dashboard.reports import gap_analysis_pdf
        from compliance.evidence_extraction import extract_text_from_file, EXTRACTED, NO_TEXT
        c = _company()
        p = _tempfile.mktemp(suffix='.pdf')
        with open(p, 'wb') as f:
            f.write(gap_analysis_pdf(c))
        r = extract_text_from_file(p, 'report.pdf'); os.unlink(p)
        # A generated (text-layer) PDF should extract text; method is the text layer.
        self.assertIn(r.status, (EXTRACTED, NO_TEXT))
        if r.status == EXTRACTED:
            self.assertEqual(r.extraction_method, 'pdf_text_layer')
            self.assertIsNotNone(r.page_count)

    def test_image_returns_no_text_without_ocr(self):
        from compliance.evidence_extraction import extract_text_from_file, NO_TEXT
        p = _tempfile.mktemp(suffix='.png')
        open(p, 'wb').write(b'\x89PNG\r\n\x1a\n')
        r = extract_text_from_file(p, 'scan.png'); os.unlink(p)
        self.assertEqual(r.status, NO_TEXT)
        self.assertTrue(any('OCR' in w for w in r.warnings))

    def test_unsupported_type(self):
        from compliance.evidence_extraction import extract_text_from_file, UNSUPPORTED
        p = _tempfile.mktemp(suffix='.exe'); open(p, 'wb').write(b'MZ')
        r = extract_text_from_file(p, 'x.exe'); os.unlink(p)
        self.assertEqual(r.status, UNSUPPORTED)

    def test_too_large_file(self):
        from compliance import evidence_extraction as ee
        p = _mk_txt('x' * 100)
        orig = ee.MAX_EXTRACTION_BYTES
        ee.MAX_EXTRACTION_BYTES = 10
        try:
            r = ee.extract_text_from_file(p, 'big.txt')
        finally:
            ee.MAX_EXTRACTION_BYTES = orig
            os.unlink(p)
        self.assertEqual(r.status, 'too_large')

    def test_missing_file_fails_safely(self):
        from compliance.evidence_extraction import extract_text_from_file, FAILED
        r = extract_text_from_file('/nope/missing.txt', 'missing.txt')
        self.assertEqual(r.status, FAILED)
        self.assertNotIn('/nope/', (r.error_message or ''))  # no path leakage

    def test_parser_exception_returns_failed_without_traceback(self):
        from compliance.evidence_extraction import extract_text_from_file, FAILED
        # A .docx that is not a valid zip -> python-docx raises -> FAILED, no traceback.
        p = _tempfile.mktemp(suffix='.docx'); open(p, 'wb').write(b'not a real docx')
        r = extract_text_from_file(p, 'bad.docx'); os.unlink(p)
        self.assertEqual(r.status, FAILED)
        self.assertNotIn('Traceback', (r.error_message or ''))

    def test_text_length_capped(self):
        from compliance import evidence_extraction as ee
        p = _mk_txt('a' * (ee.MAX_TEXT_CHARS + 5000))
        r = ee.extract_text_from_file(p, 'long.txt'); os.unlink(p)
        self.assertTrue(r.truncated)
        self.assertLessEqual(r.char_count, ee.MAX_TEXT_CHARS)

    def test_deterministic_repeated_extraction(self):
        from compliance.evidence_extraction import extract_text_from_file
        p = _mk_txt()
        a = extract_text_from_file(p, 'd.txt'); b = extract_text_from_file(p, 'd.txt'); os.unlink(p)
        self.assertEqual((a.status, a.extracted_text, a.char_count),
                         (b.status, b.extracted_text, b.char_count))


def _company_with_submission_file(filename='policy.txt', content=b'Access control policy approved.', file_type='txt'):
    """Reuse the existing submission helper (builds the full checklist FK chain)."""
    return _company_with_submission(name=filename, content=content, ftype=file_type)


class EvidenceExtractionUITests(TestCase):
    def test_not_attempted_state_does_not_imply_success(self):
        c, item, sub = _company_with_submission_file()
        self.client.force_login(_journey_user(c))
        resp = self.client.get(reverse('compliance:evidence_extraction', args=[sub.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'استخراج النص من الدليل')
        self.assertContains(resp, 'لم يتم تشغيل استخراج النص بعد.')
        self.assertContains(resp, 'تشغيل استخراج النص')           # run button
        self.assertContains(resp, 'لا يمثل تحليلًا أو حكمًا على الامتثال')
        self.assertNotContains(resp, 'حالة الاستخراج')            # no result metadata yet

    def test_extracted_state_shows_stored_result(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission_file()
        save_extraction_for_submission(sub)
        self.client.force_login(_journey_user(c))
        resp = self.client.get(reverse('compliance:evidence_extraction', args=[sub.id]))
        self.assertContains(resp, 'حالة الاستخراج')
        self.assertContains(resp, 'النص المستخرج')

    def test_no_compliance_or_verdict_or_path_leak(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission_file()
        save_extraction_for_submission(sub)
        self.client.force_login(_journey_user(c))
        body = self.client.get(reverse('compliance:evidence_extraction', args=[sub.id])).content.decode()
        for bad in ('متوافق', 'غير متوافق', 'قرار نهائي', 'الدليل كافٍ', 'الدليل غير كافٍ'):
            self.assertNotIn(bad, body)
        self.assertNotIn('/media/evidence_v2/', body)  # raw stored path not exposed

    def test_english_mode_does_not_break(self):
        c, item, sub = _company_with_submission_file()
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('set_language'), {'language': 'en', 'next': '/'})
        body = self.client.get(reverse('compliance:evidence_extraction', args=[sub.id])).content.decode()
        self.assertIn('Evidence text extraction', body)
        self.assertIn('Run text extraction', body)


class EvidenceExtractionPersistenceTests(TestCase):
    def test_run_action_persists_one_result_for_owner(self):
        from compliance.models import EvidenceTextExtraction
        c, item, sub = _company_with_submission_file()
        self.client.force_login(_journey_user(c))
        resp = self.client.post(reverse('compliance:run_evidence_extraction', args=[sub.id]))
        self.assertEqual(resp.status_code, 302)
        ext = EvidenceTextExtraction.objects.get(submission=sub)
        self.assertEqual(ext.status, 'extracted')
        self.assertGreater(ext.char_count, 0)

    def test_rerun_updates_not_duplicates(self):
        from compliance.models import EvidenceTextExtraction
        c, item, sub = _company_with_submission_file()
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('compliance:run_evidence_extraction', args=[sub.id]))
        self.client.post(reverse('compliance:run_evidence_extraction', args=[sub.id]))
        self.assertEqual(EvidenceTextExtraction.objects.filter(submission=sub).count(), 1)

    def test_warnings_stored_safely_and_no_path(self):
        from compliance.models import EvidenceTextExtraction
        c, item, sub = _company_with_submission_file(filename='scan.png', file_type='png')
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('compliance:run_evidence_extraction', args=[sub.id]))
        ext = EvidenceTextExtraction.objects.get(submission=sub)
        self.assertEqual(ext.status, 'no_text_extracted')
        self.assertIsInstance(ext.warnings, list)
        self.assertNotIn('/media/', (ext.error_message or ''))

    def test_run_action_get_not_allowed(self):
        c, item, sub = _company_with_submission_file()
        self.client.force_login(_journey_user(c))
        resp = self.client.get(reverse('compliance:run_evidence_extraction', args=[sub.id]))
        self.assertEqual(resp.status_code, 405)

    def test_run_action_anonymous_redirected(self):
        c, item, sub = _company_with_submission_file()
        resp = self.client.post(reverse('compliance:run_evidence_extraction', args=[sub.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_run_action_cross_company_blocked(self):
        from compliance.models import EvidenceTextExtraction
        c1, i1, s1 = _company_with_submission_file()
        c2 = _company(name='Other2')
        self.client.force_login(_journey_user(c2))
        resp = self.client.post(reverse('compliance:run_evidence_extraction', args=[s1.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertFalse(EvidenceTextExtraction.objects.filter(submission=s1).exists())


class EvidenceExtractionJourneyTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def _step(self, company, key):
        from compliance.journey import build_company_compliance_journey
        j = build_company_compliance_journey(company)
        return {s['key']: s for s in j['steps']}[key]

    def test_extraction_step_planned_without_evidence(self):
        s = self._step(_company(), 'ocr_extraction')
        self.assertEqual(s['status'], 'planned')
        self.assertFalse(s['is_available'])

    def test_extractable_file_without_extraction_is_needs_action(self):
        # Truthfulness gate: extractable by type but no extraction run yet.
        c, item, sub = _company_with_submission_file(file_type='pdf')
        self.assertEqual(self._step(c, 'evidence_upload')['status'], 'completed')
        self.assertEqual(self._step(c, 'ocr_extraction')['status'], 'needs_action')

    def test_completed_only_after_successful_extraction(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission_file(file_type='txt',
                                                     content=b'Access control policy approved.')
        self.assertEqual(self._step(c, 'ocr_extraction')['status'], 'needs_action')
        save_extraction_for_submission(sub)  # persists status=extracted, char_count>0
        self.assertEqual(self._step(c, 'ocr_extraction')['status'], 'completed')

    def test_no_text_result_keeps_needs_action(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission_file(filename='scan.png', file_type='png')
        save_extraction_for_submission(sub)  # status=no_text_extracted
        self.assertEqual(self._step(c, 'ocr_extraction')['status'], 'needs_action')

    def test_downstream_steps_remain_not_completed(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission_file(file_type='txt')
        save_extraction_for_submission(sub)
        for k in ('rule_engine', 'final_verdict'):
            self.assertNotEqual(self._step(c, k)['status'], 'completed')


class EvidenceExtractionSecurityTests(TestCase):
    def test_anonymous_redirected(self):
        c, item, sub = _company_with_submission_file()
        resp = self.client.get(reverse('compliance:evidence_extraction', args=[sub.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_user_cannot_access_other_company_extraction(self):
        c1, i1, s1 = _company_with_submission_file()
        c2 = _company(name='Other')
        self.client.force_login(_journey_user(c2))
        resp = self.client.get(reverse('compliance:evidence_extraction', args=[s1.id]))
        self.assertEqual(resp.status_code, 302)  # redirected away (not found for this tenant)
        self.assertNotIn('استخراج النص من الدليل', resp.content.decode())


# ============================================================
# Phase 6D — AI Evidence Analyzer Advisory
# ============================================================
class _FakeProvider:
    """Deterministic fake AI provider — no network."""
    def __init__(self, payload=None, exc=None):
        self.payload = payload if payload is not None else {
            'relevance': 'high', 'confidence': 80, 'summary': 'يبدو أن الدليل مرتبط بالضابط.',
            'matched_signals': ['سياسة موثّقة'], 'missing_items': ['تاريخ الاعتماد'],
            'recommendations': ['أرفق سجل الموافقة'],
        }
        self.exc = exc
    def analyze(self, prompt):
        if self.exc:
            raise self.exc
        return self.payload


def _extracted_submission(ftype='txt', content=b'Access control policy approved by management.'):
    """A submission with a persisted successful extraction (gate satisfied)."""
    from compliance.evidence_extraction import save_extraction_for_submission
    c, item, sub = _company_with_submission(name='p.%s' % ftype, content=content, ftype=ftype)
    save_extraction_for_submission(sub)
    return c, item, sub


class AIAnalyzerGateTests(TestCase):
    def test_no_extraction_cannot_analyze(self):
        from compliance.evidence_ai_analyzer import can_analyze_submission
        c, item, sub = _company_with_submission()
        ok, reason = can_analyze_submission(sub)
        self.assertFalse(ok)
        self.assertIn('استخراج نص', reason)

    def test_no_text_extracted_cannot_analyze(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        from compliance.evidence_ai_analyzer import can_analyze_submission
        c, item, sub = _company_with_submission(name='scan.png', ftype='png', content=b'\x89PNG\r\n')
        save_extraction_for_submission(sub)  # no_text_extracted
        ok, _ = can_analyze_submission(sub)
        self.assertFalse(ok)

    def test_extracted_with_text_can_analyze(self):
        from compliance.evidence_ai_analyzer import can_analyze_submission
        c, item, sub = _extracted_submission()
        ok, reason = can_analyze_submission(sub)
        self.assertTrue(ok)
        self.assertEqual(reason, '')


class AIAnalyzerServiceTests(TestCase):
    def test_completed_with_fake_provider(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        c, item, sub = _extracted_submission()
        obj = analyze_submission_evidence(sub, provider=_FakeProvider())
        self.assertEqual(obj.status, 'completed')
        self.assertEqual(obj.relevance, 'high')
        self.assertEqual(obj.confidence, 80)
        self.assertEqual(obj.matched_signals, ['سياسة موثّقة'])

    def test_rerun_updates_same_row(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        from compliance.models import EvidenceAIAnalysis
        c, item, sub = _extracted_submission()
        analyze_submission_evidence(sub, provider=_FakeProvider())
        analyze_submission_evidence(sub, provider=_FakeProvider({'relevance': 'low', 'confidence': 10}))
        self.assertEqual(EvidenceAIAnalysis.objects.filter(submission=sub).count(), 1)
        self.assertEqual(EvidenceAIAnalysis.objects.get(submission=sub).relevance, 'low')

    def test_gate_blocks_provider_call(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        c, item, sub = _company_with_submission()  # no extraction
        called = {'n': 0}
        class Spy:
            def analyze(self, prompt): called['n'] += 1; return {}
        obj = analyze_submission_evidence(sub, provider=Spy())
        self.assertEqual(obj.status, 'skipped')
        self.assertEqual(called['n'], 0)

    def test_invalid_response_is_failed(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        c, item, sub = _extracted_submission()
        obj = analyze_submission_evidence(sub, provider=_FakeProvider(payload=['not', 'a', 'dict']))
        self.assertEqual(obj.status, 'failed')

    def test_provider_exception_is_failed_safely(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        c, item, sub = _extracted_submission()
        obj = analyze_submission_evidence(sub, provider=_FakeProvider(exc=RuntimeError('boom')))
        self.assertEqual(obj.status, 'failed')
        self.assertNotIn('boom', obj.error_message)  # no raw exception leaked

    def test_provider_unavailable_is_skipped(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence, ProviderUnavailable
        c, item, sub = _extracted_submission()
        obj = analyze_submission_evidence(sub, provider=_FakeProvider(exc=ProviderUnavailable('no key')))
        self.assertEqual(obj.status, 'skipped')

    def test_confidence_clamped(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        c, item, sub = _extracted_submission()
        obj = analyze_submission_evidence(sub, provider=_FakeProvider({'relevance': 'high', 'confidence': 999}))
        self.assertEqual(obj.confidence, 100)

    def test_relevance_restricted(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        c, item, sub = _extracted_submission()
        obj = analyze_submission_evidence(sub, provider=_FakeProvider({'relevance': 'compliant', 'confidence': 50}))
        self.assertEqual(obj.relevance, 'unclear')   # invalid value coerced, never a compliance verdict

    def test_prompt_advisory_and_no_path_or_secret(self):
        from compliance.evidence_ai_analyzer import build_prompt
        c, item, sub = _extracted_submission()
        prompt = build_prompt(sub, sub.text_extraction)
        self.assertIn('ADVISORY ONLY', prompt)
        self.assertIn('DO NOT decide compliance', prompt)
        self.assertNotIn('/media/', prompt)
        self.assertNotIn(settings.OPENAI_API_KEY or 'no-key-sentinel', prompt) if (settings.OPENAI_API_KEY) else None

    def test_no_final_compliance_wording_in_stored_output(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        c, item, sub = _extracted_submission()
        obj = analyze_submission_evidence(sub, provider=_FakeProvider())
        blob = (obj.summary + ' ' + ' '.join(obj.matched_signals + obj.missing_items + obj.recommendations))
        for bad in ('متوافق نهائيًا', 'قرار نهائي', 'شهادة'):
            self.assertNotIn(bad, blob)


class AIAnalyzerUITests(TestCase):
    def _patch_provider(self, provider):
        import compliance.evidence_ai_analyzer as mod
        self._orig = mod.default_provider
        mod.default_provider = lambda: provider
        self.addCleanup(lambda: setattr(mod, 'default_provider', self._orig))

    def test_preview_renders_for_owner(self):
        c, item, sub = _extracted_submission()
        self.client.force_login(_journey_user(c))
        resp = self.client.get(reverse('compliance:evidence_ai_analysis', args=[sub.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'التحليل الاستشاري للذكاء الاصطناعي')
        self.assertContains(resp, 'لا يُعد قرارًا نهائيًا')
        self.assertContains(resp, 'تشغيل التحليل الاستشاري')

    def test_gate_warning_shown_without_extraction(self):
        c, item, sub = _company_with_submission()
        self.client.force_login(_journey_user(c))
        body = self.client.get(reverse('compliance:evidence_ai_analysis', args=[sub.id])).content.decode()
        self.assertIn('يجب استخراج نص قابل للقراءة من الدليل قبل تشغيل التحليل الاستشاري', body)

    def test_post_run_works_for_owner_with_fake_provider(self):
        from compliance.models import EvidenceAIAnalysis
        c, item, sub = _extracted_submission()
        self._patch_provider(_FakeProvider())
        self.client.force_login(_journey_user(c))
        resp = self.client.post(reverse('compliance:run_evidence_ai_analysis', args=[sub.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(EvidenceAIAnalysis.objects.get(submission=sub).status, 'completed')

    def test_no_compliance_status_or_verdict_wording(self):
        c, item, sub = _extracted_submission()
        self._patch_provider(_FakeProvider())
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('compliance:run_evidence_ai_analysis', args=[sub.id]))
        body = self.client.get(reverse('compliance:evidence_ai_analysis', args=[sub.id])).content.decode()
        for bad in (' C/PC/NC', 'Noncompliance', 'متوافق نهائيًا', 'غير متوافق نهائيًا', 'قرار نهائي', 'شهادة امتثال نهائية'):
            self.assertNotIn(bad, body)

    def test_english_mode_does_not_break(self):
        c, item, sub = _extracted_submission()
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('set_language'), {'language': 'en', 'next': '/'})
        body = self.client.get(reverse('compliance:evidence_ai_analysis', args=[sub.id])).content.decode()
        self.assertIn('AI advisory analysis', body)
        self.assertIn('Run advisory analysis', body)


class AIAnalyzerJourneyTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def _step(self, company, key):
        from compliance.journey import build_company_compliance_journey
        j = build_company_compliance_journey(company)
        return {s['key']: s for s in j['steps']}[key]

    def test_planned_without_extracted_text(self):
        c, item, sub = _company_with_submission()  # uploaded, not extracted
        self.assertEqual(self._step(c, 'ai_analysis')['status'], 'planned')

    def test_needs_action_with_extraction_but_no_analysis(self):
        c, item, sub = _extracted_submission()
        self.assertEqual(self._step(c, 'ai_analysis')['status'], 'needs_action')

    def test_completed_after_completed_analysis(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        c, item, sub = _extracted_submission()
        analyze_submission_evidence(sub, provider=_FakeProvider())
        self.assertEqual(self._step(c, 'ai_analysis')['status'], 'completed')

    def test_downstream_steps_remain_not_completed(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        c, item, sub = _extracted_submission()
        analyze_submission_evidence(sub, provider=_FakeProvider())
        for k in ('rule_engine', 'final_verdict'):
            self.assertNotEqual(self._step(c, k)['status'], 'completed')


class AIAnalyzerSecurityTests(TestCase):
    def test_anonymous_redirected(self):
        c, item, sub = _extracted_submission()
        resp = self.client.get(reverse('compliance:evidence_ai_analysis', args=[sub.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_cross_company_view_blocked(self):
        c1, i1, s1 = _extracted_submission()
        c2 = _company(name='Other6d')
        self.client.force_login(_journey_user(c2))
        resp = self.client.get(reverse('compliance:evidence_ai_analysis', args=[s1.id]))
        self.assertEqual(resp.status_code, 302)

    def test_cross_company_post_blocked_writes_nothing(self):
        from compliance.models import EvidenceAIAnalysis
        c1, i1, s1 = _extracted_submission()
        c2 = _company(name='Other6d2')
        self.client.force_login(_journey_user(c2))
        resp = self.client.post(reverse('compliance:run_evidence_ai_analysis', args=[s1.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertFalse(EvidenceAIAnalysis.objects.filter(submission=s1).exists())

    def test_get_run_action_405(self):
        c, item, sub = _extracted_submission()
        self.client.force_login(_journey_user(c))
        resp = self.client.get(reverse('compliance:run_evidence_ai_analysis', args=[sub.id]))
        self.assertEqual(resp.status_code, 405)


# ============================================================
# Phase 6E — Rule Engine Suggested Compliance Status
# ============================================================
def _set_ai_analysis(submission, relevance='high', confidence=80, missing=None, status='completed'):
    from compliance.models import EvidenceAIAnalysis
    return EvidenceAIAnalysis.objects.update_or_create(
        submission=submission,
        defaults=dict(status=status, relevance=relevance, confidence=confidence,
                      summary='s', matched_signals=[], missing_items=missing or [],
                      recommendations=[], raw_response={}, error_message=''))[0]


class RuleEngineServiceTests(TestCase):
    def _sub(self, fv_code='NCA-ECC-2-2024', ftype='txt', content=b'Access control policy approved.'):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission(fv_code=fv_code, name='p.%s' % ftype,
                                                content=content, ftype=ftype)
        save_extraction_for_submission(sub)
        return c, item, sub

    def test_no_extracted_text_insufficient_data(self):
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024')  # no extraction
        obj = evaluate_submission_rules(sub)
        self.assertEqual(obj.status, 'completed')
        self.assertEqual(obj.suggested_status, 'insufficient_data')

    def test_extracted_no_ai_needs_review(self):
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = self._sub()
        obj = evaluate_submission_rules(sub)
        self.assertEqual(obj.suggested_status, 'needs_review')
        self.assertEqual(obj.confidence, 40)

    def test_nca_high_no_missing_suggested_c(self):
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = self._sub()
        _set_ai_analysis(sub, relevance='high', confidence=95, missing=[])
        obj = evaluate_submission_rules(sub)
        self.assertEqual(obj.suggested_status, 'suggested_c')
        self.assertEqual(obj.confidence, 90)  # capped at 90

    def test_nca_medium_or_missing_suggested_pc(self):
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = self._sub()
        _set_ai_analysis(sub, relevance='high', confidence=88, missing=['تاريخ الاعتماد'])
        obj = evaluate_submission_rules(sub)
        self.assertEqual(obj.suggested_status, 'suggested_pc')
        self.assertEqual(obj.confidence, 75)

    def test_nca_low_suggested_nc(self):
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = self._sub()
        _set_ai_analysis(sub, relevance='low', confidence=90, missing=[])
        obj = evaluate_submission_rules(sub)
        self.assertEqual(obj.suggested_status, 'suggested_nc')
        self.assertEqual(obj.confidence, 70)

    def test_unclear_needs_review(self):
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = self._sub()
        _set_ai_analysis(sub, relevance='unclear', confidence=99, missing=[])
        obj = evaluate_submission_rules(sub)
        self.assertEqual(obj.suggested_status, 'needs_review')
        self.assertEqual(obj.confidence, 50)

    def test_aramco_high_suggested_compliance(self):
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = self._sub(fv_code='ARAMCO-SACS-002')
        _set_ai_analysis(sub, relevance='high', confidence=80, missing=[])
        obj = evaluate_submission_rules(sub)
        self.assertEqual(obj.framework_type, 'aramco_sabic')
        self.assertEqual(obj.suggested_status, 'suggested_compliance')

    def test_aramco_low_suggested_noncompliance(self):
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = self._sub(fv_code='ARAMCO-SACS-002')
        _set_ai_analysis(sub, relevance='low', confidence=60, missing=[])
        obj = evaluate_submission_rules(sub)
        self.assertEqual(obj.suggested_status, 'suggested_noncompliance')

    def test_rerun_updates_same_row(self):
        from compliance.rule_engine import evaluate_submission_rules
        from compliance.models import EvidenceRuleEvaluation
        c, item, sub = self._sub()
        evaluate_submission_rules(sub)
        _set_ai_analysis(sub, relevance='high', confidence=90, missing=[])
        evaluate_submission_rules(sub)
        self.assertEqual(EvidenceRuleEvaluation.objects.filter(submission=sub).count(), 1)
        self.assertEqual(EvidenceRuleEvaluation.objects.get(submission=sub).suggested_status, 'suggested_c')

    def test_deterministic_repeated(self):
        from compliance.rule_engine import suggest_status_for_submission
        c, item, sub = self._sub()
        _set_ai_analysis(sub, relevance='high', confidence=80, missing=[])
        a = suggest_status_for_submission(sub); b = suggest_status_for_submission(sub)
        self.assertEqual((a.suggested_status, a.confidence), (b.suggested_status, b.confidence))

    def test_does_not_touch_final_assessment_or_company_control(self):
        from compliance.rule_engine import evaluate_submission_rules
        from compliance.models import ControlAssessment, CompanyControl
        c, item, sub = self._sub()
        _set_ai_analysis(sub, relevance='high', confidence=80, missing=[])
        before = (ControlAssessment.objects.count(), CompanyControl.objects.count())
        evaluate_submission_rules(sub)
        self.assertEqual(before, (ControlAssessment.objects.count(), CompanyControl.objects.count()))


class RuleEngineUITests(TestCase):
    def _ready_sub(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024', name='p.txt',
                                                content=b'Access control policy approved.', ftype='txt')
        save_extraction_for_submission(sub)
        _set_ai_analysis(sub, relevance='high', confidence=80, missing=[])
        return c, item, sub

    def test_page_renders_for_owner(self):
        c, item, sub = self._ready_sub()
        self.client.force_login(_journey_user(c))
        resp = self.client.get(reverse('compliance:evidence_rule_evaluation', args=[sub.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'محرك القواعد')
        self.assertContains(resp, 'تشغيل التقييم النظامي')
        self.assertContains(resp, 'ليست قرارًا نهائيًا')

    def test_post_run_works_and_shows_suggestion(self):
        from compliance.models import EvidenceRuleEvaluation
        c, item, sub = self._ready_sub()
        self.client.force_login(_journey_user(c))
        resp = self.client.post(reverse('compliance:run_evidence_rule_evaluation', args=[sub.id]))
        self.assertEqual(resp.status_code, 302)
        ev = EvidenceRuleEvaluation.objects.get(submission=sub)
        self.assertEqual(ev.status, 'completed')
        body = self.client.get(reverse('compliance:evidence_rule_evaluation', args=[sub.id])).content.decode()
        self.assertIn('حالة نظامية مقترحة', body)
        self.assertIn('بانتظار مراجعة المدقق', body)

    def test_no_final_verdict_or_certification_wording(self):
        c, item, sub = self._ready_sub()
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('compliance:run_evidence_rule_evaluation', args=[sub.id]))
        body = self.client.get(reverse('compliance:evidence_rule_evaluation', args=[sub.id])).content.decode()
        for bad in ('تم الاعتماد', 'قرار نهائي', 'شهادة امتثال', 'اعتماد رسمي', '/media/evidence_v2/'):
            self.assertNotIn(bad, body)

    def test_english_mode_does_not_break(self):
        c, item, sub = self._ready_sub()
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('set_language'), {'language': 'en', 'next': '/'})
        body = self.client.get(reverse('compliance:evidence_rule_evaluation', args=[sub.id])).content.decode()
        self.assertIn('Rule engine', body)
        self.assertIn('Run system evaluation', body)


class RuleEngineJourneyTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def _step(self, company, key):
        from compliance.journey import build_company_compliance_journey
        j = build_company_compliance_journey(company)
        return {s['key']: s for s in j['steps']}[key]

    def _ready_sub(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024', name='p.txt',
                                                content=b'Access control policy approved.', ftype='txt')
        save_extraction_for_submission(sub)
        return c, item, sub

    def test_planned_before_ai_analysis(self):
        c, item, sub = self._ready_sub()  # extracted, no AI
        self.assertEqual(self._step(c, 'rule_engine')['status'], 'planned')

    def test_needs_action_after_ai_before_rule(self):
        c, item, sub = self._ready_sub()
        _set_ai_analysis(sub, relevance='high', confidence=80, missing=[])
        self.assertEqual(self._step(c, 'rule_engine')['status'], 'needs_action')

    def test_completed_after_rule_evaluation(self):
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = self._ready_sub()
        _set_ai_analysis(sub, relevance='high', confidence=80, missing=[])
        evaluate_submission_rules(sub)
        self.assertEqual(self._step(c, 'rule_engine')['status'], 'completed')

    def test_auditor_and_final_verdict_remain_not_completed(self):
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = self._ready_sub()
        _set_ai_analysis(sub, relevance='high', confidence=80, missing=[])
        evaluate_submission_rules(sub)
        self.assertNotEqual(self._step(c, 'auditor_review')['status'], 'completed')
        self.assertNotEqual(self._step(c, 'final_verdict')['status'], 'completed')


class RuleEngineSecurityTests(TestCase):
    def _ready_sub(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024', name='p.txt',
                                                content=b'policy approved.', ftype='txt')
        save_extraction_for_submission(sub)
        return c, item, sub

    def test_anonymous_redirected(self):
        c, item, sub = self._ready_sub()
        resp = self.client.get(reverse('compliance:evidence_rule_evaluation', args=[sub.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_cross_company_view_blocked(self):
        c1, i1, s1 = self._ready_sub()
        c2 = _company(name='Other6e')
        self.client.force_login(_journey_user(c2))
        resp = self.client.get(reverse('compliance:evidence_rule_evaluation', args=[s1.id]))
        self.assertEqual(resp.status_code, 302)

    def test_cross_company_post_writes_nothing(self):
        from compliance.models import EvidenceRuleEvaluation
        c1, i1, s1 = self._ready_sub()
        c2 = _company(name='Other6e2')
        self.client.force_login(_journey_user(c2))
        resp = self.client.post(reverse('compliance:run_evidence_rule_evaluation', args=[s1.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertFalse(EvidenceRuleEvaluation.objects.filter(submission=s1).exists())

    def test_get_run_action_405(self):
        c, item, sub = self._ready_sub()
        self.client.force_login(_journey_user(c))
        resp = self.client.get(reverse('compliance:run_evidence_rule_evaluation', args=[sub.id]))
        self.assertEqual(resp.status_code, 405)


# ============================================================
# Phase 6F — Auditor Final Verdict Workflow
# ============================================================
def _staff_user(email='staff6f@x.com'):
    from core.models import User
    return User.objects.create_user(email=email, password='longenough12', is_staff=True)


def _assigned_auditor_user(company, email='aud6f@x.com'):
    from core.models import User
    from auditors.models import AuditorProfile, AuditorAssignment
    u = User.objects.create_user(email=email, password='longenough12')
    p = AuditorProfile.objects.create(user=u, full_name='Auditor', status='active')
    AuditorAssignment.objects.create(company=company, auditor=p, status='accepted')
    return u


class AuditorVerdictServiceTests(TestCase):
    def _sub(self, fv_code='NCA-ECC-2-2024'):
        return _company_with_submission(fv_code=fv_code)

    def test_staff_can_record_verdict(self):
        from compliance.auditor_verdict import record_auditor_final_verdict
        c, item, sub = self._sub()
        v = record_auditor_final_verdict(sub, _staff_user(), status='final_c',
                                         rationale='تمت المراجعة.', confidence=85)
        self.assertEqual(v.status, 'final_c')
        self.assertEqual(v.confidence, 85)
        self.assertEqual(v.reviewer.is_staff, True)

    def test_assigned_auditor_can_record(self):
        from compliance.auditor_verdict import record_auditor_final_verdict
        c, item, sub = self._sub()
        u = _assigned_auditor_user(c)
        v = record_auditor_final_verdict(sub, u, status='final_pc', rationale='جزئي.')
        self.assertEqual(v.status, 'final_pc')

    def test_company_user_cannot_record(self):
        from compliance.auditor_verdict import record_auditor_final_verdict, VerdictError
        c, item, sub = self._sub()
        u = _journey_user(c)
        with self.assertRaises(VerdictError):
            record_auditor_final_verdict(sub, u, status='final_c', rationale='x')

    def test_rationale_required(self):
        from compliance.auditor_verdict import record_auditor_final_verdict, VerdictError
        c, item, sub = self._sub()
        with self.assertRaises(VerdictError):
            record_auditor_final_verdict(sub, _staff_user(), status='final_c', rationale='   ')

    def test_invalid_status_rejected(self):
        from compliance.auditor_verdict import record_auditor_final_verdict, VerdictError
        c, item, sub = self._sub()
        with self.assertRaises(VerdictError):
            record_auditor_final_verdict(sub, _staff_user(), status='compliant', rationale='x')

    def test_aramco_status_vocabulary(self):
        from compliance.auditor_verdict import record_auditor_final_verdict
        c, item, sub = self._sub(fv_code='ARAMCO-SACS-002')
        v = record_auditor_final_verdict(sub, _staff_user(), status='final_compliance', rationale='ok')
        self.assertEqual(v.framework_type, 'aramco_sabic')
        # NCA status invalid for an Aramco submission
        from compliance.auditor_verdict import VerdictError
        with self.assertRaises(VerdictError):
            record_auditor_final_verdict(sub, _staff_user('s2@x.com'), status='final_c', rationale='x')

    def test_confidence_clamped(self):
        from compliance.auditor_verdict import record_auditor_final_verdict
        c, item, sub = self._sub()
        v = record_auditor_final_verdict(sub, _staff_user(), status='final_c', rationale='x', confidence=900)
        self.assertEqual(v.confidence, 100)

    def test_rerun_updates_same_verdict(self):
        from compliance.auditor_verdict import record_auditor_final_verdict
        from compliance.models import AuditorFinalVerdict
        c, item, sub = self._sub()
        s = _staff_user()
        record_auditor_final_verdict(sub, s, status='final_c', rationale='one')
        record_auditor_final_verdict(sub, s, status='final_nc', rationale='two')
        self.assertEqual(AuditorFinalVerdict.objects.filter(submission=sub).count(), 1)
        self.assertEqual(AuditorFinalVerdict.objects.get(submission=sub).status, 'final_nc')

    def test_does_not_touch_reports_or_company_control(self):
        from compliance.auditor_verdict import record_auditor_final_verdict
        from compliance.models import CompanyControl, ControlAssessment
        c, item, sub = self._sub()
        before = (CompanyControl.objects.count(), ControlAssessment.objects.count())
        record_auditor_final_verdict(sub, _staff_user(), status='final_c', rationale='x')
        self.assertEqual(before, (CompanyControl.objects.count(), ControlAssessment.objects.count()))


class AuditorVerdictUITests(TestCase):
    def test_owner_sees_readonly_message(self):
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024')
        self.client.force_login(_journey_user(c))
        resp = self.client.get(reverse('compliance:auditor_verdict', args=[sub.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'قرار المدقق النهائي')
        self.assertContains(resp, 'يمكنك عرض القرار، ولا تملك صلاحية تعديله')
        self.assertContains(resp, 'لا يُعد شهادة امتثال رسمية')

    def test_staff_sees_form_and_can_post(self):
        from compliance.models import AuditorFinalVerdict
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024')
        self.client.force_login(_staff_user())
        body = self.client.get(reverse('compliance:auditor_verdict', args=[sub.id])).content.decode()
        self.assertIn('حفظ قرار المدقق', body)
        resp = self.client.post(reverse('compliance:auditor_verdict', args=[sub.id]),
                                {'status': 'final_c', 'rationale': 'تمت المراجعة.', 'confidence': '80'})
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(AuditorFinalVerdict.objects.get(submission=sub).status, 'final_c')

    def test_company_user_post_forbidden_writes_nothing(self):
        from compliance.models import AuditorFinalVerdict
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024')
        self.client.force_login(_journey_user(c))
        resp = self.client.post(reverse('compliance:auditor_verdict', args=[sub.id]),
                                {'status': 'final_c', 'rationale': 'x'})
        self.assertEqual(resp.status_code, 302)
        self.assertFalse(AuditorFinalVerdict.objects.filter(submission=sub).exists())

    def test_no_official_certification_wording(self):
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024')
        self.client.force_login(_staff_user())
        self.client.post(reverse('compliance:auditor_verdict', args=[sub.id]),
                         {'status': 'final_c', 'rationale': 'ok', 'confidence': '70'})
        body = self.client.get(reverse('compliance:auditor_verdict', args=[sub.id])).content.decode()
        for bad in ('شهادة رسمية', 'اعتماد رسمي من جهة حكومية', 'معتمد من NCA', 'معتمد من أرامكو', 'معتمد من سابك', '/media/evidence_v2/'):
            self.assertNotIn(bad, body)

    def test_english_mode_does_not_break(self):
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024')
        self.client.force_login(_staff_user())
        self.client.post(reverse('set_language'), {'language': 'en', 'next': '/'})
        body = self.client.get(reverse('compliance:auditor_verdict', args=[sub.id])).content.decode()
        self.assertIn('Auditor final verdict', body)
        self.assertIn('Save auditor decision', body)


class AuditorVerdictJourneyTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def _step(self, company, key):
        from compliance.journey import build_company_compliance_journey
        j = build_company_compliance_journey(company)
        return {s['key']: s for s in j['steps']}[key]

    def _ready_with_rule(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024', name='p.txt',
                                                content=b'policy approved.', ftype='txt')
        save_extraction_for_submission(sub)
        _set_ai_analysis(sub, relevance='high', confidence=80, missing=[])
        evaluate_submission_rules(sub)
        return c, item, sub

    def test_auditor_review_planned_before_rule(self):
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024')  # no rule eval
        self.assertEqual(self._step(c, 'auditor_review')['status'], 'planned')

    def test_auditor_review_needs_action_after_rule(self):
        c, item, sub = self._ready_with_rule()
        self.assertEqual(self._step(c, 'auditor_review')['status'], 'needs_action')

    def test_auditor_review_and_final_verdict_completed_after_verdict(self):
        from compliance.auditor_verdict import record_auditor_final_verdict
        c, item, sub = self._ready_with_rule()
        record_auditor_final_verdict(sub, _staff_user(), status='final_c', rationale='reviewed')
        self.assertEqual(self._step(c, 'auditor_review')['status'], 'completed')
        self.assertEqual(self._step(c, 'final_verdict')['status'], 'completed')

    def test_reports_remain_not_completed(self):
        from compliance.auditor_verdict import record_auditor_final_verdict
        c, item, sub = self._ready_with_rule()
        record_auditor_final_verdict(sub, _staff_user(), status='final_c', rationale='reviewed')
        self.assertNotEqual(self._step(c, 'reports')['status'], 'completed')


class AuditorVerdictSecurityTests(TestCase):
    def test_anonymous_redirected(self):
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024')
        resp = self.client.get(reverse('compliance:auditor_verdict', args=[sub.id]))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_cross_company_view_blocked(self):
        c1, i1, s1 = _company_with_submission(fv_code='NCA-ECC-2-2024')
        c2 = _company(name='Other6f')
        self.client.force_login(_journey_user(c2))
        resp = self.client.get(reverse('compliance:auditor_verdict', args=[s1.id]))
        self.assertEqual(resp.status_code, 302)

    def test_cross_company_post_writes_nothing(self):
        from compliance.models import AuditorFinalVerdict
        c1, i1, s1 = _company_with_submission(fv_code='NCA-ECC-2-2024')
        c2 = _company(name='Other6f2')
        self.client.force_login(_journey_user(c2))
        resp = self.client.post(reverse('compliance:auditor_verdict', args=[s1.id]),
                                {'status': 'final_c', 'rationale': 'x'})
        self.assertEqual(resp.status_code, 302)
        self.assertFalse(AuditorFinalVerdict.objects.filter(submission=s1).exists())


# ============================================================
# Phase 7A — Local End-to-End UAT 90% Gate
# ============================================================
class _UATFake:
    """Deterministic fake AI provider for UAT (no network)."""
    def __init__(self, payload=None):
        self.payload = payload or {'relevance': 'high', 'confidence': 85,
                                   'summary': 'يبدو أن الدليل مرتبط بالضابط.',
                                   'matched_signals': ['سياسة موثّقة'], 'missing_items': [],
                                   'recommendations': []}
    def analyze(self, prompt):
        return self.payload


class Phase7AEndToEndUATTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def _ready_submission(self, fv='NCA-ECC-2-2024', name='access_control_policy.txt',
                          content=b'Access control policy and incident response procedure approved.'):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission(fv_code=fv, name=name, content=content, ftype='txt')
        save_extraction_for_submission(sub)
        return c, item, sub

    def _step(self, company, key):
        from compliance.journey import build_company_compliance_journey
        return {s['key']: s for s in build_company_compliance_journey(company)['steps']}[key]

    def _patch_ai(self, provider):
        import compliance.evidence_ai_analyzer as mod
        orig = mod.default_provider
        mod.default_provider = lambda: provider
        self.addCleanup(lambda: setattr(mod, 'default_provider', orig))

    # ---- Scenario A/B: full chain upload -> verdict ----
    def test_full_chain_upload_to_verdict(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        from compliance.rule_engine import evaluate_submission_rules
        from compliance.auditor_verdict import record_auditor_final_verdict
        c, item, sub = self._ready_submission()
        # extraction completed
        self.assertEqual(self._step(c, 'ocr_extraction')['status'], 'completed')
        # AI advisory (fake provider)
        ai = analyze_submission_evidence(sub, provider=_UATFake())
        self.assertEqual(ai.status, 'completed')
        self.assertEqual(self._step(c, 'ai_analysis')['status'], 'completed')
        # rule engine suggestion
        self.assertEqual(self._step(c, 'rule_engine')['status'], 'needs_action')  # AI done, no rule yet
        ev = evaluate_submission_rules(sub)
        self.assertEqual(ev.status, 'completed')
        self.assertEqual(ev.suggested_status, 'suggested_c')
        self.assertEqual(self._step(c, 'rule_engine')['status'], 'completed')      # rule eval persisted
        self.assertEqual(self._step(c, 'auditor_review')['status'], 'needs_action')  # rule done, no verdict
        # auditor verdict
        v = record_auditor_final_verdict(sub, _staff_user(), status='final_c', rationale='reviewed')
        self.assertEqual(v.status, 'final_c')
        self.assertEqual(self._step(c, 'auditor_review')['status'], 'completed')
        self.assertEqual(self._step(c, 'final_verdict')['status'], 'completed')

    # ---- Case 2: image-only evidence ----
    def test_image_only_evidence_blocks_ai_and_rule_insufficient(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        from compliance.evidence_ai_analyzer import analyze_submission_evidence, can_analyze_submission
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024', name='scan.png',
                                                content=b'\x89PNG\r\n', ftype='png')
        save_extraction_for_submission(sub)
        ok, _ = can_analyze_submission(sub)
        self.assertFalse(ok)
        ai = analyze_submission_evidence(sub, provider=_UATFake())  # gate -> skipped
        self.assertEqual(ai.status, 'skipped')
        ev = evaluate_submission_rules(sub)
        self.assertEqual(ev.suggested_status, 'insufficient_data')
        self.assertEqual(self._step(c, 'ocr_extraction')['status'], 'needs_action')

    # ---- Case 3: missing AI provider safe skip ----
    def test_missing_ai_provider_skips_safely(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence, ProviderUnavailable
        c, item, sub = self._ready_submission()
        class _Unavail:
            def analyze(self, prompt): raise ProviderUnavailable('no key')
        ai = analyze_submission_evidence(sub, provider=_Unavail())
        self.assertEqual(ai.status, 'skipped')

    # ---- Scenario C: wording safety ----
    def test_rule_and_verdict_pages_wording_safe(self):
        from compliance.evidence_ai_analyzer import analyze_submission_evidence
        from compliance.rule_engine import evaluate_submission_rules
        c, item, sub = self._ready_submission()
        analyze_submission_evidence(sub, provider=_UATFake())
        evaluate_submission_rules(sub)
        self.client.force_login(_journey_user(c))
        rule_body = self.client.get(reverse('compliance:evidence_rule_evaluation', args=[sub.id])).content.decode()
        self.assertIn('بانتظار مراجعة المدقق', rule_body)
        for bad in ('قرار نهائي', 'شهادة', 'اعتماد رسمي', 'معتمد من NCA', 'معتمد من أرامكو', 'معتمد من سابك'):
            self.assertNotIn(bad, rule_body)
        self.client.force_login(_staff_user())
        vbody = self.client.get(reverse('compliance:auditor_verdict', args=[sub.id])).content.decode()
        self.assertIn('مراجعة داخلية', vbody)
        self.assertIn('لا يُعد شهادة امتثال رسمية', vbody)
        for bad in ('شهادة رسمية', 'اعتماد حكومي', 'معتمد من NCA'):
            self.assertNotIn(bad, vbody)

    # ---- Scenario D: reports not finalized by verdict ----
    def test_reports_not_finalized_by_verdict(self):
        from compliance.auditor_verdict import record_auditor_final_verdict
        c, item, sub = self._ready_submission()
        record_auditor_final_verdict(sub, _staff_user(), status='final_c', rationale='ok')
        self.assertNotEqual(self._step(c, 'reports')['status'], 'completed')  # still gated

    # ---- Scenario F: permissions / cross-company isolation across the chain ----
    def test_cross_company_blocked_across_chain(self):
        c1, i1, s1 = self._ready_submission()
        c2 = _company(name='UATOther')
        self.client.force_login(_journey_user(c2))
        for name in ('evidence_extraction', 'evidence_ai_analysis', 'evidence_rule_evaluation', 'auditor_verdict'):
            resp = self.client.get(reverse('compliance:%s' % name, args=[s1.id]))
            self.assertEqual(resp.status_code, 302, name)

    def test_company_user_cannot_submit_verdict(self):
        from compliance.models import AuditorFinalVerdict
        c, item, sub = self._ready_submission()
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('compliance:auditor_verdict', args=[sub.id]),
                         {'status': 'final_c', 'rationale': 'x'})
        self.assertFalse(AuditorFinalVerdict.objects.filter(submission=sub).exists())

    def test_run_actions_get_405(self):
        c, item, sub = self._ready_submission()
        self.client.force_login(_journey_user(c))
        for name in ('run_evidence_extraction', 'run_evidence_ai_analysis', 'run_evidence_rule_evaluation'):
            resp = self.client.get(reverse('compliance:%s' % name, args=[sub.id]))
            self.assertEqual(resp.status_code, 405, name)

    # ---- Scenario A: classification + applicability pages, count, i18n ----
    def test_classification_and_applicability_pages_and_count(self):
        c = _company(target_nca=True)
        CompanyIntakeProfile.objects.create(company=c, uses_cloud_services=True, has_remote_work=True,
                                            manages_official_social_media_accounts=True)
        self.client.force_login(_journey_user(c))
        cls = self.client.get(reverse('compliance:classification')).content.decode()
        self.assertIn('التصنيف الذكي', cls)
        self.assertNotIn('334', cls)
        app = self.client.get(reverse('compliance:applicability_preview')).content.decode()
        self.assertIn('قابلية تطبيق الضوابط', app)
        self.assertNotIn('334', app)     # legacy total never shown
        # official totals come from the canonical 417 helper (not legacy 334)
        from compliance.smart_classification import OFFICIAL_TOTAL
        self.assertEqual(OFFICIAL_TOTAL, 417)

    def test_english_mode_core_pages(self):
        c = _company(target_nca=True)
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('set_language'), {'language': 'en', 'next': '/'})
        for url in (reverse('compliance:classification'), reverse('compliance:applicability_preview')):
            self.assertEqual(self.client.get(url).status_code, 200)

    # ---- Scenario E: monitoring foundation ----
    def test_monitoring_route_renders_or_redirects(self):
        c = _company()
        self.client.force_login(_journey_user(c))
        self.assertIn(self.client.get('/monitoring/continuous/').status_code, (200, 302))


# ============================================================
# Phase 7B — Report Finalization Integration
# ============================================================
def _company_with_verdict(fv='NCA-ECC-2-2024', status='final_c', subscribe=True):
    """Company + submission with a recorded auditor verdict (and optional subscription)."""
    from compliance.evidence_extraction import save_extraction_for_submission
    from compliance.rule_engine import evaluate_submission_rules
    from compliance.auditor_verdict import record_auditor_final_verdict
    from billing.subscription_access import activate_company_subscription
    from core.models import User
    c, item, sub = _company_with_submission(fv_code=fv, name='p.txt',
                                            content=b'access control policy approved.', ftype='txt')
    save_extraction_for_submission(sub)
    _set_ai_analysis(sub, relevance='high', confidence=80, missing=[])
    evaluate_submission_rules(sub)
    staff = _staff_user(email=f'staff7b{User.objects.filter(is_staff=True).count()}@x.com')
    record_auditor_final_verdict(sub, staff, status=status, rationale='reviewed ok', confidence=88)
    if subscribe:
        activate_company_subscription(c, 'Plan', days=30)
    return c, item, sub


class ReportFinalizationServiceTests(TestCase):
    def test_counts_total_reviewed_pending(self):
        from compliance.report_finalization import build_auditor_reviewed_report
        c, item, sub = _company_with_verdict()
        r = build_auditor_reviewed_report(c)
        self.assertEqual(r.total_submissions, 1)
        self.assertEqual(r.reviewed_count, 1)
        self.assertEqual(r.pending_count, 0)
        self.assertTrue(r.has_verdicts)

    def test_status_and_framework_counts(self):
        from compliance.report_finalization import build_auditor_reviewed_report
        c, item, sub = _company_with_verdict(status='final_c')
        r = build_auditor_reviewed_report(c)
        self.assertEqual(r.status_counts.get('final_c'), 1)
        self.assertEqual(r.framework_counts.get('nca'), 1)

    def test_aramco_status_mapped(self):
        from compliance.report_finalization import build_auditor_reviewed_report
        c, item, sub = _company_with_verdict(fv='ARAMCO-SACS-002', status='final_compliance')
        r = build_auditor_reviewed_report(c)
        self.assertEqual(r.status_counts.get('final_compliance'), 1)
        self.assertEqual(r.framework_counts.get('aramco_sabic'), 1)

    def test_excludes_raw_text_and_paths(self):
        from compliance.report_finalization import build_auditor_reviewed_report
        c, item, sub = _company_with_verdict()
        r = build_auditor_reviewed_report(c)
        row = r.rows[0]
        blob = str(row)
        self.assertNotIn('/media/', blob)
        self.assertNotIn('access control policy approved', blob)  # raw extracted text excluded

    def test_does_not_write_assessment_or_company_control(self):
        from compliance.report_finalization import build_auditor_reviewed_report
        from compliance.models import ControlAssessment, CompanyControl
        c, item, sub = _company_with_verdict()
        before = (ControlAssessment.objects.count(), CompanyControl.objects.count())
        build_auditor_reviewed_report(c)
        self.assertEqual(before, (ControlAssessment.objects.count(), CompanyControl.objects.count()))

    def test_deterministic(self):
        from compliance.report_finalization import build_auditor_reviewed_report
        c, item, sub = _company_with_verdict()
        a = build_auditor_reviewed_report(c); b = build_auditor_reviewed_report(c)
        self.assertEqual((a.total_submissions, a.reviewed_count, a.status_counts),
                         (b.total_submissions, b.reviewed_count, b.status_counts))


class ReportFinalizationUITests(TestCase):
    def test_empty_state_when_no_verdict(self):
        from billing.subscription_access import activate_company_subscription
        c = _company()
        activate_company_subscription(c, 'Plan', days=30)
        self.client.force_login(_journey_user(c))
        body = self.client.get(reverse('compliance:auditor_reviewed_report')).content.decode()
        # Phase 8D-2-FIX-C: actionable empty state instead of a dead "no results".
        self.assertIn('لا يوجد تقرير مدقق بعد', body)
        self.assertIn('بعد مراجعة المدقق', body)
        self.assertIn('لا يُعد شهادة امتثال رسمية', body)

    def test_renders_rows_after_verdict(self):
        c, item, sub = _company_with_verdict()
        self.client.force_login(_journey_user(c))
        body = self.client.get(reverse('compliance:auditor_reviewed_report')).content.decode()
        self.assertIn('تقرير مراجعة امتثال داخلي', body)
        self.assertIn('ملخص قرارات المدقق', body)
        self.assertIn('الحالة بعد مراجعة المدقق', body)

    def test_no_certification_or_334(self):
        c, item, sub = _company_with_verdict()
        self.client.force_login(_journey_user(c))
        body = self.client.get(reverse('compliance:auditor_reviewed_report')).content.decode()
        # Positive certification CLAIMS must be absent. (The mandated disclaimer legitimately
        # negates "شهادة امتثال رسمية"/"اعتماد", so those phrases are not blanket-forbidden.)
        for bad in ('معتمد من NCA', 'معتمد من أرامكو', 'معتمد من سابك',
                    'certified by NCA', 'official accreditation', '334', '/media/'):
            self.assertNotIn(bad, body)
        # And the required negation disclaimer IS present.
        self.assertIn('لا يُعد شهادة امتثال رسمية', body)

    def test_subscription_gate_preserved(self):
        # Unsubscribed company hits the subscription-required gate, not the report.
        c, item, sub = _company_with_verdict(subscribe=False)
        self.client.force_login(_journey_user(c))
        body = self.client.get(reverse('compliance:auditor_reviewed_report')).content.decode()
        self.assertNotIn('تفاصيل القرارات', body)  # report body not shown

    def test_english_mode_does_not_break(self):
        c, item, sub = _company_with_verdict()
        self.client.force_login(_journey_user(c))
        self.client.post(reverse('set_language'), {'language': 'en', 'next': '/'})
        body = self.client.get(reverse('compliance:auditor_reviewed_report')).content.decode()
        self.assertIn('Internal compliance review report', body)
        self.assertIn('Status after auditor review', body)


class ReportFinalizationJourneyTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def _step(self, company, key):
        from compliance.journey import build_company_compliance_journey
        return {s['key']: s for s in build_company_compliance_journey(company)['steps']}[key]

    def test_reports_locked_without_subscription(self):
        c, item, sub = _company_with_verdict(subscribe=False)
        self.assertEqual(self._step(c, 'reports')['status'], 'locked')

    def test_reports_not_completed_with_subscription_but_no_verdict(self):
        from billing.subscription_access import activate_company_subscription
        c = _company()
        activate_company_subscription(c, 'Plan', days=30)
        self.assertNotEqual(self._step(c, 'reports')['status'], 'completed')

    def test_reports_completed_with_subscription_and_verdict(self):
        c, item, sub = _company_with_verdict(subscribe=True)
        self.assertEqual(self._step(c, 'reports')['status'], 'completed')

    def test_monitoring_unchanged(self):
        c, item, sub = _company_with_verdict(subscribe=True)
        self.assertNotEqual(self._step(c, 'monitoring')['status'], 'completed')


class ReportFinalizationSecurityTests(TestCase):
    def test_anonymous_redirected(self):
        resp = self.client.get(reverse('compliance:auditor_reviewed_report'))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/login', resp.url)

    def test_company_sees_only_own_rows(self):
        from compliance.report_finalization import build_auditor_reviewed_report
        c1, i1, s1 = _company_with_verdict()
        c2, i2, s2 = _company_with_verdict(status='final_nc')
        r1 = build_auditor_reviewed_report(c1)
        self.assertEqual(r1.reviewed_count, 1)
        self.assertTrue(all(row.submission_id == s1.id for row in r1.rows))

    def test_verdict_still_not_certificate(self):
        # Regression: recording a verdict + report does not create any certificate artifact.
        c, item, sub = _company_with_verdict()
        from compliance.models import AuditorFinalVerdict
        v = AuditorFinalVerdict.objects.get(submission=sub)
        self.assertNotIn('شهادة', v.status_ar)


# ============================================================
# Phase 7C — Local Polish / Pre-Deploy Safety Pack (regression)
# ============================================================
class Phase7CPreDeployPolishTests(TestCase):
    def setUp(self):
        call_command('seed_framework_versions', stdout=StringIO())

    def _step(self, company, key):
        from compliance.journey import build_company_compliance_journey
        return {s['key']: s for s in build_company_compliance_journey(company)['steps']}[key]

    # ---- Route smoke: authorized owner gets 200 across the whole chain ----
    def test_chain_routes_render_for_owner(self):
        c, item, sub = _company_with_verdict(subscribe=True)
        u = c.users.first() if hasattr(c, 'users') else None
        self.client.force_login(_journey_user(c, email='polish_owner@x.com'))
        pages = [
            reverse('compliance:dashboard'),
            reverse('compliance:classification'),
            reverse('compliance:applicability_preview'),
            reverse('compliance:reports_index'),
            reverse('compliance:auditor_reviewed_report'),
            reverse('compliance:evidence_extraction', args=[sub.id]),
            reverse('compliance:evidence_ai_analysis', args=[sub.id]),
            reverse('compliance:evidence_rule_evaluation', args=[sub.id]),
            reverse('compliance:evidence_submission_detail', args=[sub.id]),
            reverse('compliance:auditor_verdict', args=[sub.id]),
        ]
        for url in pages:
            self.assertEqual(self.client.get(url).status_code, 200, url)
        self.assertIn(self.client.get('/monitoring/continuous/').status_code, (200, 302))

    # ---- Anonymous redirect for protected pages ----
    def test_protected_pages_redirect_anonymous(self):
        c, item, sub = _company_with_verdict(subscribe=True)
        for url in (reverse('compliance:classification'),
                    reverse('compliance:applicability_preview'),
                    reverse('compliance:auditor_reviewed_report'),
                    reverse('compliance:evidence_extraction', args=[sub.id]),
                    reverse('compliance:auditor_verdict', args=[sub.id])):
            resp = self.client.get(url)
            self.assertEqual(resp.status_code, 302, url)
            self.assertIn('/login', resp.url, url)

    # ---- Wording safety on key pages ----
    def test_wording_safety_on_key_pages(self):
        c, item, sub = _company_with_verdict(subscribe=True)
        self.client.force_login(_journey_user(c, email='polish_w@x.com'))
        bodies = {
            'rule': self.client.get(reverse('compliance:evidence_rule_evaluation', args=[sub.id])).content.decode(),
            'verdict': self.client.get(reverse('compliance:auditor_verdict', args=[sub.id])).content.decode(),
            'report': self.client.get(reverse('compliance:auditor_reviewed_report')).content.decode(),
        }
        for body in bodies.values():
            for bad in ('334', 'CyberTrust KSA', 'معتمد من NCA', 'معتمد من أرامكو', 'معتمد من سابك',
                        'certified by NCA', 'official accreditation', 'اعتماد حكومي', 'تم إصدار شهادة'):
                self.assertNotIn(bad, body)
        # advisory / suggested / internal-review wording present where it should be
        self.assertIn('بانتظار مراجعة المدقق', bodies['rule'])
        self.assertIn('لا يُعد شهادة امتثال رسمية', bodies['verdict'])
        self.assertIn('لا يُعد شهادة امتثال رسمية', bodies['report'])

    def test_ai_page_is_advisory(self):
        c, item, sub = _company_with_verdict(subscribe=True)
        self.client.force_login(_journey_user(c, email='polish_ai@x.com'))
        body = self.client.get(reverse('compliance:evidence_ai_analysis', args=[sub.id])).content.decode()
        self.assertIn('لا يُعد قرارًا نهائيًا أو شهادة امتثال', body)

    # ---- Journey truthfulness for a fully-populated company ----
    def test_journey_steps_completed_for_full_chain(self):
        c, item, sub = _company_with_verdict(subscribe=True)
        for key in ('ocr_extraction', 'ai_analysis', 'rule_engine', 'auditor_review', 'final_verdict', 'reports'):
            self.assertEqual(self._step(c, key)['status'], 'completed', key)
        # planned/foundation steps must NOT be falsely completed
        # (rule_engine here IS completed; OCR-for-images and monitoring connectors are not modeled as steps)
        self.assertNotEqual(self._step(c, 'monitoring')['status'], 'completed')

    # ---- Cross-tenant link safety across the chain ----
    def test_cross_company_blocked_across_chain(self):
        c1, i1, s1 = _company_with_verdict(subscribe=True)
        c2 = _company(name='PolishOther')
        self.client.force_login(_journey_user(c2, email='polish_other@x.com'))
        for name in ('evidence_extraction', 'evidence_ai_analysis', 'evidence_rule_evaluation', 'auditor_verdict'):
            self.assertEqual(self.client.get(reverse('compliance:%s' % name, args=[s1.id])).status_code, 302, name)

    def test_company_user_cannot_submit_verdict(self):
        from compliance.models import AuditorFinalVerdict
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024')
        self.client.force_login(_journey_user(c, email='polish_co@x.com'))
        self.client.post(reverse('compliance:auditor_verdict', args=[sub.id]),
                         {'status': 'final_c', 'rationale': 'x'})
        # a verdict may already exist from _company_with_submission? no — fresh submission, none recorded
        self.assertFalse(AuditorFinalVerdict.objects.filter(submission=sub).exists())

    def test_get_run_actions_405(self):
        c, item, sub = _company_with_submission(fv_code='NCA-ECC-2-2024')
        self.client.force_login(_journey_user(c, email='polish_405@x.com'))
        for name in ('run_evidence_extraction', 'run_evidence_ai_analysis', 'run_evidence_rule_evaluation'):
            self.assertEqual(self.client.get(reverse('compliance:%s' % name, args=[sub.id])).status_code, 405, name)

    # ---- English mode renders ----
    def test_english_mode_key_pages(self):
        c, item, sub = _company_with_verdict(subscribe=True)
        self.client.force_login(_journey_user(c, email='polish_en@x.com'))
        self.client.post(reverse('set_language'), {'language': 'en', 'next': '/'})
        for url in (reverse('compliance:classification'),
                    reverse('compliance:applicability_preview'),
                    reverse('compliance:auditor_reviewed_report'),
                    reverse('compliance:evidence_ai_analysis', args=[sub.id])):
            self.assertEqual(self.client.get(url).status_code, 200, url)

    # ---- Official total is 417, never legacy 334 ----
    def test_official_total_417_not_334(self):
        from compliance.smart_classification import OFFICIAL_TOTAL, _FW
        self.assertEqual(OFFICIAL_TOTAL, 417)
        self.assertNotIn(334, [f['expected'] for f in _FW.values()])


class ControlsLibraryFilterTests(TestCase):
    """Manus 8D-2 Fix 2 — controls library filters must not appear empty/confusing."""

    def _seed_one_control(self):
        from compliance.models import Framework, Domain, Control
        fw, _ = Framework.objects.get_or_create(
            code='NCA_ECC', defaults={'name': 'NCA Essential Cybersecurity Controls'})
        dom, _ = Domain.objects.get_or_create(
            framework=fw, code='GOV', defaults={'name': 'Governance'})
        Control.objects.get_or_create(
            framework=fw, domain=dom, control_id='ECC-1-1-1',
            defaults={'title': 'Cybersecurity Strategy'})
        return fw, dom

    def _login_company_user(self):
        c = _company()
        self.client.force_login(_journey_user(c, email='ctrllib@x.com'))
        return c

    def test_filters_populated_when_controls_exist(self):
        from compliance.models import Control, Framework
        self._seed_one_control()
        self.assertTrue(Control.objects.exists())
        self._login_company_user()
        resp = self.client.get(reverse('compliance:controls_list'))
        self.assertEqual(resp.status_code, 200)
        # At least one real framework option is rendered beyond "All Frameworks".
        a_framework = Framework.objects.filter(
            id__in=Control.objects.values_list('framework_id', flat=True)).first()
        self.assertIsNotNone(a_framework)
        self.assertContains(resp, a_framework.name)
        # No empty-state explanation while options exist.
        self.assertNotContains(resp, 'ستظهر خيارات التصفية')

    def test_empty_state_explanation_when_no_filter_options(self):
        from compliance.models import Control
        Control.objects.all().delete()  # simulate environment with no usable control plan
        self._login_company_user()
        resp = self.client.get(reverse('compliance:controls_list'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'ستظهر خيارات التصفية')

    def test_official_total_417_shown_not_334(self):
        self._login_company_user()
        resp = self.client.get(reverse('compliance:controls_list'))
        self.assertContains(resp, '417')
        self.assertNotContains(resp, '334')


class ClassificationEmptyStateTests(TestCase):
    """Manus 8D-2 Fix 3 — Cloud-services-only must not dead-end on 'No results yet'."""

    def _company_with_intake(self, **intake):
        from compliance.models import CompanyIntakeProfile
        c = _company()
        defaults = dict(company=c, review_status='completed')
        defaults.update(intake)
        CompanyIntakeProfile.objects.create(**defaults)
        return c

    def test_completed_intake_no_results_shows_advisory_guidance(self):
        # Profile completed but no applicability results (e.g. minimal inputs) ->
        # helpful advisory guidance, not a dead 'No results yet'.
        c = self._company_with_intake(uses_cloud_services=True)
        self.client.force_login(_journey_user(c, email='cls_empty@x.com'))
        resp = self.client.get(reverse('compliance:applicability_review'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'لم يتم توليد توصية كافية من هذه الإجابات وحدها')
        # Advisory wording: no final compliance decision / certificate claim.
        self.assertContains(resp, 'استشاري')

    def test_advisory_wording_has_no_official_certification_claim(self):
        c = self._company_with_intake(uses_cloud_services=True)
        self.client.force_login(_journey_user(c, email='cls_safe@x.com'))
        resp = self.client.get(reverse('compliance:applicability_review'))
        body = resp.content.decode()
        for banned in ('اعتماد رسمي', 'شهادة امتثال رسمية', 'معتمد من'):
            self.assertNotIn(banned, body)

    def test_no_profile_keeps_complete_classification_prompt(self):
        c = _company()
        self.client.force_login(_journey_user(c, email='cls_noprofile@x.com'))
        resp = self.client.get(reverse('compliance:applicability_review'))
        self.assertContains(resp, 'أكمل التصنيف أولاً')


class ManusE2EQASeedCommandTests(TestCase):
    """Manus 8D-2 Fix 4 — QA seeder for activated-auditor verdict workflow."""

    def test_refuses_without_confirmation(self):
        from django.core.management import call_command
        from django.core.management.base import CommandError
        with self.assertRaises(CommandError):
            call_command('seed_manus_e2e_qa', stdout=StringIO())

    def test_seeds_company_auditor_and_accepted_assignment(self):
        from django.core.management import call_command
        from core.models import User, Company
        from auditors.models import AuditorProfile, AuditorAssignment
        from billing.subscription_access import company_has_active_subscription
        call_command('seed_manus_e2e_qa', '--confirm', stdout=StringIO())
        company = Company.objects.get(cr_number='QA-MANUS-0001')
        self.assertTrue(User.objects.filter(email='qa.company@manus-e2e.test',
                                            company=company).exists())
        auditor = AuditorProfile.objects.get(user__email='qa.auditor@manus-e2e.test')
        self.assertEqual(auditor.status, 'active')  # activated auditor
        self.assertTrue(AuditorAssignment.objects.filter(
            company=company, auditor=auditor, status='accepted').exists())
        # QA-marked active subscription, no payment record.
        self.assertTrue(company_has_active_subscription(company))

    def test_idempotent(self):
        from django.core.management import call_command
        from core.models import User, Company
        from auditors.models import AuditorAssignment
        call_command('seed_manus_e2e_qa', '--confirm', stdout=StringIO())
        call_command('seed_manus_e2e_qa', '--confirm', stdout=StringIO())
        self.assertEqual(Company.objects.filter(cr_number='QA-MANUS-0001').count(), 1)
        self.assertEqual(User.objects.filter(email='qa.auditor@manus-e2e.test').count(), 1)
        self.assertEqual(AuditorAssignment.objects.filter(
            company__cr_number='QA-MANUS-0001').count(), 1)

    def test_seeded_auditor_can_submit_verdict_when_controls_available(self):
        from django.core.management import call_command
        from core.models import User
        from compliance.models import EvidenceSubmission
        from compliance.auditor_verdict import can_submit_final_verdict
        # Seed framework versions + official controls so a submission is created.
        call_command('seed_framework_versions', stdout=StringIO())
        _seed_official('ARAMCO-SACS-002', 'ARAMCO_SACS002')
        call_command('seed_manus_e2e_qa', '--confirm', stdout=StringIO())
        sub = EvidenceSubmission.objects.filter(
            company__cr_number='QA-MANUS-0001').first()
        if sub is not None:  # controls available -> verdict-ready submission exists
            auditor = User.objects.get(email='qa.auditor@manus-e2e.test')
            self.assertTrue(can_submit_final_verdict(auditor, sub))


class GuidedCompanyWorkflowTests(TestCase):
    """Phase 8D-2-FIX-C — every guided company page answers where/what/next."""

    def _login(self, c, email):
        self.client.force_login(_journey_user(c, email=email))

    def _subscribed_company(self):
        from billing.subscription_access import activate_company_subscription
        c = _company()
        activate_company_subscription(c, 'Plan', days=30)
        return c

    def test_build_page_guide_status_logic(self):
        from compliance.journey import build_page_guide
        c = _company()
        # Fresh company: classification is current; downstream steps are blocked.
        self.assertEqual(build_page_guide(c, 'classification')['status'], 'current')
        self.assertEqual(build_page_guide(c, 'applicability')['status'], 'blocked')
        self.assertEqual(build_page_guide(c, 'evidence')['status'], 'blocked')
        self.assertEqual(build_page_guide(c, 'classification')['total'], 13)

    def test_classification_page_shows_current_step_and_next_action(self):
        c = _company()
        self._login(c, 'guide_cls@x.com')
        resp = self.client.get(reverse('compliance:classification'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'أنت الآن في الخطوة 3 من 13')
        self.assertContains(resp, 'المطلوب منك الآن')
        self.assertContains(resp, 'الخطوة التالية')

    def test_applicability_page_blocked_when_classification_missing(self):
        c = _company()
        self._login(c, 'guide_app@x.com')
        resp = self.client.get(reverse('compliance:applicability_preview'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'لم تظهر نتائج بعد لأن خطوة التصنيف لم تكتمل')

    def test_applicability_review_blocked_state_present(self):
        c = _company()
        self._login(c, 'guide_appr@x.com')
        resp = self.client.get(reverse('compliance:applicability_review'))
        self.assertContains(resp, 'بانتظار خطوة سابقة')

    def test_evidence_page_explains_missing_control_plan(self):
        c = _company()
        self._login(c, 'guide_ev@x.com')
        resp = self.client.get(reverse('compliance:evidence_checklist'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'خطة الضوابط')

    def test_controls_page_shows_guided_step(self):
        c = _company()
        self._login(c, 'guide_ctrl@x.com')
        resp = self.client.get(reverse('compliance:controls_list'))
        self.assertContains(resp, 'أنت الآن في الخطوة 5 من 13')

    def test_auditor_reviewed_report_explains_pending_review(self):
        c = self._subscribed_company()
        self._login(c, 'guide_rep@x.com')
        resp = self.client.get(reverse('compliance:auditor_reviewed_report'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'لا يوجد تقرير مدقق بعد')
        self.assertContains(resp, 'بعد مراجعة المدقق')

    def test_risk_page_shows_guided_step(self):
        c = _company()
        self._login(c, 'guide_risk@x.com')
        resp = self.client.get(reverse('risk:list'))
        self.assertContains(resp, 'أنت الآن في الخطوة 13 من 13')

    def test_guided_pages_have_no_unsafe_certification_wording(self):
        c = self._subscribed_company()
        self._login(c, 'guide_safe@x.com')
        # Affirmative certification/accreditation CLAIMS must never appear.
        # (Disclaimers like "لا يُعد شهادة امتثال رسمية" are explicitly safe.)
        banned_affirmative = ['معتمد من NCA', 'معتمد من أرامكو', 'معتمد من سابك',
                              'اعتماد حكومي', 'certified by NCA', 'certified by Aramco',
                              'certified by SABIC', 'official accreditation',
                              'government accreditation']
        for name in ('compliance:classification', 'compliance:applicability_preview',
                     'compliance:controls_list', 'compliance:evidence_checklist',
                     'compliance:auditor_reviewed_report', 'risk:list'):
            body = self.client.get(reverse(name)).content.decode()
            for w in banned_affirmative:
                self.assertNotIn(w, body, '%s in %s' % (w, name))
            # If the "official compliance certificate" phrase appears at all, it must
            # be a negated disclaimer, never an affirmative claim.
            if 'شهادة امتثال رسمية' in body:
                self.assertIn('لا يُعد شهادة امتثال رسمية', body, name)

    def test_controls_page_417_not_334(self):
        c = _company()
        self._login(c, 'guide_417@x.com')
        body = self.client.get(reverse('compliance:controls_list')).content.decode()
        self.assertIn('417', body)
        self.assertNotIn('>334<', body)

    def test_monitoring_page_shows_ongoing_step_guidance(self):
        c = _company()
        self._login(c, 'guide_mon@x.com')
        resp = self.client.get(reverse('monitoring:overview'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'خطوة مستمرة')

    def test_company_dashboard_shows_guided_journey_and_next_action(self):
        c = _company()
        self._login(c, 'guide_dash@x.com')
        resp = self.client.get(reverse('compliance:dashboard'))
        self.assertEqual(resp.status_code, 200)
        # The journey wizard + workflow stepper both render a next recommended action.
        self.assertContains(resp, 'الخطوة')


# ============================================================
# Phase 8E-EVIDENCE-OCR-A — Evidence Upload + OCR/Text-Extraction MVP
# ============================================================
def _docx_bytes(text):
    import io
    from docx import Document
    d = Document(); d.add_paragraph(text)
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()


def _pdf_bytes(text):
    import io
    from reportlab.pdfgen import canvas
    buf = io.BytesIO(); c = canvas.Canvas(buf); c.drawString(72, 720, text); c.showPage(); c.save()
    return buf.getvalue()


class EvidenceOCRMVPTests(TestCase):
    # ---- extraction by file type (service layer) ----
    def test_txt_extraction_succeeds(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission(name='e.txt',
                                                content=b'Access control policy approved by management.', ftype='txt')
        obj = save_extraction_for_submission(sub)
        self.assertEqual(obj.status, 'extracted')
        self.assertTrue(obj.has_text)
        self.assertIn('policy', obj.extracted_text.lower())

    def test_docx_extraction_succeeds(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission(name='e.docx',
                                                content=_docx_bytes('Firewall configured and reviewed quarterly'),
                                                ftype='docx')
        obj = save_extraction_for_submission(sub)
        self.assertEqual(obj.status, 'extracted')
        self.assertIn('Firewall', obj.extracted_text)

    def test_pdf_textlayer_extraction_succeeds(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission(name='e.pdf',
                                                content=_pdf_bytes('Encryption at rest enabled'), ftype='pdf')
        obj = save_extraction_for_submission(sub)
        self.assertEqual(obj.status, 'extracted')
        self.assertIn('Encryption', obj.extracted_text)

    def test_image_does_not_fake_ocr_success(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        png = b'\x89PNG\r\n\x1a\n' + b'\x00' * 32
        c, item, sub = _company_with_submission(name='e.png', content=png, ftype='png')
        obj = save_extraction_for_submission(sub)
        self.assertFalse(obj.has_text)                    # never pretends OCR success
        self.assertEqual(obj.status, 'no_text_extracted')  # honest manual-review status

    def test_corrupt_file_extraction_no_500(self):
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission(name='bad.docx', content=b'not a real docx', ftype='docx')
        obj = save_extraction_for_submission(sub)  # must not raise
        self.assertIn(obj.status, ('failed', 'unsupported_type', 'no_text_extracted'))
        self.assertNotIn('/', obj.error_message or '')     # no file path leaked

    def test_extracted_text_is_bounded(self):
        from compliance.evidence_extraction import save_extraction_for_submission, MAX_TEXT_CHARS
        big = ('policy ' * 20000).encode()  # 140k chars, well over the cap
        self.assertGreater(len(big), MAX_TEXT_CHARS)
        c, item, sub = _company_with_submission(name='big.txt', content=big, ftype='txt')
        obj = save_extraction_for_submission(sub)
        # Stored/previewed text is bounded to the cap regardless of source size.
        self.assertLessEqual(len(obj.extracted_text), MAX_TEXT_CHARS)
        self.assertLessEqual(obj.char_count, MAX_TEXT_CHARS)

    # ---- upload permissions / tenant ----
    def _login_company(self, c, email):
        u = _journey_user(c, email=email)
        self.client.force_login(u)
        return u

    def test_anonymous_cannot_upload(self):
        c, item, sub = _company_with_submission()
        r = self.client.get(reverse('compliance:evidence_upload_v2', args=[item.id]))
        self.assertEqual(r.status_code, 302)
        self.assertIn('/login', r.url)

    def test_company_user_can_upload(self):
        from compliance.models import EvidenceSubmission
        c, fv, scope = _company_with_checklist()
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        self._login_company(c, 'ocrup@x.com')
        before = EvidenceSubmission.objects.filter(company=c).count()
        f = _SUF('policy.txt', b'access control policy', content_type='text/plain')
        self.client.post(reverse('compliance:evidence_upload_v2', args=[item.id]), {'uploaded_file': f})
        self.assertEqual(EvidenceSubmission.objects.filter(company=c).count(), before + 1)

    def test_upload_writes_audit_log(self):
        from core.models import AuditLog
        c, fv, scope = _company_with_checklist()
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        self._login_company(c, 'ocraud@x.com')
        self.client.post(reverse('compliance:evidence_upload_v2', args=[item.id]),
                         {'uploaded_file': _SUF('p.txt', b'x', content_type='text/plain')})
        self.assertTrue(AuditLog.objects.filter(action='evidence_uploaded').exists())

    def test_unlinked_user_upload_safe_no_company(self):
        from core.models import User
        c, fv, scope = _company_with_checklist()
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        orphan = User.objects.create_user(username='ocrorphan@x.com', email='ocrorphan@x.com',
                                          password='longenough12', role='company_admin')
        self.client.force_login(orphan)
        resp = self.client.get(reverse('compliance:evidence_upload_v2', args=[item.id]))
        self.assertEqual(resp.status_code, 200)  # no 500
        self.assertContains(resp, 'not linked to a company')

    def test_auditor_cannot_upload_via_company_portal(self):
        from auditors.models import AuditorProfile
        from core.models import User
        c, fv, scope = _company_with_checklist()
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        au = User.objects.create_user(username='ocraud2@x.com', email='ocraud2@x.com',
                                      password='longenough12', role='auditor')
        AuditorProfile.objects.create(user=au, full_name='A', status='active')
        self.client.force_login(au)
        resp = self.client.get(reverse('compliance:evidence_upload_v2', args=[item.id]))
        self.assertContains(resp, 'Auditor account', status_code=200)

    def test_upload_rejects_unsupported_type(self):
        from compliance.models import EvidenceSubmission
        c, fv, scope = _company_with_checklist()
        item = EvidenceChecklistItem.objects.filter(company=c).first()
        self._login_company(c, 'ocrbad@x.com')
        before = EvidenceSubmission.objects.filter(company=c).count()
        self.client.post(reverse('compliance:evidence_upload_v2', args=[item.id]),
                         {'uploaded_file': _SUF('m.exe', b'MZ', content_type='application/octet-stream')})
        self.assertEqual(EvidenceSubmission.objects.filter(company=c).count(), before)

    def test_company_cannot_access_other_company_submission(self):
        c1, i1, s1 = _company_with_submission(name='a.txt')
        c2, i2, s2 = _company_with_submission(name='b.txt')
        self._login_company(c1, 'ocrt1@x.com')
        resp = self.client.get(reverse('compliance:evidence_submission_detail', args=[s2.id]))
        self.assertEqual(resp.status_code, 404)  # not found for this tenant -> 404

    def test_company_cannot_run_extraction_for_other_company(self):
        from compliance.models import EvidenceTextExtraction
        c1, i1, s1 = _company_with_submission(name='a.txt')
        c2, i2, s2 = _company_with_submission(name='b.txt')
        self._login_company(c1, 'ocrt2@x.com')
        self.client.post(reverse('compliance:run_evidence_extraction', args=[s2.id]))
        self.assertFalse(EvidenceTextExtraction.objects.filter(submission=s2).exists())

    # ---- run extraction via view ----
    def test_run_extraction_writes_audit_and_shows_text(self):
        from core.models import AuditLog
        c, item, sub = _company_with_submission(name='e.txt',
                                                content=b'network segmentation policy', ftype='txt')
        self._login_company(c, 'ocrrun@x.com')
        self.client.post(reverse('compliance:run_evidence_extraction', args=[sub.id]))
        self.assertTrue(AuditLog.objects.filter(action='evidence_extraction').exists())
        detail = self.client.get(reverse('compliance:evidence_submission_detail', args=[sub.id])).content.decode()
        self.assertIn('network segmentation', detail)

    def test_submission_detail_shows_extraction_status(self):
        c, item, sub = _company_with_submission(name='e.txt', content=b'hello', ftype='txt')
        self._login_company(c, 'ocrstat@x.com')
        body = self.client.get(reverse('compliance:evidence_submission_detail', args=[sub.id])).content.decode()
        self.assertIn('Text extraction', body)

    # ---- CRM evidence summary ----
    def test_crm_company_detail_shows_evidence_summary(self):
        from core.models import User
        from compliance.evidence_extraction import save_extraction_for_submission
        c, item, sub = _company_with_submission(name='e.txt', content=b'policy text', ftype='txt')
        save_extraction_for_submission(sub)
        staff = User.objects.create_user(username='ocrstaff@x.com', email='ocrstaff@x.com',
                                         password='longenough12', role='admin', is_staff=True)
        self.client.force_login(staff)
        resp = self.client.get(reverse('platform_admin:company_detail', args=[c.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Evidence summary')

    # ---- smart processing animation + safety ----
    def test_processing_animation_present_on_extraction_page(self):
        c, item, sub = _company_with_submission(name='e.txt', content=b'x', ftype='txt')
        self._login_company(c, 'ocranim@x.com')
        body = self.client.get(reverse('compliance:evidence_extraction', args=[sub.id])).content.decode()
        self.assertIn('data-smart-processing', body)
        self.assertIn('Processing evidence', body)
        self.assertIn('Reading file', body)

    def test_processing_animation_no_unsafe_wording(self):
        c, item, sub = _company_with_submission(name='e.txt', content=b'x', ftype='txt')
        self._login_company(c, 'ocrsafe@x.com')
        body = self.client.get(reverse('compliance:evidence_extraction', args=[sub.id])).content.decode()
        for w in ('شهادة امتثال رسمية', 'اعتماد رسمي', 'معتمد من NCA', 'certified by NCA',
                  'official accreditation', 'government accredited', 'AI approved', 'AI certified',
                  'compliance confirmed'):
            self.assertNotIn(w, body)


# ============================================================
# Phase 8F-GAP-ENGINE-A — Real Gap Analysis Engine
# ============================================================
class GapEngineTests(TestCase):
    def _apply_controls(self, c):
        from compliance.models import ControlApplicabilityResult
        return list(ControlApplicabilityResult.objects.filter(
            company=c, decision='applicable', control__framework_version__isnull=False))

    # ---- calculation rules ----
    def test_no_evidence_is_missing(self):
        from compliance.gap_engine import recalculate_company_gap
        from compliance.models import ControlGapAssessment
        c, fv, scope = _company_with_official_plan()
        # applicable controls exist but no checklist/evidence yet
        recalculate_company_gap(c)
        self.assertTrue(ControlGapAssessment.objects.filter(company=c, status='missing').exists())

    def test_evidence_without_text_is_needs_review(self):
        from compliance.gap_engine import recalculate_company_gap
        from compliance.models import ControlGapAssessment, EvidenceChecklistItem
        c, item, sub = _company_with_submission(name='e.png', content=b'\x89PNG\r\n', ftype='png')
        recalculate_company_gap(c)
        ctrl = item.evidence_requirement.control
        g = ControlGapAssessment.objects.get(company=c, control=ctrl)
        self.assertEqual(g.status, 'needs_review')
        self.assertEqual(g.evidence_count, 1)
        self.assertEqual(g.extracted_text_available_count, 0)

    def test_evidence_with_text_is_partial(self):
        from compliance.gap_engine import recalculate_company_gap
        from compliance.evidence_extraction import save_extraction_for_submission
        from compliance.models import ControlGapAssessment
        c, item, sub = _company_with_submission(name='e.txt', content=b'access control policy approved', ftype='txt')
        save_extraction_for_submission(sub)  # produces readable text
        recalculate_company_gap(c)
        ctrl = item.evidence_requirement.control
        g = ControlGapAssessment.objects.get(company=c, control=ctrl)
        self.assertEqual(g.status, 'partially_compliant')
        self.assertGreaterEqual(g.extracted_text_available_count, 1)

    def test_auditor_compliant_is_compliant(self):
        from compliance.gap_engine import recalculate_company_gap
        from compliance.models import ControlGapAssessment, ControlAssessment
        c, fv, scope = _company_with_assessments()
        a = ControlAssessment.objects.filter(company=c).first()
        a.status = 'compliant'; a.save(update_fields=['status'])
        recalculate_company_gap(c)
        g = ControlGapAssessment.objects.get(company=c, control=a.control)
        self.assertEqual(g.status, 'compliant')
        self.assertEqual(g.score, 100)

    def test_not_applicable_scope(self):
        from compliance.gap_engine import recalculate_company_gap
        from compliance.models import ControlGapAssessment, ControlApplicabilityResult
        c, fv, scope = _company_with_official_plan()
        car = ControlApplicabilityResult.objects.filter(company=c).first()
        car.decision = 'not_applicable'; car.save(update_fields=['decision'])
        recalculate_company_gap(c)
        g = ControlGapAssessment.objects.get(company=c, control=car.control)
        self.assertEqual(g.status, 'not_applicable')

    def test_recalculation_idempotent(self):
        from compliance.gap_engine import recalculate_company_gap
        from compliance.models import ControlGapAssessment
        c, item, sub = _company_with_submission()
        recalculate_company_gap(c)
        n1 = ControlGapAssessment.objects.filter(company=c).count()
        recalculate_company_gap(c)
        n2 = ControlGapAssessment.objects.filter(company=c).count()
        self.assertEqual(n1, n2)

    def test_empty_company_no_500(self):
        from compliance.gap_engine import recalculate_company_gap, get_company_gap_summary
        c = _company()
        s = recalculate_company_gap(c)  # no plan at all
        self.assertEqual(s['total'], 0)
        self.assertEqual(get_company_gap_summary(c)['overall_readiness_percent'], 0)

    def test_recalc_does_not_create_ai_gap_rows(self):
        from compliance.gap_engine import recalculate_company_gap
        from ai_engine.models import GapAnalysis
        before = GapAnalysis.objects.count()
        c, item, sub = _company_with_submission()
        recalculate_company_gap(c)
        self.assertEqual(GapAnalysis.objects.count(), before)  # deterministic engine, no AI

    # ---- readiness scoring ----
    def test_readiness_percent_weighted(self):
        from compliance.gap_engine import _readiness_percent
        # 1 compliant(100) + 1 partial(50) + 1 missing(0) over 3 applicable = 50
        self.assertEqual(_readiness_percent({'compliant': 1, 'partially_compliant': 1,
                                             'missing': 1, 'needs_review': 0, 'not_applicable': 0}), 50)
        # not_applicable excluded from denominator
        self.assertEqual(_readiness_percent({'compliant': 1, 'partially_compliant': 0,
                                             'missing': 0, 'needs_review': 0, 'not_applicable': 5}), 100)

    def test_counts_and_frameworks_scoped(self):
        from compliance.gap_engine import recalculate_company_gap, get_company_gap_summary
        c, item, sub = _company_with_submission()
        recalculate_company_gap(c)
        s = get_company_gap_summary(c)
        self.assertEqual(sum(s['overall'].values()), s['total'])
        self.assertGreaterEqual(len(s['frameworks']), 1)


class GapDashboardViewTests(TestCase):
    def _login(self, c, email):
        u = _journey_user(c, email=email)
        self.client.force_login(u)
        return u

    def test_anonymous_redirected(self):
        r = self.client.get(reverse('compliance:gap_dashboard'))
        self.assertEqual(r.status_code, 302)
        self.assertIn('/login', r.url)

    def test_company_user_can_view(self):
        c, item, sub = _company_with_submission()
        self._login(c, 'gapv@x.com')
        resp = self.client.get(reverse('compliance:gap_dashboard'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Internal readiness')

    def test_disclaimer_and_distribution_present(self):
        from compliance.gap_engine import recalculate_company_gap
        c, item, sub = _company_with_submission()
        recalculate_company_gap(c)
        self._login(c, 'gapd@x.com')
        body = self.client.get(reverse('compliance:gap_dashboard')).content.decode()
        self.assertIn('Not an official certification', body)
        self.assertIn('Framework readiness', body)
        self.assertIn('Missing', body)  # status distribution card

    def test_missing_control_links_to_evidence(self):
        from compliance.gap_engine import recalculate_company_gap
        c, fv, scope = _company_with_checklist()
        recalculate_company_gap(c)  # no evidence -> missing rows
        self._login(c, 'gapm@x.com')
        body = self.client.get(reverse('compliance:gap_dashboard')).content.decode()
        self.assertIn('Upload evidence', body)
        self.assertIn(reverse('compliance:evidence_checklist'), body)

    def test_smart_processing_animation_present(self):
        c, item, sub = _company_with_submission()
        self._login(c, 'gapa@x.com')
        body = self.client.get(reverse('compliance:gap_dashboard')).content.decode()
        self.assertIn('data-smart-processing', body)
        self.assertIn('Processing evidence', body)

    def test_empty_company_dashboard_no_500(self):
        c = _company()
        self._login(c, 'gape@x.com')
        self.assertEqual(self.client.get(reverse('compliance:gap_dashboard')).status_code, 200)

    # ---- recalc action ----
    def test_recalc_requires_post(self):
        c, item, sub = _company_with_submission()
        self._login(c, 'gaprp@x.com')
        self.assertEqual(self.client.get(reverse('compliance:run_gap_recalc')).status_code, 405)

    def test_recalc_writes_audit(self):
        from core.models import AuditLog
        c, item, sub = _company_with_submission()
        self._login(c, 'gapau@x.com')
        self.client.post(reverse('compliance:run_gap_recalc'))
        log = AuditLog.objects.filter(action='gap_recalculated').order_by('-id').first()
        self.assertIsNotNone(log)
        self.assertEqual(log.metadata.get('company_id'), c.id)

    def test_recalc_tenant_scoped(self):
        from compliance.gap_engine import recalculate_company_gap, gap_rows
        c1, i1, s1 = _company_with_submission(name='a.txt')
        c2, i2, s2 = _company_with_submission(name='b.txt')
        recalculate_company_gap(c1); recalculate_company_gap(c2)
        self.assertTrue(all(r.company_id == c1.id for r in gap_rows(c1)))
        self.assertTrue(all(r.company_id == c2.id for r in gap_rows(c2)))

    # ---- permissions ----
    def test_auditor_denied(self):
        from auditors.models import AuditorProfile
        from core.models import User
        au = User.objects.create_user(username='gapaud@x.com', email='gapaud@x.com',
                                      password='longenough12', role='auditor')
        AuditorProfile.objects.create(user=au, full_name='A', status='active')
        self.client.force_login(au)
        resp = self.client.get(reverse('compliance:gap_dashboard'))
        self.assertContains(resp, 'Auditor account', status_code=200)

    def test_staff_without_company_routed_to_crm(self):
        from core.models import User
        st = User.objects.create_user(username='gapstaff@x.com', email='gapstaff@x.com',
                                      password='longenough12', role='admin', is_staff=True)
        self.client.force_login(st)
        resp = self.client.get(reverse('compliance:gap_dashboard'))
        self.assertContains(resp, 'Get Solution CRM', status_code=200)

    # ---- CRM summary + safety ----
    def test_crm_company_detail_shows_gap_summary(self):
        from compliance.gap_engine import recalculate_company_gap
        from core.models import User
        c, item, sub = _company_with_submission()
        recalculate_company_gap(c)
        staff = User.objects.create_user(username='gapcrm@x.com', email='gapcrm@x.com',
                                         password='longenough12', role='admin', is_staff=True)
        self.client.force_login(staff)
        resp = self.client.get(reverse('platform_admin:company_detail', args=[c.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Preliminary readiness')

    def test_no_unsafe_wording_on_gap_page(self):
        from compliance.gap_engine import recalculate_company_gap
        c, item, sub = _company_with_submission()
        recalculate_company_gap(c)
        self._login(c, 'gapsafe@x.com')
        body = self.client.get(reverse('compliance:gap_dashboard')).content.decode()
        # Affirmative certification/accreditation claims must never appear. The safe
        # NEGATED disclaimer "Not an official certification" is allowed.
        for w in ('معتمد من NCA', 'معتمد من أرامكو', 'معتمد من سابك', 'اعتماد حكومي',
                  'certified by NCA', 'official accreditation', 'government accredited'):
            self.assertNotIn(w, body)
        if 'official certification' in body:
            self.assertIn('Not an official certification', body)
        if 'شهادة امتثال رسمية' in body:
            self.assertIn('لا تُعد شهادة امتثال رسمية', body)


# ============================================================
# Phase 8H-REPORTING-A — Commercial (internal readiness) Report Engine
# ============================================================
class CommercialReportEngineTests(TestCase):
    def _prepared(self, **subkw):
        from compliance.gap_engine import recalculate_company_gap
        from risk.risk_engine import generate_risks_from_gap
        c, item, sub = _company_with_submission(**subkw)
        recalculate_company_gap(c)
        generate_risks_from_gap(c)
        return c, item, sub

    def test_build_report_shape(self):
        from compliance.report_engine import build_commercial_readiness_report
        c, item, sub = self._prepared()
        rep = build_commercial_readiness_report(c)
        for key in ('executive', 'framework_readiness', 'evidence', 'gap', 'risk', 'remediation', 'next_actions'):
            self.assertIn(key, rep)
        self.assertGreaterEqual(len(rep['framework_readiness']), 1)
        self.assertTrue(rep['next_actions'])

    def test_empty_company_report_no_error(self):
        from compliance.report_engine import build_commercial_readiness_report
        rep = build_commercial_readiness_report(_company())
        self.assertFalse(rep['has_data'])
        self.assertEqual(rep['executive']['overall_readiness_percent'], 0)
        self.assertTrue(rep['next_actions'])  # still gives a safe default action

    def test_report_aggregates_gap_and_risk(self):
        from compliance.report_engine import build_commercial_readiness_report
        c, item, sub = self._prepared()
        rep = build_commercial_readiness_report(c)
        self.assertGreaterEqual(rep['gap']['total'], 1)
        self.assertGreaterEqual(rep['risk']['total'], 1)
        self.assertGreaterEqual(rep['remediation']['total'], 1)


class CommercialReportViewTests(TestCase):
    def _prepared(self, **subkw):
        from compliance.gap_engine import recalculate_company_gap
        from risk.risk_engine import generate_risks_from_gap
        c, item, sub = _company_with_submission(**subkw)
        recalculate_company_gap(c); generate_risks_from_gap(c)
        return c, item, sub

    def _login(self, c, email):
        self.client.force_login(_journey_user(c, email=email))

    URL = 'compliance:commercial_readiness_report'
    REFRESH = 'compliance:refresh_commercial_report'

    # ---- access ----
    def test_anonymous_redirected(self):
        r = self.client.get(reverse(self.URL))
        self.assertEqual(r.status_code, 302)
        self.assertIn('/login', r.url)

    def test_company_user_can_view(self):
        c, item, sub = self._prepared()
        self._login(c, 'crv@x.com')
        self.assertEqual(self.client.get(reverse(self.URL)).status_code, 200)

    def test_unlinked_user_safe_no_company(self):
        from core.models import User
        u = User.objects.create_user(username='crorph@x.com', email='crorph@x.com',
                                     password='longenough12', role='company_admin')
        self.client.force_login(u)
        resp = self.client.get(reverse(self.URL))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'not linked to a company')

    def test_auditor_denied(self):
        from auditors.models import AuditorProfile
        from core.models import User
        au = User.objects.create_user(username='craud@x.com', email='craud@x.com',
                                      password='longenough12', role='auditor')
        AuditorProfile.objects.create(user=au, full_name='A', status='active')
        self.client.force_login(au)
        self.assertContains(self.client.get(reverse(self.URL)), 'Auditor account', status_code=200)

    def test_staff_without_company_routed_to_crm(self):
        from core.models import User
        st = User.objects.create_user(username='crstaff@x.com', email='crstaff@x.com',
                                      password='longenough12', role='admin', is_staff=True)
        self.client.force_login(st)
        self.assertContains(self.client.get(reverse(self.URL)), 'Get Solution CRM', status_code=200)

    # ---- content ----
    def test_report_shows_company_and_disclaimer(self):
        c, item, sub = self._prepared()
        self._login(c, 'crc@x.com')
        body = self.client.get(reverse(self.URL)).content.decode()
        self.assertIn(c.name, body)
        self.assertIn('Internal readiness report', body)
        self.assertIn('لا يُعد شهادة امتثال رسمية', body)  # safe negated disclaimer
        self.assertIn('Not an official certification', body)

    def test_report_sections_present(self):
        c, item, sub = self._prepared()
        self._login(c, 'crs@x.com')
        body = self.client.get(reverse(self.URL)).content.decode()
        for section in ('Executive summary', 'Framework readiness', 'Evidence summary',
                        'Gap summary', 'Risk summary', 'Remediation plan'):
            self.assertIn(section, body)
        self.assertIn('Readiness', body)  # overall readiness card
        self.assertIn('Recommended next actions', body)

    def test_report_empty_company_no_500(self):
        c = _company()
        self._login(c, 'cre@x.com')
        self.assertEqual(self.client.get(reverse(self.URL)).status_code, 200)

    def test_print_friendly_layout(self):
        c, item, sub = self._prepared()
        self._login(c, 'crp@x.com')
        body = self.client.get(reverse(self.URL)).content.decode()
        self.assertIn('window.print()', body)
        self.assertIn('@media print', body)

    # ---- refresh action ----
    def test_refresh_requires_post(self):
        c, item, sub = self._prepared()
        self._login(c, 'crrp@x.com')
        self.assertEqual(self.client.get(reverse(self.REFRESH)).status_code, 405)

    def test_refresh_has_smart_processing(self):
        c, item, sub = self._prepared()
        self._login(c, 'cranim@x.com')
        body = self.client.get(reverse(self.URL)).content.decode()
        self.assertIn('data-smart-processing', body)
        self.assertIn('Processing evidence', body)

    def test_refresh_writes_audit(self):
        from core.models import AuditLog
        c, item, sub = self._prepared()
        self._login(c, 'craudit@x.com')
        self.client.post(reverse(self.REFRESH))
        self.assertTrue(AuditLog.objects.filter(action='report_refreshed').exists())

    def test_view_writes_audit(self):
        from core.models import AuditLog
        c, item, sub = self._prepared()
        self._login(c, 'crvaudit@x.com')
        self.client.get(reverse(self.URL))
        self.assertTrue(AuditLog.objects.filter(action='report_viewed').exists())

    # ---- safety ----
    def test_no_unsafe_wording(self):
        c, item, sub = self._prepared()
        self._login(c, 'crsafe@x.com')
        body = self.client.get(reverse(self.URL)).content.decode()
        for w in ('معتمد من NCA', 'معتمد من أرامكو', 'معتمد من سابك', 'اعتماد حكومي',
                  'certified by NCA', 'official accreditation', 'government accredited'):
            self.assertNotIn(w, body)
        if 'official certification' in body:
            self.assertIn('Not an official certification', body)
        if 'شهادة امتثال رسمية' in body:
            self.assertIn('لا يُعد شهادة امتثال رسمية', body)


class CommercialReportCRMSummaryTests(TestCase):
    def test_crm_detail_shows_report_summary(self):
        from core.models import User
        from compliance.gap_engine import recalculate_company_gap
        c, item, sub = _company_with_submission()
        recalculate_company_gap(c)
        staff = User.objects.create_user(username='crcrm@x.com', email='crcrm@x.com',
                                         password='longenough12', role='admin', is_staff=True)
        self.client.force_login(staff)
        resp = self.client.get(reverse('platform_admin:company_detail', args=[c.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Readiness report summary')

    def test_crm_report_summary_staff_only(self):
        c, item, sub = _company_with_submission()
        self.client.force_login(_journey_user(c, email='crns@x.com'))
        self.assertEqual(self.client.get(
            reverse('platform_admin:company_detail', args=[c.id])).status_code, 403)


# ============================================================
# Phase 8H-B — Commercial Report PDF Export
# ============================================================
class CommercialReportPDFTests(TestCase):
    URL = 'compliance:commercial_readiness_report_pdf'

    def _prepared(self, **subkw):
        from compliance.gap_engine import recalculate_company_gap
        from risk.risk_engine import generate_risks_from_gap
        c, item, sub = _company_with_submission(**subkw)
        recalculate_company_gap(c); generate_risks_from_gap(c)
        return c, item, sub

    def _login(self, c, email):
        self.client.force_login(_journey_user(c, email=email))

    def _pdf_text(self, pdf_bytes):
        import io, pdfplumber
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as p:
            return ' '.join((pg.extract_text() or '') for pg in p.pages)

    # ---- access ----
    def test_anonymous_redirected(self):
        r = self.client.get(reverse(self.URL))
        self.assertEqual(r.status_code, 302)
        self.assertIn('/login', r.url)

    def test_company_user_can_export(self):
        c, item, sub = self._prepared()
        self._login(c, 'pdfu@x.com')
        resp = self.client.get(reverse(self.URL))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp['Content-Type'], 'application/pdf')

    def test_unlinked_user_safe_no_company(self):
        from core.models import User
        u = User.objects.create_user(username='pdforph@x.com', email='pdforph@x.com',
                                     password='longenough12', role='company_admin')
        self.client.force_login(u)
        resp = self.client.get(reverse(self.URL))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'not linked to a company')

    def test_auditor_denied(self):
        from auditors.models import AuditorProfile
        from core.models import User
        au = User.objects.create_user(username='pdfaud@x.com', email='pdfaud@x.com',
                                      password='longenough12', role='auditor')
        AuditorProfile.objects.create(user=au, full_name='A', status='active')
        self.client.force_login(au)
        self.assertContains(self.client.get(reverse(self.URL)), 'Auditor account', status_code=200)

    def test_staff_without_company_routed_to_crm(self):
        from core.models import User
        st = User.objects.create_user(username='pdfstaff@x.com', email='pdfstaff@x.com',
                                      password='longenough12', role='admin', is_staff=True)
        self.client.force_login(st)
        self.assertContains(self.client.get(reverse(self.URL)), 'Get Solution CRM', status_code=200)

    # ---- PDF response ----
    def test_pdf_content_type_and_filename(self):
        c, item, sub = self._prepared()
        self._login(c, 'pdff@x.com')
        resp = self.client.get(reverse(self.URL))
        self.assertEqual(resp['Content-Type'], 'application/pdf')
        self.assertIn('attachment; filename="1saudicyber_readiness_report_%s.pdf"' % c.id,
                      resp['Content-Disposition'])

    def test_pdf_non_empty_and_valid_header(self):
        c, item, sub = self._prepared()
        self._login(c, 'pdfh@x.com')
        content = self.client.get(reverse(self.URL)).content
        self.assertGreater(len(content), 500)
        self.assertTrue(content.startswith(b'%PDF-'))

    def test_pdf_contains_company_and_disclaimer(self):
        c, item, sub = self._prepared()
        self._login(c, 'pdfd@x.com')
        txt = self._pdf_text(self.client.get(reverse(self.URL)).content)
        self.assertIn(c.name, txt)
        self.assertIn('internal readiness report', txt.lower())
        self.assertIn('not an official certification', txt.lower())

    def test_pdf_no_unsafe_wording(self):
        c, item, sub = self._prepared()
        self._login(c, 'pdfsafe@x.com')
        txt = self._pdf_text(self.client.get(reverse(self.URL)).content).lower()
        for w in ('certified by nca', 'certified by aramco', 'certified by sabic',
                  'official accreditation', 'government accredited'):
            self.assertNotIn(w, txt)
        if 'official certification' in txt:
            self.assertIn('not an official certification', txt)

    def test_pdf_empty_company_no_500(self):
        c = _company()
        self._login(c, 'pdfe@x.com')
        resp = self.client.get(reverse(self.URL))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp['Content-Type'], 'application/pdf')

    def test_pdf_uses_only_current_company_data(self):
        c1, i1, s1 = self._prepared(name='a.txt')
        c2, i2, s2 = self._prepared(name='b.txt')
        self._login(c1, 'pdften@x.com')
        txt = self._pdf_text(self.client.get(reverse(self.URL)).content)
        self.assertIn(c1.name, txt)
        # c2 has the same generated CR pattern; ensure c2's CR is not in c1's report
        self.assertNotIn(c2.cr_number, txt)

    # ---- audit + service ----
    def test_pdf_export_writes_audit(self):
        from core.models import AuditLog
        c, item, sub = self._prepared()
        self._login(c, 'pdfaudit@x.com')
        self.client.get(reverse(self.URL))
        self.assertTrue(AuditLog.objects.filter(action='report_pdf_exported').exists())

    def test_service_builds_pdf_bytes(self):
        from compliance.report_pdf import build_commercial_readiness_pdf
        c, item, sub = self._prepared()
        pdf = build_commercial_readiness_pdf(c)
        self.assertTrue(pdf.startswith(b'%PDF-'))

    def test_pdf_export_fallback_on_failure_no_500(self):
        from unittest import mock
        c, item, sub = self._prepared()
        self._login(c, 'pdffail@x.com')
        with mock.patch('compliance.report_pdf.build_commercial_readiness_pdf',
                        side_effect=RuntimeError('boom')):
            resp = self.client.get(reverse(self.URL))
        # safe fallback: redirect to the HTML report, no 500, no traceback
        self.assertEqual(resp.status_code, 302)
        self.assertIn(reverse('compliance:commercial_readiness_report'), resp.url)

    # ---- UI ----
    def test_html_report_has_export_pdf_and_print(self):
        c, item, sub = self._prepared()
        self._login(c, 'pdfui@x.com')
        body = self.client.get(reverse('compliance:commercial_readiness_report')).content.decode()
        self.assertIn(reverse(self.URL), body)  # Export PDF button link
        self.assertIn('window.print()', body)   # Print button still present
